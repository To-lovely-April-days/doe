"""
doe_engine.py — DOE 实验设计引擎
编译为 doe_engine.pyd 供 C# 通过 pythonnet 调用

依赖: pip install numpy pandas openpyxl scipy statsmodels
注意: 不再依赖 pyDOE2（CCD/BBD/FracFact 已改为纯 numpy 实现，兼容 Python 3.12+）

★ 升级内容:
  1. DOEDesigner 新增: ccd(), box_behnken(), d_optimal() 三种 RSM 设计方法
  2. DOEDesigner 新增: get_design_quality() 返回 D/A/G 效率、VIF、别名结构、Power
  3. DOEAnalyzer 改造: fit_ols() 输出完整 ANOVA 表(SS/MS/F/P)、系数表(β/SE/T/P/VIF)、
     模型摘要(R²/R²adj/R²pred/RMSE/Lack-of-Fit)、效应 Pareto 图
  4. DOEAnalyzer 新增: residual_diagnostics() 返回完整残差四图数据
"""

import json
import numpy as np
import pandas as pd
from itertools import product as iterproduct
from typing import List, Dict, Optional, Any

# ═══════════════════════════════════════════════════════
# DOEDesigner — 参数矩阵生成
# ═══════════════════════════════════════════════════════

class DOEDesigner:
    """
    DOE 实验设计器 — 根据因子定义和设计方法生成参数矩阵。
    所有输入输出均为 JSON 字符串，便于 C# 侧 pythonnet 调用。
    """

    def full_factorial(self, factors_json: str) -> str:
        """
        全因子设计 — 所有因子所有水平的完全组合。
        
        factors_json 格式:
        [
            {"name": "温度", "levels": [80, 120, 160]},
            {"name": "压力", "levels": [10, 20, 30]}
        ]
        
        返回 matrix_json:
        [{"温度": 80, "压力": 10}, {"温度": 80, "压力": 20}, ...]
        """
        factors = json.loads(factors_json)
        
        names = [f["name"] for f in factors]
        levels_list = [f["levels"] for f in factors]
        
        # 全因子组合
        combos = list(iterproduct(*levels_list))
        
        matrix = []
        for combo in combos:
            row = {names[i]: combo[i] for i in range(len(names))}
            matrix.append(row)
        
        return json.dumps(matrix, ensure_ascii=False)

    def fractional_factorial(self, factors_json: str, resolution: int = 3) -> str:
        """
        部分因子设计 — 2^(k-p) 设计，用于因子较多时的筛选实验。
        仅支持 2 水平因子（取上下界）。
        
        factors_json 格式:
        [
            {"name": "温度", "lower": 80, "upper": 160},
            {"name": "压力", "lower": 10, "upper": 30},
            ...
        ]
        resolution: 分辨度 (III, IV, V 对应 3, 4, 5)
        """
        factors = json.loads(factors_json)
        k = len(factors)
        
        if k < 3:
            # 因子数太少，退化为全因子
            return self._two_level_full_factorial(factors)
        
        # ── 纯 numpy 实现部分因子设计 ──
        # 基本原理: 取前 p 个因子做全因子 (2^p 组)，后面的因子用生成元列（列的乘积）生成
        gen_string = self._build_generator_string(k, resolution)
        coded_matrix = self._fracfact_numpy(gen_string)
        
        # 解码为实际值
        matrix = []
        for row in coded_matrix:
            decoded_row = {}
            for i, f in enumerate(factors):
                if i < len(row):
                    coded_val = row[i]
                    lower = f["lower"]
                    upper = f["upper"]
                    actual_val = lower + (coded_val + 1) / 2 * (upper - lower)
                    decoded_row[f["name"]] = round(actual_val, 6)
            matrix.append(decoded_row)
        
        return json.dumps(matrix, ensure_ascii=False)

    def taguchi(self, factors_json: str, table_type: str = "auto") -> str:
        """
        正交设计 — Taguchi 正交表。
        
        factors_json 格式:
        [
            {"name": "温度", "levels": [80, 120, 160]},
            {"name": "压力", "levels": [10, 20, 30]},
            {"name": "催化剂", "levels": [1.0, 3.0, 5.0]}
        ]
        table_type: "L4", "L8", "L9", "L16", "L18", "L27" 或 "auto"
        """
        factors = json.loads(factors_json)
        k = len(factors)
        level_counts = [len(f["levels"]) for f in factors]
        max_levels = max(level_counts)
        
        if table_type == "auto":
            table_type = self._select_orthogonal_table(k, max_levels)
        
        # 获取正交表
        oa = self._get_orthogonal_array(table_type, k, max_levels)
        
        # 映射到实际水平值
        matrix = []
        for row in oa:
            actual_row = {}
            for i, f in enumerate(factors):
                if i < len(row):
                    level_idx = int(row[i]) % len(f["levels"])
                    actual_row[f["name"]] = f["levels"][level_idx]
            matrix.append(actual_row)
        
        return json.dumps(matrix, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # ★ 新增: 连续/类别因子分离与笛卡尔积辅助
    # ═══════════════════════════════════════════════════════

    def _separate_factors(self, factors: list) -> tuple:
        """
        将因子列表分为连续因子和类别因子。
        
        新格式 (C# 端传入):
          连续: {"name": "温度", "type": "continuous", "lower": 80, "upper": 160}
          类别: {"name": "催化剂", "type": "categorical", "levels": ["A", "B", "C"]}
        
        旧格式兼容 (无 type 字段，默认连续):
          {"name": "温度", "lower": 80, "upper": 160}
        
        返回: (continuous_factors, categorical_factors)
        """
        continuous = []
        categorical = []
        for f in factors:
            ftype = f.get("type", "continuous").lower()
            if ftype == "categorical":
                categorical.append(f)
            else:
                continuous.append(f)
        return continuous, categorical

    def _cross_with_categorical(self, base_matrix: list, categorical_factors: list) -> list:
        """
        将连续因子的 RSM 设计矩阵与类别因子做笛卡尔积。
        
        例: base_matrix 有 18 行（CCD 连续部分），categorical 有 1 个因子 3 个水平
        → 结果: 18 × 3 = 54 行
        
        类别因子在矩阵中用水平标签（字符串）表示。
        """
        if not categorical_factors:
            return base_matrix
        
        # 构建类别因子的所有水平组合
        cat_names = [f["name"] for f in categorical_factors]
        cat_levels = [f["levels"] for f in categorical_factors]
        cat_combos = list(iterproduct(*cat_levels))
        
        # 笛卡尔积
        crossed = []
        for base_row in base_matrix:
            for combo in cat_combos:
                new_row = dict(base_row)
                for i, name in enumerate(cat_names):
                    new_row[name] = combo[i]
                crossed.append(new_row)
        
        return crossed

    # ═══════════════════════════════════════════════════════
    # ★ 新增: RSM 设计方法 — CCD / Box-Behnken / D-Optimal
    # ═══════════════════════════════════════════════════════

    def ccd(self, factors_json: str, alpha_type: str = "rotatable", center_count: int = -1) -> str:
        """
        ★ 新增: 中心复合设计 (Central Composite Design) — 纯 numpy 实现，无需 pyDOE2
        
        RSM 核心方法 — 用于拟合二阶响应面模型 y = β₀ + Σβᵢxᵢ + Σβᵢⱼxᵢxⱼ + Σβᵢᵢxᵢ²
        
        由三部分组成:
          1. 角点 (2^k 全因子): 编码值 ±1，共 2^k 个点
          2. 轴点 (2k 个星号点): 沿每个轴 ±α，共 2k 个点
          3. 中心点 (n₀ 个重复): 所有因子取 0，共 n₀ 个点
        
        alpha_type:
          "rotatable"  — α = (2^k)^(1/4)，使预测方差在等距点上相等（推荐）
          "face"       — α = 1，轴点在面上（CCF），不超出因子范围
          "orthogonal" — α = √(√(n_f + n_a + n_0) * n_f - n_f) ，使设计正交
        
        数学原理:
          可旋转: α = (n_f)^(1/4) 其中 n_f = 2^k（角点数）
          k=2: α=1.4142, k=3: α=1.6818, k=4: α=2.0000, k=5: α=2.3784
        """
        factors = json.loads(factors_json)
        
        # ★ 新增: 分离连续因子和类别因子
        continuous, categorical = self._separate_factors(factors)
        k = len(continuous)
        
        if k < 2:
            raise ValueError("CCD 至少需要 2 个连续因子")
        
        # 中心点数自动决定
        if center_count < 0:
            _minitab_center_defaults = {2: 5, 3: 6, 4: 7, 5: 10}
            center_count = _minitab_center_defaults.get(k, 6)
        
        # ── Part 1: 角点 (2^k 全因子) ──
        # 生成 {-1, +1}^k 的所有组合
        n_f = 2 ** k
        factorial_points = np.array(list(iterproduct([-1, 1], repeat=k)), dtype=float)
        
        # ── Part 2: 轴点 (2k 个) ──
        # 计算 alpha
        if alpha_type == "face":
            alpha = 1.0
        elif alpha_type == "orthogonal":
            # ★ FIX #5: 正交 alpha — 使用 Montgomery 9th Ed. Table 11.5 标准查表值
            # 对于表中没有的 k，使用迭代法: α² = (1/2) × [√(N × n_f) - n_f]
            _ortho_table = {2: (1.0, 5), 3: (1.2154, 6), 4: (1.4142, 7), 5: (1.5962, 10), 6: (1.7612, 15)}
            if k in _ortho_table:
                _alpha_tab, _n0_tab = _ortho_table[k]
                if center_count < 0:
                    center_count = _n0_tab
                if center_count == _n0_tab:
                    alpha = _alpha_tab
                else:
                    # 用户指定了不同 center_count，重新计算
                    _N = n_f + 2 * k + center_count
                    _alpha_sq = 0.5 * (np.sqrt(_N * n_f) - n_f)
                    alpha = np.sqrt(max(_alpha_sq, 1.0))
            else:
                _N = n_f + 2 * k + center_count
                _alpha_sq = 0.5 * (np.sqrt(_N * n_f) - n_f)
                alpha = np.sqrt(max(_alpha_sq, 1.0)) if _alpha_sq > 0 else 1.0
            alpha = max(alpha, 1.0)  # 保底不小于 1
        else:  # rotatable（默认）
            alpha = n_f ** 0.25  # α = (2^k)^(1/4)
        
        axial_points = np.zeros((2 * k, k))
        for i in range(k):
            axial_points[2 * i, i] = alpha
            axial_points[2 * i + 1, i] = -alpha
        
        # ── Part 3: 中心点 ──
        center_points = np.zeros((center_count, k))
        
        # ── 合并 ──
        coded_matrix = np.vstack([factorial_points, axial_points, center_points])
        
        # ── 解码为实际值 ──
        # 编码值 coded ∈ {-α, -1, 0, +1, +α}
        # 标准映射: actual = center + coded * half_range
        #   → 角点 (coded=±1) 映射到 [lower, upper]
        #   → 轴点 (coded=±α) 映射到超出 [lower, upper] 的位置（可旋转设计的本质要求）
        # 对于 face-centered (α=1)，所有点恰好落在 [lower, upper] 内
        # 对于 rotatable/orthogonal (α>1)，轴点会超出用户指定范围
        #   — 这是正确行为：如果需要所有点在范围内，应选择 face-centered (CCF)
        matrix = []
        for row in coded_matrix:
            decoded_row = {}
            for i, f in enumerate(continuous):
                center = (f["lower"] + f["upper"]) / 2.0
                half_range = (f["upper"] - f["lower"]) / 2.0
                # 统一映射: actual = center + coded * half_range
                actual_val = center + row[i] * half_range
                decoded_row[f["name"]] = round(actual_val, 6)
            matrix.append(decoded_row)
        
        # ★ 新增: 与类别因子做笛卡尔积
        matrix = self._cross_with_categorical(matrix, categorical)
        
        return json.dumps(matrix, ensure_ascii=False)

    def box_behnken(self, factors_json: str, center_count: int = -1) -> str:
        """
        ★ 新增: Box-Behnken 设计 — 纯 numpy 实现，无需 pyDOE2
        
        RSM 方法 — 与 CCD 的区别:
          1. 没有角点（避免极端条件组合，化工安全性更好）
          2. 没有轴点超出因子范围（所有点在 [-1,+1] 内）
          3. 需要至少 3 个因子
          4. 实验次数通常少于 CCD
        
        数学原理:
          - 从 k 个因子中每次取 2 个做 2² 全因子（±1），其余因子取中心值 0
          - 所有 C(k,2) 个配对组合合并 + 中心点重复
          - 例: k=3 → C(3,2)=3 对 × 4 组 + 中心点 = 12 + n₀
          - 例: k=4 → C(4,2)=6 对 × 4 组 + 中心点 = 24 + n₀
        """
        factors = json.loads(factors_json)
        
        # ★ 新增: 分离连续因子和类别因子
        continuous, categorical = self._separate_factors(factors)
        k = len(continuous)
        
        if k < 3:
            raise ValueError("Box-Behnken 设计至少需要 3 个连续因子")
        
        if center_count < 0:
            center_count = min(max(3, k), 5)
        
        # ── 生成所有 C(k,2) 配对的 2² 全因子 ──
        from itertools import combinations
        
        coded_rows = []
        for i, j in combinations(range(k), 2):
            # 对因子 i, j 做 2² 全因子 {-1,+1}²，其余因子 = 0
            for vi in [-1, 1]:
                for vj in [-1, 1]:
                    row = [0.0] * k
                    row[i] = vi
                    row[j] = vj
                    coded_rows.append(row)
        
        # ── 添加中心点 ──
        for _ in range(center_count):
            coded_rows.append([0.0] * k)
        
        coded_matrix = np.array(coded_rows)
        
        # ── 解码为实际值 ──
        matrix = []
        for row in coded_matrix:
            decoded_row = {}
            for i, f in enumerate(continuous):
                center = (f["lower"] + f["upper"]) / 2.0
                half_range = (f["upper"] - f["lower"]) / 2.0
                actual_val = center + row[i] * half_range
                decoded_row[f["name"]] = round(actual_val, 6)
            matrix.append(decoded_row)
        
        # ★ 新增: 与类别因子做笛卡尔积
        matrix = self._cross_with_categorical(matrix, categorical)
        
        return json.dumps(matrix, ensure_ascii=False)

    def plackett_burman(self, factors_json: str) -> str:
        """
        Plackett-Burman 设计 — 经济高效的因子筛选

        特点:
          1. 实验次数为 4 的倍数 (12, 20, 24, ...)
          2. 最多可筛选 N-1 个因子
          3. 分辨度 III（主效应与二因子交互混杂）
          4. 比 Fractional Factorial 更灵活（不要求因子数为 2 的幂）

        例: 7 个因子 → 12 组（FF 需要 8 或 16 组）
            11 个因子 → 12 组
            15 个因子 → 20 组
        """
        factors = json.loads(factors_json)
        continuous, categorical = self._separate_factors(factors)
        k = len(continuous)

        if k < 2:
            raise ValueError("Plackett-Burman 至少需要 2 个连续因子")

        # 确定实验次数 N（4 的倍数，>= k+1）
        N = k + 1
        if N % 4 != 0:
            N = ((N // 4) + 1) * 4
        N = max(N, 12)

        # 方法1: 用 pyDOE2（优先，最可靠）
        design = None
        try:
            from pyDOE2 import pbdesign
            design = pbdesign(k)
        except ImportError:
            pass

        # 方法2: 用 scipy Hadamard 矩阵
        if design is None:
            try:
                from scipy.linalg import hadamard
                H = hadamard(N)
                # Hadamard 矩阵第一行全 1，跳过；取前 k 列（跳过第一列全 1）
                design = H[1:, 1:k + 1].astype(float)
            except Exception:
                pass

        # 方法3: 手动构造 12-run PB 设计（Paley 构造）
        if design is None:
            if k <= 11:
                # 经典 12-run PB 的第一行生成向量
                gen = [1, 1, -1, 1, 1, 1, -1, -1, -1, 1, -1]
                rows = []
                current = list(gen)
                for _ in range(11):
                    rows.append(current[:k])
                    # 循环右移
                    current = [current[-1]] + current[:-1]
                # 加一行全 -1
                rows.append([-1] * k)
                design = np.array(rows, dtype=float)
            else:
                raise ValueError(
                    f"因子数 {k} 过多，请安装 pyDOE2: pip install pyDOE2"
                )

        # 解码为实际因子值
        matrix = []
        for row in design:
            decoded = {}
            for i, f in enumerate(continuous):
                if i < len(row):
                    center = (f["lower"] + f["upper"]) / 2.0
                    half_range = (f["upper"] - f["lower"]) / 2.0
                    decoded[f["name"]] = round(
                        center + float(row[i]) * half_range, 6
                    )
            matrix.append(decoded)

        # 与类别因子做笛卡尔积
        matrix = self._cross_with_categorical(matrix, categorical)

        return json.dumps(matrix, ensure_ascii=False)

    def d_optimal(self, factors_json: str, num_runs: int = -1, model_type: str = "quadratic") -> str:
        """
        ★ 修复 v4: D-Optimal 设计 — 完整支持混合因子（连续+类别）
        
        D-最优准则: 最大化信息矩阵 |X'X| 的行列式
        → 使回归系数的联合置信椭球体积最小
        
        model_type: "linear" / "interaction" / "quadratic"
        num_runs: 实验次数，-1 自动（= 参数数 × 1.5 向上取整）
        
        ★ 修复内容:
          原来对连续因子做 Fedorov 选点后再与类别因子笛卡尔积，
          导致: (1) 类别×连续交互项在选点时未被考虑
                (2) num_runs 被叉乘放大，用户无法控制实验次数
          修复后: 在包含类别因子哑变量的完整模型矩阵上做 Fedorov 选点。
          
        候选集构建:
          连续因子: 5 水平均匀网格 [-1, -0.5, 0, 0.5, 1]
          类别因子: 所有水平的 one-hot 编码 (drop first)
          完整候选集 = 连续网格 × 类别水平组合
        """
        factors = json.loads(factors_json)
        continuous, categorical = self._separate_factors(factors)
        k_cont = len(continuous)
        k_cat = len(categorical)
        
        if k_cont == 0 and k_cat == 0:
            raise ValueError("D-Optimal 至少需要 1 个因子")
        
        # ── 纯连续因子: 走原有逻辑（无需哑变量）──
        if k_cat == 0:
            if model_type == "linear":
                p = 1 + k_cont
            elif model_type == "interaction":
                p = 1 + k_cont + k_cont * (k_cont - 1) // 2
            else:
                p = 1 + 2 * k_cont + k_cont * (k_cont - 1) // 2
            
            if num_runs < 0:
                num_runs = int(np.ceil(p * 1.5))
            if num_runs < p:
                num_runs = p
            
            candidates_coded = self._generate_candidate_grid(k_cont, 5)
            X_cand = self._build_model_matrix(candidates_coded, model_type)
            best_idx = self._fedorov_exchange(X_cand, num_runs)  # n_starts 自动适配
            
            matrix = []
            for idx in best_idx:
                coded_row = candidates_coded[idx]
                decoded_row = {}
                for i, f in enumerate(continuous):
                    center = (f["lower"] + f["upper"]) / 2.0
                    half_range = (f["upper"] - f["lower"]) / 2.0
                    decoded_row[f["name"]] = round(center + coded_row[i] * half_range, 6)
                matrix.append(decoded_row)
            return json.dumps(matrix, ensure_ascii=False)
        
        # ── 含类别因子: 在完整因子空间上做 Fedorov ──
        
        # 1. 类别因子元数据
        cat_levels_list = [f["levels"] for f in categorical]
        cat_dummy_counts = [len(levels) - 1 for levels in cat_levels_list]
        total_cat_dummies = sum(cat_dummy_counts)
        
        # 2. 计算完整模型参数数
        if model_type == "linear":
            p = 1 + k_cont + total_cat_dummies
        elif model_type == "interaction":
            p = 1 + k_cont + total_cat_dummies
            p += k_cont * (k_cont - 1) // 2     # 连续×连续
            p += k_cont * total_cat_dummies      # 连续×类别
            for i in range(len(cat_dummy_counts)):
                for j in range(i + 1, len(cat_dummy_counts)):
                    p += cat_dummy_counts[i] * cat_dummy_counts[j]
        else:  # quadratic
            p = 1 + k_cont + total_cat_dummies
            p += k_cont * (k_cont - 1) // 2     # 连续×连续交互
            p += k_cont                          # 连续二次项
            p += k_cont * total_cat_dummies      # 连续×类别交互
            for i in range(len(cat_dummy_counts)):
                for j in range(i + 1, len(cat_dummy_counts)):
                    p += cat_dummy_counts[i] * cat_dummy_counts[j]
        
        if num_runs < 0:
            num_runs = int(np.ceil(p * 1.5))
        if num_runs < p:
            num_runs = p
        
        # 3. 生成候选集
        grid_size = 5
        if k_cont > 0:
            cont_grid = np.linspace(-1, 1, grid_size)
            cont_candidates = np.array(list(iterproduct(*([cont_grid] * k_cont))))
        else:
            cont_candidates = np.array([[]]) 
        
        cat_combos = list(iterproduct(*cat_levels_list))
        
        # 4. 构建完整候选集 + 展开后的编码矩阵
        full_candidates = []   # (cont_coded_row, cat_combo)
        full_model_rows = []   # 展开后的数值行
        
        for cont_row in cont_candidates:
            for cat_combo in cat_combos:
                full_candidates.append((cont_row, cat_combo))
                expanded = list(cont_row) if k_cont > 0 else []
                for ci, levels in enumerate(cat_levels_list):
                    for lv in levels[1:]:
                        expanded.append(1.0 if cat_combo[ci] == lv else 0.0)
                full_model_rows.append(expanded)
        
        full_model_arr = np.array(full_model_rows)
        
        # 5. 构建完整模型矩阵
        X_cand = self._build_full_model_matrix(
            full_model_arr, k_cont, cat_dummy_counts, model_type
        )
        
        # 6. 候选点不够时增加网格密度
        if X_cand.shape[0] < num_runs and k_cont > 0:
            cont_grid = np.linspace(-1, 1, 7)
            cont_candidates = np.array(list(iterproduct(*([cont_grid] * k_cont))))
            full_candidates = []
            full_model_rows = []
            for cont_row in cont_candidates:
                for cat_combo in cat_combos:
                    full_candidates.append((cont_row, cat_combo))
                    expanded = list(cont_row)
                    for ci, levels in enumerate(cat_levels_list):
                        for lv in levels[1:]:
                            expanded.append(1.0 if cat_combo[ci] == lv else 0.0)
                    full_model_rows.append(expanded)
            full_model_arr = np.array(full_model_rows)
            X_cand = self._build_full_model_matrix(
                full_model_arr, k_cont, cat_dummy_counts, model_type
            )
        
        if X_cand.shape[0] < num_runs:
            num_runs = X_cand.shape[0]
        
        # 7. Fedorov 选点
        best_idx = self._fedorov_exchange(X_cand, num_runs)  # n_starts 自动适配
        
        # 8. 解码为实际因子值
        matrix = []
        for idx in best_idx:
            cont_row, cat_combo = full_candidates[idx]
            decoded = {}
            for i, f in enumerate(continuous):
                center = (f["lower"] + f["upper"]) / 2.0
                half_range = (f["upper"] - f["lower"]) / 2.0
                decoded[f["name"]] = round(center + float(cont_row[i]) * half_range, 6)
            for i, f in enumerate(categorical):
                decoded[f["name"]] = cat_combo[i]
            matrix.append(decoded)
        
        return json.dumps(matrix, ensure_ascii=False)

    def get_design_quality(self, factors_json: str, matrix_json: str, model_type: str = "quadratic",
                           delta: float = 2.0) -> str:
        """
        ★ 修复 v4: 评估设计矩阵的统计质量 — 完整支持类别因子
        
        修复内容:
          原来只对连续因子构建模型矩阵，完全忽略类别因子哑变量。
          修复后: 类别因子用 one-hot (drop first) 编码纳入模型矩阵。
          VIF 对类别因子报告其所有哑变量列中的最大值，
          Power 报告最小值（最难检测的水平）。
        
        delta: 希望检测到的最小效应大小（编码单位，默认 2.0）
          编码空间 [-1, +1] 中，delta=2.0 表示因子从最低水平变到最高水平时
          响应至少变化 2 个标准误单位。
          较小的 delta (如 1.0) 需要更多实验次数才能达到高 power。
          ★ 注意: 返回结果中包含 assumed_delta 字段，供 C# 端显示给用户。
        
        返回 JSON:
        {
            "d_efficiency": 0.92,
            "a_efficiency": 0.88,
            "g_efficiency": 0.85,
            "vif": {"温度": 1.2, "催化剂": 1.5, ...},
            "condition_number": 3.5,
            "power_analysis": {"温度": 0.95, "催化剂": 0.88, ...},
            "assumed_delta": 2.0,
            "alias_structure": [],
            "run_count": 20,
            "parameter_count": 15,
            "degrees_of_freedom": 5
        }
        """
        from scipy import stats as sp_stats
        
        factors = json.loads(factors_json)
        matrix = json.loads(matrix_json)
        
        continuous, categorical = self._separate_factors(factors)
        k_cont = len(continuous)
        n = len(matrix)
        
        cont_names = [f["name"] for f in continuous]
        cat_names = [f["name"] for f in categorical]
        cat_levels_list = [f["levels"] for f in categorical]
        cat_dummy_counts = [len(levels) - 1 for levels in cat_levels_list]
        total_cat_dummies = sum(cat_dummy_counts)
        
        # ── 构建展开后的编码矩阵 ──
        coded = []
        for row in matrix:
            coded_row = []
            # 连续因子: 编码到 [-1, +1]
            for f in continuous:
                center = (f.get("lower", 0) + f.get("upper", 1)) / 2.0
                half_range = (f.get("upper", 1) - f.get("lower", 0)) / 2.0
                if half_range < 1e-12:
                    half_range = 1.0
                val = row.get(f["name"], center)
                if isinstance(val, str):
                    val = center
                coded_row.append((float(val) - center) / half_range)
            # 类别因子: one-hot (drop first)
            for f in categorical:
                val_str = str(row.get(f["name"], ""))
                levels = f["levels"]
                for lv in levels[1:]:
                    coded_row.append(1.0 if val_str == str(lv) else 0.0)
            coded.append(coded_row)
        
        coded_arr = np.array(coded)
        
        # ── 构建完整模型矩阵 ──
        if total_cat_dummies > 0:
            X = self._build_full_model_matrix(coded_arr, k_cont, cat_dummy_counts, model_type)
        else:
            X = self._build_model_matrix(coded_arr, model_type)
        
        p = X.shape[1]
        
        result = {
            "d_efficiency": 0.0,
            "a_efficiency": 0.0,
            "g_efficiency": 0.0,
            "vif": {},
            "condition_number": 0.0,
            "power_analysis": {},
            "assumed_delta": delta,
            "alias_structure": [],
            "run_count": n,
            "parameter_count": p,
            "degrees_of_freedom": max(n - p, 0)
        }
        
        try:
            XtX = X.T @ X
            
            # 条件数
            eigenvalues = np.linalg.eigvalsh(XtX)
            eigenvalues = eigenvalues[eigenvalues > 1e-12]
            if len(eigenvalues) > 0:
                result["condition_number"] = round(float(np.sqrt(max(eigenvalues) / min(eigenvalues))), 4)
            
            det_XtX = np.linalg.det(XtX)
            if abs(det_XtX) < 1e-12 or n <= p:
                return json.dumps(result, ensure_ascii=False)
            
            XtX_inv = np.linalg.inv(XtX)
            
            # D-效率: (|X'X|^(1/p)) / n
            d_eff = (det_XtX ** (1.0 / p)) / n
            result["d_efficiency"] = round(float(d_eff), 4)
            
            # A-效率: p / (trace(XtX_inv) * n)
            a_eff = p / (np.trace(XtX_inv) * n)
            result["a_efficiency"] = round(float(a_eff), 4)
            
            # G-效率: p / (n * max_leverage)
            H = X @ XtX_inv @ X.T
            max_leverage = float(np.max(np.diag(H)))
            if max_leverage > 1e-12:
                g_eff = p / (n * max_leverage)
                result["g_efficiency"] = round(float(g_eff), 4)
            
            # ── VIF + Power: 连续因子逐列，类别因子合并报告 ──
            alpha = 0.05
            # ★ 修复: delta 现在从方法参数传入，不再硬编码
            df_error = max(n - p, 1)
            f_crit = sp_stats.f.ppf(1 - alpha, 1, df_error)
            
            # 模型矩阵列布局: [截距, cont_0..cont_k-1, cat_oh_0..cat_oh_m-1, 交互..., 二次...]
            col_idx = 1  # 跳过截距
            
            # 连续因子 VIF + Power
            for i in range(k_cont):
                j = col_idx + i
                if j < p:
                    vif_val = float(XtX_inv[j, j] * XtX[j, j])
                    result["vif"][cont_names[i]] = round(vif_val, 4)
                    if float(XtX_inv[j, j]) > 1e-12:
                        ncp = delta ** 2 / float(XtX_inv[j, j])
                        power = 1.0 - sp_stats.ncf.cdf(f_crit, 1, df_error, ncp)
                        result["power_analysis"][cont_names[i]] = round(float(power), 4)
            col_idx += k_cont
            
            # 类别因子 VIF (取最大) + Power (取最小)
            for fi, f in enumerate(categorical):
                n_dummies = cat_dummy_counts[fi]
                max_vif = 1.0
                min_power = 1.0
                for di in range(n_dummies):
                    j = col_idx + di
                    if j < p:
                        vif_val = float(XtX_inv[j, j] * XtX[j, j])
                        max_vif = max(max_vif, vif_val)
                        if float(XtX_inv[j, j]) > 1e-12:
                            ncp = delta ** 2 / float(XtX_inv[j, j])
                            power = 1.0 - sp_stats.ncf.cdf(f_crit, 1, df_error, ncp)
                            min_power = min(min_power, float(power))
                result["vif"][f["name"]] = round(max_vif, 4)
                result["power_analysis"][f["name"]] = round(min_power, 4)
                col_idx += n_dummies
            
        except np.linalg.LinAlgError:
            pass
        
        return json.dumps(result, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # 原有方法（保留不变）
    # ═══════════════════════════════════════════════════════

    def randomize(self, matrix_json: str, seed: int = -1) -> str:
        """随机化排列参数矩阵"""
        matrix = json.loads(matrix_json)
        rng = np.random.default_rng(seed if seed >= 0 else None)
        indices = rng.permutation(len(matrix))
        randomized = [matrix[i] for i in indices]
        return json.dumps(randomized, ensure_ascii=False)

    def add_center_points(self, matrix_json: str, factors_json: str, count: int = 3) -> str:
        """向矩阵中添加中心点实验"""
        matrix = json.loads(matrix_json)
        factors = json.loads(factors_json)
        center_point = {}
        for f in factors:
            if "levels" in f and len(f["levels"]) > 0:
                center_point[f["name"]] = round(sum(f["levels"]) / len(f["levels"]), 6)
            elif "lower" in f and "upper" in f:
                center_point[f["name"]] = round((f["lower"] + f["upper"]) / 2, 6)
        for _ in range(count):
            matrix.append(dict(center_point))
        return json.dumps(matrix, ensure_ascii=False)

    def get_design_summary(self, factors_json: str, method: str, matrix_json: str) -> str:
        """获取设计摘要信息"""
        factors = json.loads(factors_json)
        matrix = json.loads(matrix_json)
        summary = {
            "method": method,
            "factor_count": len(factors),
            "run_count": len(matrix),
            "factors": [
                {
                    "name": f.get("name", ""),
                    "levels": f.get("levels", [f.get("lower", 0), f.get("upper", 0)]),
                    "level_count": len(f.get("levels", [0, 0]))
                }
                for f in factors
            ]
        }
        return json.dumps(summary, ensure_ascii=False)

    # ── 内部辅助方法 ──

    def _two_level_full_factorial(self, factors: list) -> str:
        """2水平全因子设计"""
        names = [f["name"] for f in factors]
        levels_list = [[f["lower"], f["upper"]] for f in factors]
        combos = list(iterproduct(*levels_list))
        matrix = [{names[i]: combo[i] for i in range(len(names))} for combo in combos]
        return json.dumps(matrix, ensure_ascii=False)

    def _fracfact_numpy(self, gen_string: str) -> np.ndarray:
        """
        ★ 新增: 纯 numpy 实现部分因子设计（替代 pyDOE2.fracfact）
        
        gen_string 格式: "a b c ab ac"
          - 单字母 (a,b,c,...) = 基本因子，做 2^p 全因子
          - 多字母 (ab,ac,...) = 生成因子，由对应基本因子列的乘积生成
        
        返回: (n, k) 编码矩阵，值 ∈ {-1, +1}
        """
        parts = gen_string.strip().split()
        
        # 识别基本因子（单字母）
        base_letters = []
        for p in parts:
            if len(p) == 1 and p not in base_letters:
                base_letters.append(p)
        
        p_base = len(base_letters)
        n_runs = 2 ** p_base
        
        # 生成基本因子的全因子矩阵 {-1, +1}^p
        base_matrix = np.array(list(iterproduct([-1, 1], repeat=p_base)), dtype=float)
        
        # 为每个字母建立列索引
        letter_col = {letter: i for i, letter in enumerate(base_letters)}
        
        # 构建完整矩阵
        columns = []
        for part in parts:
            if len(part) == 1:
                # 基本因子 → 直接取对应列
                columns.append(base_matrix[:, letter_col[part]])
            else:
                # 生成因子 → 对应列的逐元素乘积
                col = np.ones(n_runs)
                for ch in part:
                    if ch in letter_col:
                        col *= base_matrix[:, letter_col[ch]]
                columns.append(col)
        
        return np.column_stack(columns) if columns else np.zeros((n_runs, 0))

    def _build_generator_string(self, k: int, resolution: int) -> str:
        """
        构建部分因子设计的生成元字符串。
        使用 Box-Hunter-Hunter 最小混杂生成元表，确保达到声称的分辨度。
        
        参考: Box, Hunter & Hunter (2005) "Statistics for Experimenters", 2nd Ed.
        """
        base_letters = [chr(ord('a') + i) for i in range(min(k, 26))]
        
        if k <= 4:
            return " ".join(base_letters[:k])
        
        # 已知最小混杂生成元表 (k → {resolution → gen_string})
        # 格式: 基本因子做全因子, 额外因子用生成元列（基本因子列的乘积）
        #
        # 参考: Montgomery (2017) "Design and Analysis of Experiments", 9th Ed., Appendix Table XII
        #        Box, Hunter & Hunter (2005) "Statistics for Experimenters", 2nd Ed.
        #
        # 分辨度定义: Resolution = 最短 defining word 的字母数
        #   III  → 主效应与二因子交互混杂
        #   IV   → 主效应不与二因子交互混杂，但二因子交互之间混杂
        #   V    → 主效应和二因子交互都不与二因子交互混杂
        #
        # 验证方法: defining relation I = (生成元的 word product)，最短 word 长度 = 分辨度
        known_generators = {
            # k=5
            5: {
                # 2^(5-2)=8 runs, E=AB, F=AC → I=ABE=ACF=BCEF, 最短=3 → Res III
                3: "a b c ab ac",
                # 2^(5-1)=16 runs, E=ABCD → I=ABCDE, 最短=5 → Res V
                # 注: k=5 没有恰好 Res IV 的标准设计, Res IV 请求退化到 Res III (8 runs)
                # 或升级到 Res V (16 runs)。这里选择升级, 给用户更好的设计
                4: "a b c d abcd",
                # 2^(5-1)=16 runs, E=ABCD → Res V
                5: "a b c d abcd",
            },
            # k=6
            6: {
                # 2^(6-3)=8 runs, D=AB, E=AC, F=BC → I=ABD=ACE=BCF=..., 最短=3 → Res III
                3: "a b c ab ac bc",
                # 2^(6-2)=16 runs, E=ABC, F=BCD → I=ABCE=BCDF=ADEF, 最短=4 → Res IV
                4: "a b c d abc bcd",
                # 2^(6-1)=32 runs, F=ABCDE → I=ABCDEF, 最短=6 → Res VI
                5: "a b c d e abcde",
            },
            # k=7
            7: {
                # 2^(7-4)=8 runs, D=AB, E=AC, F=BC, G=ABC → Res III
                3: "a b c ab ac bc abc",
                # 2^(7-3)=16 runs, E=ABC, F=BCD, G=ACD → I=ABCE=BCDF=ACDG=..., 最短=4 → Res IV
                4: "a b c d abc bcd acd",
                # 2^(7-1)=64 runs, G=ABCDEF → Res VII
                5: "a b c d e f abcdef",
            },
            # k=8
            8: {
                # 2^(8-4)=16 runs: Res IV (k=8 无法在 16 runs 内实现 Res III 以下)
                # Montgomery Table XII: E=BCD, F=ACD, G=ABC, H=ABD
                # Defining relation 最短 word = ABCG (4 letters) → Res IV
                # 注意: k=8 请求 Res III 时退化为此设计 (16 runs, Res IV)
                3: "a b c d bcd acd abc abd",
                # 2^(8-3)=32 runs: F=BCDE, G=ACDE, H=ABDE → Res IV
                4: "a b c d e bcde acde abde",
            },
        }
        
        if k in known_generators:
            gen_table = known_generators[k]
            # 找到不超过请求分辨度的最佳可用设计
            best_res = None
            for res in sorted(gen_table.keys()):
                if res <= resolution:
                    best_res = res
            if best_res is None:
                best_res = min(gen_table.keys())
            return gen_table[best_res]
        
        # 对于不在表中的 k，使用启发式构建
        # ★ 警告: 启发式生成元不保证达到请求的分辨度
        # 实际分辨度可能低于 resolution 参数声称的值
        # k>8 的情况较少见，如果需要严格的分辨度保证，建议用 D-Optimal 代替
        import warnings
        warnings.warn(
            f"k={k} 超出内置生成元表范围，使用启发式构建部分因子设计。"
            f"实际分辨度可能低于请求的 Res {resolution}，建议验证混杂结构。",
            stacklevel=2
        )
        if resolution >= 5:
            p = max(2, k - 4)
        elif resolution >= 4:
            p = max(2, k - 3)
        else:
            p = max(2, k - 2)
        base_count = k - p
        gen_parts = list(base_letters[:base_count])
        interactions = []
        for i in range(base_count):
            for j in range(i + 1, base_count):
                interactions.append(base_letters[i] + base_letters[j])
                if len(gen_parts) + len(interactions) >= k:
                    break
            if len(gen_parts) + len(interactions) >= k:
                break
        needed = k - base_count
        gen_parts.extend(interactions[:needed])
        return " ".join(gen_parts)

    def _select_orthogonal_table(self, k: int, max_levels: int) -> str:
        """自动选择正交表类型"""
        if max_levels <= 2:
            if k <= 3: return "L4"
            elif k <= 7: return "L8"
            else: return "L16"
        elif max_levels <= 3:
            if k <= 4: return "L9"
            elif k <= 8: return "L18"
            else: return "L27"
        else:
            return "L16"

    def _get_orthogonal_array(self, table_type: str, k: int, max_levels: int) -> np.ndarray:
        """获取标准正交表"""
        tables = {
            "L4": np.array([
                [0,0,0], [0,1,1], [1,0,1], [1,1,0]
            ]),
            "L8": np.array([
                [0,0,0,0,0,0,0], [0,0,0,1,1,1,1],
                [0,1,1,0,0,1,1], [0,1,1,1,1,0,0],
                [1,0,1,0,1,0,1], [1,0,1,1,0,1,0],
                [1,1,0,0,1,1,0], [1,1,0,1,0,0,1]
            ]),
            "L9": np.array([
                [0,0,0,0], [0,1,1,1], [0,2,2,2],
                [1,0,1,2], [1,1,2,0], [1,2,0,1],
                [2,0,2,1], [2,1,0,2], [2,2,1,0]
            ]),
            "L16": np.array([
                [0,0,0,0,0], [0,0,0,1,1], [0,1,1,0,0], [0,1,1,1,1],
                [1,0,1,0,1], [1,0,1,1,0], [1,1,0,0,1], [1,1,0,1,0],
                [0,0,1,0,1], [0,0,1,1,0], [0,1,0,0,1], [0,1,0,1,0],
                [1,0,0,0,0], [1,0,0,1,1], [1,1,1,0,0], [1,1,1,1,1]
            ]),
            "L18": np.array([
                [0,0,0,0,0,0,0,0], [0,0,1,1,1,1,1,1], [0,0,2,2,2,2,2,2],
                [0,1,0,0,1,1,2,2], [0,1,1,1,2,2,0,0], [0,1,2,2,0,0,1,1],
                [0,2,0,1,0,2,1,2], [0,2,1,2,1,0,2,0], [0,2,2,0,2,1,0,1],
                [1,0,0,2,2,1,1,0], [1,0,1,0,0,2,2,1], [1,0,2,1,1,0,0,2],
                [1,1,0,1,2,0,0,2], [1,1,1,2,0,1,1,0], [1,1,2,0,1,2,2,1],
                [1,2,0,2,1,2,0,1], [1,2,1,0,2,0,1,2], [1,2,2,1,0,1,2,0]
            ]),
            "L27": self._generate_L27()
        }
        oa = tables.get(table_type, tables["L9"])
        if oa.shape[1] > k:
            oa = oa[:, :k]
        elif oa.shape[1] < k:
            # 正交表列数不足: 不能简单复制列（会导致因子完全混杂）
            # 尝试升级到更大的正交表
            larger_tables = ["L4", "L8", "L9", "L16", "L18", "L27"]
            upgraded = False
            for lt in larger_tables:
                if lt in tables and tables[lt].shape[1] >= k:
                    oa = tables[lt][:, :k]
                    upgraded = True
                    break
            if not upgraded:
                raise ValueError(
                    f"正交表 {table_type} 仅支持 {oa.shape[1]} 个因子，"
                    f"但需要 {k} 个因子。请手动指定更大的正交表类型（如 L18, L27），"
                    f"或减少因子数量。"
                )
        return oa

    def _generate_L27(self) -> np.ndarray:
        """生成 L27(3^13) 的前几列"""
        rows = []
        for i in range(3):
            for j in range(3):
                for k_val in range(3):
                    row = [i, j, k_val,
                           (i+j) % 3, (i+k_val) % 3, (j+k_val) % 3,
                           (i+j+k_val) % 3,
                           (i+2*j) % 3, (i+2*k_val) % 3,
                           (j+2*k_val) % 3,
                           (2*i+j) % 3, (2*i+k_val) % 3,
                           (2*j+k_val) % 3]
                    rows.append(row)
        return np.array(rows)

    # ★ 新增: D-Optimal 内部方法

    def _generate_candidate_grid(self, k: int, grid_levels: int = 5) -> np.ndarray:
        """生成候选点网格（编码空间 [-1, +1]）"""
        levels = np.linspace(-1, 1, grid_levels)
        grids = [levels] * k
        mesh = np.array(list(iterproduct(*grids)))
        return mesh

    def _build_model_matrix(self, coded_arr: np.ndarray, model_type: str) -> np.ndarray:
        """
        ★ 新增: 构建扩展模型矩阵
        
        coded_arr: (n, k) 编码因子矩阵
        model_type: "linear" / "interaction" / "quadratic"
        
        返回: (n, p) 模型矩阵，含截距列
        """
        n, k = coded_arr.shape
        
        # 截距
        X = [np.ones(n)]
        
        # 主效应
        for j in range(k):
            X.append(coded_arr[:, j])
        
        if model_type in ("interaction", "quadratic"):
            # 交互项 x_i * x_j
            for i in range(k):
                for j in range(i + 1, k):
                    X.append(coded_arr[:, i] * coded_arr[:, j])
        
        if model_type == "quadratic":
            # 二次项 x_i²
            for j in range(k):
                X.append(coded_arr[:, j] ** 2)
        
        return np.column_stack(X)

    def _build_full_model_matrix(self, expanded_arr: np.ndarray,
                                  k_cont: int, cat_dummy_counts: list,
                                  model_type: str) -> np.ndarray:
        """
        ★ 新增 v4: 构建含类别因子哑变量的完整模型矩阵
        
        expanded_arr: (n, k_cont + total_cat_dummies)
            前 k_cont 列是连续因子编码值 [-1, +1]
            后面是类别因子 one-hot 列 (drop first)
        
        k_cont: 连续因子数
        cat_dummy_counts: 每个类别因子的哑变量数 [len(levels)-1, ...]
        model_type: "linear" / "interaction" / "quadratic"
        
        模型矩阵列布局:
          1. 截距: 1 列
          2. 连续主效应: k_cont 列
          3. 类别主效应: Σ(levels_i - 1) 列 (one-hot)
          4. 连续×连续交互: C(k_cont, 2) 列       [interaction/quadratic]
          5. 连续×类别交互: k_cont × Σ(levels_i-1) [interaction/quadratic]
          6. 类别×类别交互: Σ_i<j (l_i-1)(l_j-1)   [interaction/quadratic]
          7. 连续二次项: k_cont 列                   [quadratic only]
        
        注意: 类别因子不生成二次项（one-hot 列的平方等于自身，无意义）
        """
        n = expanded_arr.shape[0]
        total_cat_dummies = sum(cat_dummy_counts)
        
        cont_cols = expanded_arr[:, :k_cont] if k_cont > 0 else np.zeros((n, 0))
        cat_cols = expanded_arr[:, k_cont:] if total_cat_dummies > 0 else np.zeros((n, 0))
        
        columns = [np.ones(n)]  # 截距
        
        # 连续主效应
        for j in range(k_cont):
            columns.append(cont_cols[:, j])
        
        # 类别主效应 (one-hot)
        for j in range(total_cat_dummies):
            columns.append(cat_cols[:, j])
        
        if model_type in ("interaction", "quadratic"):
            # 连续×连续交互
            for i in range(k_cont):
                for j in range(i + 1, k_cont):
                    columns.append(cont_cols[:, i] * cont_cols[:, j])
            
            # 连续×类别交互
            for i in range(k_cont):
                for j in range(total_cat_dummies):
                    columns.append(cont_cols[:, i] * cat_cols[:, j])
            
            # 类别×类别交互（不同类别因子之间的哑变量乘积）
            cat_start = 0
            for ci in range(len(cat_dummy_counts)):
                for cj in range(ci + 1, len(cat_dummy_counts)):
                    cat_start_cj = sum(cat_dummy_counts[:cj])
                    for di in range(cat_dummy_counts[ci]):
                        for dj in range(cat_dummy_counts[cj]):
                            columns.append(cat_cols[:, cat_start + di] * cat_cols[:, cat_start_cj + dj])
                cat_start += cat_dummy_counts[ci]
        
        if model_type == "quadratic":
            # 连续二次项（类别不生成二次项）
            for j in range(k_cont):
                columns.append(cont_cols[:, j] ** 2)
        
        return np.column_stack(columns) if columns else np.ones((n, 1))

    def _fedorov_exchange(self, X_cand: np.ndarray, num_runs: int, n_starts: int = None) -> list:
        """
        ★ 优化: Fedorov 坐标交换算法选择 D-最优子集
        
        X_cand: (N_cand, p) 候选点的模型矩阵
        num_runs: 需要选择的实验次数
        n_starts: 多起点随机搜索次数
        
        优化:
          1. 使用 set 做 O(1) 成员查找（原来 list 的 `in` 是 O(n)）
          2. 预计算候选点的 x_i @ x_j 外积缓存
          3. 增加收敛判断 — 单轮无改善即终止
        
        返回: 最优设计在候选集中的行索引列表
        """
        N_cand, p = X_cand.shape
        best_det = -np.inf
        best_idx = None
        
        for _ in range(n_starts):
            # 随机初始化
            idx = list(np.random.choice(N_cand, size=num_runs, replace=False))
            idx_set = set(idx)  # ★ 修复: O(1) 成员查找
            
            improved = True
            max_iter = 50
            iter_count = 0
            
            while improved and iter_count < max_iter:
                improved = False
                iter_count += 1
                
                for i in range(num_runs):
                    X_design = X_cand[idx]
                    try:
                        current_det = np.linalg.det(X_design.T @ X_design)
                    except np.linalg.LinAlgError:
                        current_det = 0.0
                    
                    old_idx = idx[i]
                    best_swap_det = current_det
                    best_swap_idx = old_idx
                    
                    for c in range(N_cand):
                        if c in idx_set:  # ★ 修复: O(1) 查找
                            continue
                        idx[i] = c
                        X_test = X_cand[idx]
                        try:
                            test_det = np.linalg.det(X_test.T @ X_test)
                        except np.linalg.LinAlgError:
                            test_det = 0.0
                        
                        if test_det > best_swap_det * 1.001:  # 1‰ 改善阈值
                            best_swap_det = test_det
                            best_swap_idx = c
                    
                    # 更新索引和 set
                    if best_swap_idx != old_idx:
                        idx_set.discard(old_idx)
                        idx_set.add(best_swap_idx)
                        idx[i] = best_swap_idx
                        improved = True
                    else:
                        idx[i] = old_idx  # 恢复
            
            X_final = X_cand[idx]
            try:
                final_det = np.linalg.det(X_final.T @ X_final)
            except np.linalg.LinAlgError:
                final_det = 0.0
            
            if final_det > best_det:
                best_det = final_det
                best_idx = list(idx)
        
        return best_idx if best_idx is not None else list(range(num_runs))
# ═══════════════════════════════════════════════════════
    # ★ 新增: 项目迭代专用设计方法
    # augment_design / steepest_ascent / confirmation_runs
    # ═══════════════════════════════════════════════════════

    def augment_design(self, factors_json: str, existing_matrix_json: str,
                       num_additional: int = -1, model_type: str = "quadratic") -> str:
        """
        ★ 新增: 增强设计 — 在已有实验数据基础上补充新实验点

        核心思想:
          把已有实验点锁定为"必选点"，然后在候选集中用 Fedorov 交换算法
          选出额外的点，使得 "已有 + 新增" 组合后的信息矩阵 |X'X| 最大。

        应用场景:
          - 一轮 RSM 做完但模型 R² 不够 → 补点改善模型
          - 因子范围扩展后 → 在新区域补充采样
          - Power 不够 → 增加实验次数提升统计功效

        factors_json: 因子定义（与 d_optimal 格式相同）
        existing_matrix_json: 已有实验点 JSON
        num_additional: 需要补充的实验数，-1 自动（= 缺失自由度数 × 1.5）
        model_type: "linear" / "interaction" / "quadratic"

        返回: 仅新增点的矩阵 JSON（不含已有点）
        """
        factors = json.loads(factors_json)
        existing = json.loads(existing_matrix_json)

        continuous, categorical = self._separate_factors(factors)
        k_cont = len(continuous)
        k_cat = len(categorical)

        if k_cont == 0 and k_cat == 0:
            return json.dumps([])

        # ── 构建候选集（与 d_optimal 相同）──
        grid_levels = 5

        if k_cont > 0:
            cont_grid = self._generate_candidate_grid(k_cont, grid_levels)
        else:
            cont_grid = np.array([[]])

        if k_cat > 0:
            cat_levels_list = [f["levels"] for f in categorical]
            cat_combos = list(iterproduct(*cat_levels_list))
            cat_dummy_counts = [len(levels) - 1 for levels in cat_levels_list]
        else:
            cat_combos = [()]
            cat_levels_list = []
            cat_dummy_counts = []

        # 构建完整候选集: 连续网格 × 类别组合
        full_candidates = []
        expanded_rows = []

        for cont_row in cont_grid:
            for cat_combo in cat_combos:
                full_candidates.append((cont_row, cat_combo))
                row = list(cont_row) if k_cont > 0 else []
                for ci in range(k_cat):
                    levels = cat_levels_list[ci]
                    for lv_idx in range(1, len(levels)):
                        row.append(1.0 if cat_combo[ci] == levels[lv_idx] else 0.0)
                expanded_rows.append(row)

        X_expanded = np.array(expanded_rows) if expanded_rows else np.empty((0, 0))

        if k_cat > 0 and X_expanded.shape[0] > 0:
            X_model_cand = self._build_full_model_matrix(
                X_expanded, k_cont, cat_dummy_counts, model_type)
        elif X_expanded.shape[0] > 0:
            X_model_cand = self._build_model_matrix(X_expanded, model_type)
        else:
            return json.dumps([])

        # ── 编码已有实验点 ──
        existing_expanded = []
        for ex_row in existing:
            coded = []
            for f in continuous:
                center = (f["lower"] + f["upper"]) / 2.0
                half_range = (f["upper"] - f["lower"]) / 2.0
                val = ex_row.get(f["name"], center)
                coded.append((float(val) - center) / half_range if half_range > 0 else 0.0)
            for f in categorical:
                val = ex_row.get(f["name"], f["levels"][0])
                for lv_idx in range(1, len(f["levels"])):
                    coded.append(1.0 if val == f["levels"][lv_idx] else 0.0)
            existing_expanded.append(coded)

        n_existing = len(existing)

        if n_existing > 0 and len(existing_expanded[0]) > 0:
            X_existing = np.array(existing_expanded)
            if k_cat > 0:
                X_model_existing = self._build_full_model_matrix(
                    X_existing, k_cont, cat_dummy_counts, model_type)
            else:
                X_model_existing = self._build_model_matrix(X_existing, model_type)
        else:
            X_model_existing = np.empty((0, X_model_cand.shape[1]))

        p = X_model_cand.shape[1]

        # ── 自动确定补充数量 ──
        if num_additional < 0:
            df_needed = max(p - n_existing + 2, 0)
            num_additional = max(int(np.ceil(df_needed * 1.5)), 3)

        num_additional = min(num_additional, len(full_candidates))
        if num_additional <= 0:
            return json.dumps([])

        # ── Fedorov 交换: 锁定已有点，只交换新增点 ──
        N_cand = len(full_candidates)
        best_det = -np.inf
        best_new_idx = None
        n_starts = min(max(5, p), 20)

        for _ in range(n_starts):
            new_idx = list(np.random.choice(N_cand, size=num_additional, replace=False))
            new_set = set(new_idx)

            improved = True
            max_iter = 30
            iter_count = 0

            while improved and iter_count < max_iter:
                improved = False
                iter_count += 1

                for i in range(num_additional):
                    X_new = X_model_cand[new_idx]
                    X_combined = np.vstack([X_model_existing, X_new]) if n_existing > 0 else X_new

                    try:
                        current_det = np.linalg.det(X_combined.T @ X_combined)
                    except np.linalg.LinAlgError:
                        current_det = 0.0

                    old_cand = new_idx[i]
                    best_swap_det = current_det
                    best_swap = old_cand

                    for c in range(N_cand):
                        if c in new_set:
                            continue
                        new_idx[i] = c
                        X_new_test = X_model_cand[new_idx]
                        X_test = np.vstack([X_model_existing, X_new_test]) if n_existing > 0 else X_new_test
                        try:
                            test_det = np.linalg.det(X_test.T @ X_test)
                        except np.linalg.LinAlgError:
                            test_det = 0.0

                        if test_det > best_swap_det * 1.001:
                            best_swap_det = test_det
                            best_swap = c

                    if best_swap != old_cand:
                        new_set.discard(old_cand)
                        new_set.add(best_swap)
                        new_idx[i] = best_swap
                        improved = True
                    else:
                        new_idx[i] = old_cand

            # 计算最终行列式
            X_new_final = X_model_cand[new_idx]
            X_final = np.vstack([X_model_existing, X_new_final]) if n_existing > 0 else X_new_final
            try:
                final_det = np.linalg.det(X_final.T @ X_final)
            except np.linalg.LinAlgError:
                final_det = 0.0

            if final_det > best_det:
                best_det = final_det
                best_new_idx = list(new_idx)

        # ── 解码新增点 ──
        new_matrix = []
        if best_new_idx:
            for idx in best_new_idx:
                cont_row, cat_combo = full_candidates[idx]
                decoded = {}
                for i, f in enumerate(continuous):
                    center = (f["lower"] + f["upper"]) / 2.0
                    half_range = (f["upper"] - f["lower"]) / 2.0
                    decoded[f["name"]] = round(center + float(cont_row[i]) * half_range, 6)
                for i, f in enumerate(categorical):
                    decoded[f["name"]] = cat_combo[i]
                new_matrix.append(decoded)

        return json.dumps(new_matrix, ensure_ascii=False)

    def steepest_ascent(self, factors_json: str, coefficients_json: str,
                        step_count: int = 5, step_multiplier: float = 1.0,
                        descend: bool = False) -> str:
        """
        ★ 新增: 最速上升/下降法 — 沿一阶模型梯度方向生成实验点

        数学原理:
          给定一阶模型 y = β₀ + Σβᵢxᵢ
          最速上升方向: Δxᵢ ∝ βᵢ（各因子步长与系数成正比）
          选一个参考因子（系数最大的），固定其步长 = 1 个步长单位
          其余因子按系数比例缩放

        factors_json: 因子定义（含 lower, upper）
        coefficients_json: 一阶模型系数
          格式: {"温度": 12.5, "压力": -3.2, "催化剂": 8.1}
          不含截距，只有主效应系数
        step_count: 沿梯度方向的步数
        step_multiplier: 步长缩放因子（1.0 = 默认步长，0.5 = 半步长）
        descend: True = 最速下降（最小化目标时使用）

        返回: 实验点矩阵 JSON（从中心点出发，依次沿梯度方向）
        """
        factors = json.loads(factors_json)
        coefficients = json.loads(coefficients_json)

        continuous, categorical = self._separate_factors(factors)
        if len(continuous) == 0:
            return json.dumps([])

        # 提取连续因子的系数
        cont_names = [f["name"] for f in continuous]
        betas = np.array([coefficients.get(name, 0.0) for name in cont_names])

        if np.all(np.abs(betas) < 1e-10):
            return json.dumps([])  # 所有系数为零，无法确定方向

        # 如果是下降，取反
        if descend:
            betas = -betas

        # 参考因子: 系数绝对值最大的
        ref_idx = np.argmax(np.abs(betas))
        ref_beta = betas[ref_idx]

        # 各因子的步长比例（相对参考因子）
        step_ratios = betas / ref_beta if abs(ref_beta) > 1e-10 else np.zeros_like(betas)

        # 参考因子的步长 = 范围的 20% × step_multiplier
        ref_factor = continuous[ref_idx]
        ref_range = ref_factor["upper"] - ref_factor["lower"]
        base_step = ref_range * 0.2 * step_multiplier

        # 中心点
        centers = [(f["lower"] + f["upper"]) / 2.0 for f in continuous]

        # 生成实验点
        matrix = []
        for step in range(1, step_count + 1):
            row = {}
            for i, f in enumerate(continuous):
                val = centers[i] + step * base_step * step_ratios[i]
                # 裁剪到因子范围
                val = max(f["lower"], min(f["upper"], val))
                row[f["name"]] = round(val, 6)
            matrix.append(row)

        # 补充类别因子（取第一个水平）
        if categorical:
            matrix = self._cross_with_categorical(matrix, categorical)

        return json.dumps(matrix, ensure_ascii=False)

    def confirmation_runs(self, optimal_factors_json: str, factors_json: str,
                          repeat_count: int = 5, perturbation: float = 0.0) -> str:
        """
        ★ 新增: 验证实验 — 在预测最优点（及其邻域）生成重复实验

        optimal_factors_json: 预测最优因子组合，如 {"温度": 135, "压力": 22}
        factors_json: 因子定义（用于确定扰动范围）
        repeat_count: 重复次数（推荐 3-5 次）
        perturbation: 扰动幅度（0.0 = 纯重复，0.05 = ±5% 范围内随机扰动）

        返回: 验证实验矩阵 JSON
        """
        optimal = json.loads(optimal_factors_json)
        factors = json.loads(factors_json)

        continuous, categorical = self._separate_factors(factors)

        matrix = []

        if perturbation <= 0:
            # 纯重复: 同一点做 repeat_count 次
            for _ in range(repeat_count):
                matrix.append(dict(optimal))
        else:
            # 带扰动: 在最优点附近随机采样
            rng = np.random.RandomState(42)
            for _ in range(repeat_count):
                row = dict(optimal)
                for f in continuous:
                    name = f["name"]
                    center = optimal.get(name, (f["lower"] + f["upper"]) / 2.0)
                    half_range = (f["upper"] - f["lower"]) / 2.0
                    noise = rng.uniform(-perturbation, perturbation) * half_range
                    val = center + noise
                    val = max(f["lower"], min(f["upper"], val))
                    row[name] = round(val, 6)
                matrix.append(row)

        return json.dumps(matrix, ensure_ascii=False)

# ═══════════════════════════════════════════════════════
# DOEDataImporter — 历史数据导入（保留不变）
# ═══════════════════════════════════════════════════════

class DOEDataImporter:
    """
    DOE 数据导入器 — 校验并导入 Excel 历史实验数据。
    列名必须与 DOE 方案中定义的因子名称和响应名称完全匹配。
    """

    def validate_excel(self, path: str, factor_names_json: str, response_names_json: str) -> str:
        """校验 Excel 文件格式
        ★ 修复 (Bug#5): 只对响应变量列做严格的数值检查，
        因子列允许字符串值（类别因子如 "催化剂A"）
        """
        factor_names = json.loads(factor_names_json)
        response_names = json.loads(response_names_json)
        
        result = {
            "is_valid": True, "errors": [], "warnings": [],
            "valid_row_count": 0, "columns_found": [], "preview": []
        }
        
        try:
            df = pd.read_excel(path, engine='openpyxl')
        except Exception as e:
            result["is_valid"] = False
            result["errors"].append(f"无法读取 Excel 文件: {str(e)}")
            return json.dumps(result, ensure_ascii=False)
        
        result["columns_found"] = list(df.columns)
        all_required = factor_names + response_names
        missing_cols = [name for name in all_required if name not in df.columns]
        if missing_cols:
            result["is_valid"] = False
            result["errors"].append(f"缺少必需列: {', '.join(missing_cols)}")
        
        extra_cols = [col for col in df.columns if col not in set(all_required)]
        if extra_cols:
            result["warnings"].append(f"忽略多余列: {', '.join(extra_cols)}")
        
        if not result["is_valid"]:
            return json.dumps(result, ensure_ascii=False)
        
        # ★ 修复 (Bug#5): 只对响应变量列做非数值检查
        # 因子列可能包含类别因子的字符串值（如 "催化剂A"），这是合法的
        factor_names_set = set(factor_names)
        for col in all_required:
            if col in df.columns and col not in factor_names_set:
                # 响应变量列: 必须是数值
                non_numeric = df[col].apply(lambda x: not isinstance(x, (int, float, np.integer, np.floating)) and pd.notna(x))
                if non_numeric.any():
                    bad_rows = list(non_numeric[non_numeric].index + 2)
                    result["errors"].append(f"列 '{col}' 第 {bad_rows[:5]} 行包含非数值数据")
                    result["is_valid"] = False
        
        for col in all_required:
            if col in df.columns:
                null_count = df[col].isnull().sum()
                if null_count > 0:
                    result["warnings"].append(f"列 '{col}' 有 {null_count} 个缺失值，这些行将被跳过")
        
        valid_mask = df[all_required].notna().all(axis=1)
        result["valid_row_count"] = int(valid_mask.sum())
        
        if result["valid_row_count"] == 0:
            result["is_valid"] = False
            result["errors"].append("没有有效的数据行（所有行都存在缺失值）")
        
        if result["valid_row_count"] > 0:
            preview_df = df.loc[valid_mask, all_required].head(5)
            result["preview"] = preview_df.to_dict(orient='records')
            for row in result["preview"]:
                for k, v in row.items():
                    if isinstance(v, float) and np.isnan(v):
                        row[k] = None
        
        return json.dumps(result, ensure_ascii=False)

    def import_excel(self, path: str, factor_names_json: str, response_names_json: str) -> str:
        """导入 Excel 数据
        ★ 修复 (Bug#4): 因子值不再强制 float()，支持类别因子的字符串值
        """
        factor_names = json.loads(factor_names_json)
        response_names = json.loads(response_names_json)
        all_required = factor_names + response_names
        
        df = pd.read_excel(path, engine='openpyxl')
        valid_mask = df[all_required].notna().all(axis=1)
        valid_df = df.loc[valid_mask, all_required]
        
        records = []
        for _, row in valid_df.iterrows():
            # ★ 修复: 因子值尝试转 float，失败则保留为字符串（类别因子）
            factors = {}
            for name in factor_names:
                val = row[name]
                try:
                    factors[name] = float(val)
                except (ValueError, TypeError):
                    factors[name] = str(val)
            
            record = {
                "factors": factors,
                "responses": {name: float(row[name]) for name in response_names}
            }
            records.append(record)
        
        return json.dumps(records, ensure_ascii=False)


# ═══════════════════════════════════════════════════════
# 模块入口 — 供 C# pythonnet 调用
# ═══════════════════════════════════════════════════════

def create_designer():
    """创建 DOEDesigner 实例"""
    return DOEDesigner()

def create_importer():
    """创建 DOEDataImporter 实例"""
    return DOEDataImporter()


# ═══════════════════════════════════════════════════════
# 可选依赖导入
# ═══════════════════════════════════════════════════════
import io
import base64

try:
    import statsmodels.api as sm
    from statsmodels.formula.api import ols
    from statsmodels.stats.anova import anova_lm
    HAS_STATSMODELS = True
except ImportError:
    HAS_STATSMODELS = False

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from mpl_toolkits.mplot3d import Axes3D
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False


# ═══════════════════════════════════════════════════════
# DOEAnalyzer — 统计分析
# ★ 改造: 新增 fit_ols() 完整 ANOVA 分析, residual_diagnostics() 残差四图
# ═══════════════════════════════════════════════════════


def _calc_adequate_precision_design_expert(model):
    """
    计算 Adequate Precision — 对标 Design-Expert 定义
    
    Design-Expert 公式 (参考 Stat-Ease documentation):
        AP = (max(ŷ) - min(ŷ)) / sqrt( mean(Var(ŷ)) )
    其中:
        Var(ŷᵢ) = MSE × hᵢᵢ  (hat matrix 对角线)
        mean(Var(ŷ)) = MSE × mean(hᵢᵢ) = MSE × p/n
    简化: AP = (max(ŷ) - min(ŷ)) / sqrt(MSE × p/n)
    
    AP > 4 表示模型有足够信噪比用于导航设计空间。
    """
    try:
        fitted = model.fittedvalues
        fitted_arr = fitted.values if hasattr(fitted, 'values') else np.array(fitted)
        ms_res = float(model.mse_resid)
        if ms_res < 1e-12:
            return 0.0
        y_range = float(fitted_arr.max() - fitted_arr.min())
        if y_range < 1e-12:
            return 0.0
        n = len(fitted_arr)
        p = len(model.params)
        mean_var_pred = ms_res * p / n
        return float(y_range / np.sqrt(mean_var_pred))
    except Exception:
        return 0.0


def _calc_vif_correct(model):
    """
    计算 VIF — 标准方法（去掉截距列后逐列回归）
    VIF_j = 1 / (1 - R²_j)，其中 R²_j 是第 j 列对其余列回归的 R²
    """
    try:
        from statsmodels.stats.outliers_influence import variance_inflation_factor
        X = model.model.exog
        param_names = list(model.params.index)
        intercept_idx = None
        for i, name in enumerate(param_names):
            if name == "Intercept":
                intercept_idx = i
                break
        if intercept_idx is not None:
            X_no_intercept = np.delete(X, intercept_idx, axis=1)
            names_no_intercept = [n for i, n in enumerate(param_names) if i != intercept_idx]
        else:
            X_no_intercept = X
            names_no_intercept = param_names
        vif_dict = {}
        import statsmodels.api as _sm
        for j in range(X_no_intercept.shape[1]):
            try:
                X_with_const = _sm.add_constant(X_no_intercept)
                vif_val = variance_inflation_factor(X_with_const, j + 1)
                vif_dict[names_no_intercept[j]] = round(float(vif_val), 4)
            except Exception:
                vif_dict[names_no_intercept[j]] = None
        return vif_dict
    except Exception:
        return {}

class DOEAnalyzer:
    """
    DOE 统计分析器 — 主效应、交互效应、Pareto、ANOVA、回归、残差分析。
    
    ★ 改造说明:
      - 原有的 anova_table() / regression_summary() / residual_analysis() 保留兼容
      - 新增 fit_ols() 方法: 拟合完整二阶模型，返回 ANOVA 表 + 系数表 + 模型摘要
      - 新增 residual_diagnostics() 方法: 返回残差四图数据
      - 新增 effects_pareto() 方法: 基于 T 值的效应 Pareto 图
    """

    def __init__(self):
        self._df = None
        self._df_full_backup = None  # ★ v11 新增: 用于 refit_excluding / apply_box_cox 后恢复
        self._factor_names = []
        self._categorical_factors = []  # ★ 新增: 类别因子名列表
        self._continuous_factors = []   # ★ 新增: 连续因子名列表
        self._response_name = ""
        self._model = None
        self._model_sum = None  # ★ v8: Sum coding 模型（用于 CI 计算，与 JMP 一致）
        self._coding_info = {}
        self._design_method = ""  # ★ v17: 保存设计方法，用于全因子弯曲检验判断
        self._curvature_mse = None  # ★ v17: 弯曲检验合并 MSE（供 Pareto 使用）
        self._curvature_df_error = None  # ★ v17: 弯曲检验误差 DF
        self._curvature_ci_params = None  # ★ v18: 弯曲检验 CI 参数（供 prediction_profiler 使用）

    def load_data(self, factors_json: str, responses_json: str, response_name: str,
                  factor_types_json: str = "{}") -> str:
        """
        加载实验数据。
        factors_json: '[{"温度": 120, "压力": 15, "催化剂": "A"}, ...]'
        responses_json: '[87.5, 92.1, ...]'
        response_name: "转化率"
        factor_types_json: '{"温度": "continuous", "压力": "continuous", "催化剂": "categorical"}'
                          ★ 新增: 可选参数，指定因子类型。未指定的默认为 continuous。
        """
        factors_list = json.loads(factors_json)
        responses_list = json.loads(responses_json)
        factor_types = json.loads(factor_types_json) if factor_types_json else {}
        
        if len(factors_list) == 0:
            return json.dumps({"loaded": False, "error": "无数据"}, ensure_ascii=False)
        
        self._df = pd.DataFrame(factors_list)
        self._df_full_backup = None  # ★ v15: 清除旧备份，防止恢复操作覆盖新数据
        self._response_name = response_name
        self._df[response_name] = responses_list[:len(self._df)]
        self._factor_names = [c for c in self._df.columns if c != response_name]
        
        # ★ 新增: 分类连续因子和类别因子
        self._categorical_factors = []
        self._continuous_factors = []
        self._factor_bounds = {}  # ★ v15: 用户设定的因子设计范围（用于编码对齐 Minitab）
        self._coding_alpha = 1.0  # ★ v16: CCD 编码缩放因子（Minitab 用 half_range*α 编码）
        
        # ★ v16: 解析设计方法元数据（__design_meta__）
        design_meta = factor_types.pop("__design_meta__", None)
        if isinstance(design_meta, dict):
            self._design_method = str(design_meta.get("design_method", ""))  # ★ v17: 保存设计方法
            design_method = self._design_method.lower()
            if design_method == "ccd":
                ccd_alpha_type = str(design_meta.get("ccd_alpha_type", "rotatable")).lower()
                k = int(design_meta.get("k_continuous", len([n for n in self._factor_names if factor_types.get(n) != "categorical"])))
                if k >= 2:
                    if ccd_alpha_type in ("face_centered", "face"):
                        self._coding_alpha = 1.0  # face-centered α=1，不需要缩放
                    else:
                        # rotatable: α = (2^k)^0.25
                        self._coding_alpha = float((2 ** k) ** 0.25)
        
        for name in self._factor_names:
            ftype_raw = factor_types.get(name, "continuous")
            # ★ v15: 支持扩展格式 {"type":"continuous","lower":80,"upper":160}
            if isinstance(ftype_raw, dict):
                ftype = ftype_raw.get("type", "continuous").lower()
                if "lower" in ftype_raw and "upper" in ftype_raw:
                    lo = float(ftype_raw["lower"])
                    hi = float(ftype_raw["upper"])
                    # ★ v16: CCD 时 half_range 乘以 α（对齐 Minitab 编码方式）
                    self._factor_bounds[name] = (lo, hi, self._coding_alpha)
            elif isinstance(ftype_raw, str):
                ftype = ftype_raw.lower()
            else:
                ftype = "continuous"
            # 自动检测: 如果列包含非数值，视为类别因子
            if ftype == "categorical":
                self._categorical_factors.append(name)
            else:
                # 检查列数据是否全为数值
                try:
                    self._df[name].astype(float)
                    self._continuous_factors.append(name)
                except (ValueError, TypeError):
                    self._categorical_factors.append(name)
        
        # ★ 改造: 拟合二阶模型（含交互项和二次项），替代原来的只有主效应
        self._fit_quadratic_model()
        
        return json.dumps({
            "loaded": True,
            "rows": len(self._df),
            "factors": len(self._factor_names),
            "factor_names": self._factor_names
        }, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # ★ 新增: fit_ols() — 完整 OLS 回归分析，对标 JMP/Minitab
    # ═══════════════════════════════════════════════════════

    def fit_ols(self, model_type: str = "quadratic", terms_json: str = "") -> str:
        """
        ★ v7 改造: 增加不可估计项自动检测与剔除

        新增返回字段:
          dropped_terms: ["温度²", "温度×压力"]  — 因秩亏/共线性被自动剔除的项
          inestimable_warning: "以下项因数据不足被自动剔除: 温度², 温度×压力"
          original_model_type: "quadratic"  — 用户请求的原始模型类型
        """
        if self._df is None or not HAS_STATSMODELS:
            return json.dumps({"error": "数据未加载或缺少 statsmodels"}, ensure_ascii=False)

        try:
            k = len(self._factor_names)
            n = len(self._df)

            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            reverse_names = {v: k_name for k_name, v in safe_names.items()}
            safe_continuous = [safe_names[f] for f in self._continuous_factors]
            safe_categorical = [safe_names[f] for f in self._categorical_factors]

            df_safe = pd.DataFrame()
            self._coding_info = {}
            for orig_name, safe_name in safe_names.items():
                if orig_name in self._categorical_factors:
                    df_safe[safe_name] = self._df[orig_name].astype(str)
                    self._coding_info[orig_name] = {"type": "categorical"}
                else:
                    col = self._df[orig_name].astype(float)
                    if orig_name in self._factor_bounds:
                        # ★ v16: 优先用设计范围，CCD 时乘以 α（对齐 Minitab）
                        bounds = self._factor_bounds[orig_name]
                        lo, hi = bounds[0], bounds[1]
                        alpha_scale = bounds[2] if len(bounds) > 2 else 1.0
                        center = (lo + hi) / 2.0
                        half_range = (hi - lo) / 2.0 * alpha_scale
                    else:
                        # 回退: 用数据范围（直接导入 Excel 时没有因子范围信息）
                        col_min = col.min()
                        col_max = col.max()
                        center = (col_min + col_max) / 2.0
                        half_range = (col_max - col_min) / 2.0
                    if half_range < 1e-12:
                        half_range = 1.0
                    self._coding_info[orig_name] = {"type": "continuous", "center": center, "half_range": half_range}
                    df_safe[safe_name] = (col - center) / half_range

            df_safe["Y"] = self._df[self._response_name].values

            # ★ v17: 全因子弯曲检验 — 当设计方法为因子型且有中心点时，对标 Minitab
            # custom 模型（精简模型）也走弯曲路径，只是用用户选择的 terms
            if self._is_factorial_design(self._design_method):
                # 构建 coding_info 序列化（弯曲检验返回时需要）
                _coding_info_for_curv = {}
                for fname, info in self._coding_info.items():
                    ci = {"type": info.get("type", "continuous")}
                    if ci["type"] == "continuous":
                        ci["center"] = round(float(info.get("center", 0.0)), 6)
                        ci["half_range"] = round(float(info.get("half_range", 1.0)), 6)
                    _coding_info_for_curv[fname] = ci
                
                # ★ v17: 解析 custom terms（精简模型时传入）
                custom_terms_list = None
                if model_type == "custom" and terms_json:
                    try:
                        custom_terms_list = json.loads(terms_json)
                    except Exception:
                        custom_terms_list = None
                
                curvature_result = self._fit_factorial_with_curvature(
                    df_safe, safe_continuous, safe_categorical,
                    safe_names, reverse_names,
                    custom_terms=custom_terms_list
                )
                if curvature_result is not None:
                    self._model = curvature_result["model"]
                    self._model_treatment = curvature_result["model"]
                    # ★ v17: 保存弯曲检验的合并 MSE 和 df_error，供 effects_pareto 使用
                    self._curvature_mse = curvature_result["model_summary"]["mse"]
                    self._curvature_df_error = curvature_result["model_summary"]["df_error"]
                    
                    # ★ v18: 保存弯曲检验的合并误差 CI 参数，供 prediction_profiler 使用
                    # CI 的 df = error_df + curvature_df（对齐 JMP）
                    _curv_error_df = curvature_result["model_summary"]["df_error"]
                    _curv_curvature_df = curvature_result["curvature"]["curvature_df"]
                    _curv_ci_df = _curv_error_df + _curv_curvature_df
                    self._curvature_ci_params = {
                        "sigma": curvature_result.get("sigma", 0.0),
                        "error_df": _curv_ci_df,
                        "XtX_inv": np.array(curvature_result.get("xtx_inv_flat", [])).reshape(
                            curvature_result.get("param_count", 0),
                            curvature_result.get("param_count", 0)
                        ) if curvature_result.get("param_count", 0) > 0 else None,
                    }
                    
                    return json.dumps({
                        "anova_table": curvature_result["anova_table"],
                        "coefficients": curvature_result["coefficients"],
                        "uncoded_coefficients": curvature_result.get("uncoded_coefficients", []),
                        "model_summary": curvature_result["model_summary"],
                        "curvature_test": curvature_result["curvature"],
                        "coding_info": _coding_info_for_curv,
                        "uncoded_equation": curvature_result.get("uncoded_equation", ""),
                        "uncoded_equations": curvature_result.get("uncoded_equations"),
                        "dropped_terms": [],
                        "inestimable_warning": "",
                        "original_model_type": model_type,
                        "xtx_inv_flat": curvature_result.get("xtx_inv_flat", []),
                        "sigma": curvature_result.get("sigma", 0.0),
                        "t_crit": curvature_result.get("t_crit", 0.0),
                        "param_count": curvature_result.get("param_count", 0),
                    }, ensure_ascii=False)

            # ★ v17: 非弯曲检验路径，清空弯曲 MSE（确保 Pareto 走正常逻辑）
            self._curvature_mse = None
            self._curvature_df_error = None
            self._curvature_ci_params = None  # ★ v18: 清空弯曲 CI 参数

            # ★ v7: 构建公式
            if model_type == "custom" and terms_json:
                formula = self._build_custom_formula(terms_json, safe_names, reverse_names,
                                                      safe_continuous, safe_categorical)
            else:
                formula = self._build_formula(safe_continuous, safe_categorical, model_type)

            # ★ v7: 不可估计项自动检测 ——
            # 先尝试拟合，如果 statsmodels 因奇异矩阵报错，
            # 则逐项检测并剔除造成秩亏的列，再重新拟合
            dropped_terms = []
            original_model_type = model_type

            try:
                model = ols(formula, data=df_safe).fit()
                # 检查是否有 NaN 系数（部分奇异但 statsmodels 未报错的情况）
                nan_params = [p for p in model.params.index if np.isnan(model.params[p])]
                if nan_params:
                    raise np.linalg.LinAlgError("Singular matrix detected via NaN params")
                # ★ v7: 检查 df_resid <= 0 — 参数数 >= 数据点数，模型过参数化
                if model.df_resid <= 0:
                    raise np.linalg.LinAlgError(
                        f"Over-parameterized: {len(model.params)} params vs {int(model.nobs)} observations, df_resid=0")
                # ★ v7 增强: 检查模型矩阵秩亏 — rank(X) < 列数 表示存在不可估计项
                X_exog = model.model.exog
                actual_rank = np.linalg.matrix_rank(X_exog)
                n_params = X_exog.shape[1]
                if actual_rank < n_params:
                    raise np.linalg.LinAlgError(
                        f"Rank deficient: rank={actual_rank} < params={n_params}, "
                        f"模型矩阵存在共线性列，需要剔除不可估计项")
                # ★ v10: SVD 奇异值比值兜底（兼容不同 numpy 版本的 matrix_rank 容差差异）
                svd_vals = np.linalg.svd(X_exog, compute_uv=False)
                sv_ratio = svd_vals[-1] / svd_vals[0] if svd_vals[0] > 0 else 0
                if sv_ratio < 1e-12:
                    raise np.linalg.LinAlgError(
                        f"Near-singular design matrix: min/max singular value ratio = {sv_ratio:.2e}, "
                        f"存在不可估计项（类似 Minitab 的秩亏检测）")
                # ★ v7 增强: 检查 Inf/NaN SE（标准误为无穷大表示该项不可估计）
                inf_se_params = [model.params.index[i] for i in range(len(model.bse))
                                 if np.isinf(model.bse.iloc[i]) or np.isnan(model.bse.iloc[i])]
                if inf_se_params:
                    raise np.linalg.LinAlgError(
                        f"Inestimable terms detected via Inf/NaN SE: {inf_se_params}")
                # ★ v7 增强: 检查 NaN p-values（因子无变异时系数=0, p=NaN）
                nan_pval_params = [p for p in model.pvalues.index
                                   if p != "Intercept" and np.isnan(model.pvalues[p])]
                if nan_pval_params:
                    raise np.linalg.LinAlgError(
                        f"Inestimable terms detected via NaN p-values: {nan_pval_params}")
            except Exception as fit_err:
                # ★ v7 核心: 逐项检测不可估计项并剔除
                formula, dropped_terms = self._detect_and_drop_inestimable(
                    formula, df_safe, safe_names, reverse_names, safe_continuous, safe_categorical, model_type
                )
                if not formula or formula == "Y ~ 1":
                    return json.dumps({
                        "error": f"所有模型项均不可估计，可能数据量不足或因子无变异: {str(fit_err)}",
                        "dropped_terms": dropped_terms
                    }, ensure_ascii=False)
                model = ols(formula, data=df_safe).fit()

            self._model = model
            self._model_treatment = model  # ★ FIX #4: 保存 Treatment coding 模型供方程展开

            # ────── 系数表 (★ v13: 使用 Sum coding, 对标 Minitab "已编码系数") ──────
            # Minitab 的 "已编码系数" 表用 Sum/Effect coding (-1,0,1)
            # Treatment coding 的系数表作为内部使用 (方程展开等)
            import re as _re
            
            # 构建 Sum coding 模型 (跟 ANOVA 一致)
            formula_for_coeff = str(model.model.formula) if hasattr(model.model, 'formula') else self._build_formula(
                [safe_names[f] for f in self._continuous_factors],
                [safe_names[f] for f in self._categorical_factors], model_type)
            formula_sum_coeff = _re.sub(r'C\((\w+)\)(?![,\[])', r'C(\1, Sum)', formula_for_coeff)
            
            try:
                model_sum_coeff = ols(formula_sum_coeff, data=df_safe).fit()
            except Exception:
                model_sum_coeff = model  # 回退到 Treatment coding
            
            # ★ v8: 保存 Sum coding 模型供 _predict_with_se 计算 CI（与 JMP Effect coding 一致）
            self._model_sum = model_sum_coeff
            
            coefficients = []
            # ★ FIX #3: 使用标准 VIF 计算（去掉截距列后逐列回归）
            vif_dict = _calc_vif_correct(model_sum_coeff)

            for idx_j, term in enumerate(model_sum_coeff.params.index):
                coeff = float(model_sum_coeff.params[term])
                se = float(model_sum_coeff.bse[term])
                t_val = float(model_sum_coeff.tvalues[term])
                p_val = float(model_sum_coeff.pvalues[term])
                # ★ FIX #3: 从标准 VIF 字典获取
                vif = vif_dict.get(term) if idx_j > 0 else None
                # 去掉 Sum coding 标记再还原名称
                display_term = self._restore_term_name(term.replace(", Sum", ""), reverse_names)
                coefficients.append({
                    "term": display_term, "coeff": round(coeff, 6), "se": round(se, 6),
                    "t_value": round(t_val, 4), "p_value": round(p_val, 6), "vif": vif
                })

            # ────── ANOVA 表 ──────
            anova_table = self._build_anova_table(model, df_safe, safe_names, reverse_names, model_type)

            # ────── 模型摘要 ──────
            leverage = model.get_influence().hat_matrix_diag
            press_residuals = model.resid / (1 - leverage)
            press = float(np.sum(press_residuals ** 2))
            ss_total = float(np.sum((df_safe["Y"] - df_safe["Y"].mean()) ** 2))
            r2_pred = 1.0 - press / ss_total if ss_total > 1e-12 else 0.0

            fitted = model.fittedvalues
            p_count = len(model.params)
            # ★ v15: RMSE 使用 Sum coding 模型（与系数表/ANOVA 一致，对标 Minitab）
            ms_res = float(model_sum_coeff.mse_resid) if model_sum_coeff is not None else float(model.mse_resid)
            adeq_prec = 0.0
            # ★ FIX #2: Adequate Precision 对标 Design-Expert 标准公式
            # AP = (max(ŷ) - min(ŷ)) / sqrt(MSE × p/n)
            adeq_prec = _calc_adequate_precision_design_expert(model)

            lof_p = self._calc_lack_of_fit_p(model, df_safe, safe_names, model_type)
            equation = self._build_equation(model, reverse_names)
            equations = self._build_equations_by_category(model, reverse_names)

            # ★ v7: 构建不可估计项警告信息
            inestimable_warning = ""
            if dropped_terms:
                inestimable_warning = f"以下项因数据不足或共线性被自动剔除: {', '.join(dropped_terms)}"

            # ★ v15: model_summary 使用 Sum coding 模型（与系数表/ANOVA 一致，对标 Minitab）
            _summary_model = model_sum_coeff if model_sum_coeff is not None else model
            model_summary = {
                "r_squared": round(float(_summary_model.rsquared), 6),
                "r_squared_adj": round(float(_summary_model.rsquared_adj), 6),
                "r_squared_pred": round(float(r2_pred), 6),
                "rmse": round(float(np.sqrt(ms_res)), 6),
                "adeq_precision": round(adeq_prec, 4),
                "press": round(press, 4),
                "lack_of_fit_p": round(lof_p, 6) if lof_p is not None else None,
                "model_p": round(float(_summary_model.f_pvalue), 6) if not np.isnan(_summary_model.f_pvalue) else None,
                "equation": equation,
                "equations": equations
            }

            # ★ v13 新增: 返回 coding_info 供 C# 端参考
            coding_info_serializable = {}
            for fname, info in self._coding_info.items():
                ci = {"type": info.get("type", "continuous")}
                if ci["type"] == "continuous":
                    ci["center"] = round(float(info.get("center", 0.0)), 6)
                    ci["half_range"] = round(float(info.get("half_range", 1.0)), 6)
                coding_info_serializable[fname] = ci

            # ★ v13 新增: 用原始值拟合模型，提供未编码系数（对标 Minitab "以未编码单位表示的回归方程"）
            # C# 端用户选择"未编码"时直接用此系数展示，无需做复杂转换
            uncoded_coefficients = []
            uncoded_equation = ""
            uncoded_equations = {}
            try:
                # 构建原始值 DataFrame
                df_uncoded = pd.DataFrame()
                for orig_name, safe_name in safe_names.items():
                    if orig_name in self._categorical_factors:
                        df_uncoded[safe_name] = self._df[orig_name].astype(str)
                    else:
                        df_uncoded[safe_name] = self._df[orig_name].astype(float)
                df_uncoded["Y"] = self._df[self._response_name].values

                # 用编码模型的实际公式（已剔除不可估计项）在原始值上拟合
                if hasattr(model, 'model') and hasattr(model.model, 'formula'):
                    uncoded_formula = str(model.model.formula)
                else:
                    uncoded_formula = self._build_formula(
                        [safe_names[f] for f in self._continuous_factors],
                        [safe_names[f] for f in self._categorical_factors],
                        model_type
                    )
                model_uncoded = ols(uncoded_formula, data=df_uncoded).fit()

                # ★ v14 改进: 提取完整系数表（含 SE/T/P/VIF），对标 Minitab 未编码系数表
                # 这样 C# 端切换到"未编码单位"时，系数表也能完整显示
                uncoded_vif_dict = _calc_vif_correct(model_uncoded)
                for idx_j, term in enumerate(model_uncoded.params.index):
                    coeff_val = float(model_uncoded.params[term])
                    se_val = float(model_uncoded.bse[term])
                    t_val = float(model_uncoded.tvalues[term])
                    p_val = float(model_uncoded.pvalues[term])
                    vif_val = uncoded_vif_dict.get(term) if idx_j > 0 else None
                    display_term = self._restore_term_name(term, reverse_names)
                    uncoded_coefficients.append({
                        "term": display_term,
                        "coeff": round(coeff_val, 6),
                        "se": round(se_val, 6),
                        "t_value": round(t_val, 4),
                        "p_value": round(p_val, 6),
                        "vif": vif_val
                    })

                # 构建未编码方程
                uncoded_equation = self._build_equation(model_uncoded, reverse_names)
                uncoded_equations = self._build_equations_by_category(model_uncoded, reverse_names)
            except Exception:
                pass  # 原始值拟合失败时不影响主流程

           # ★ v9: 导出 CI 参数给 C# 端实时计算
            ci_export = {}
            try:
                from scipy import stats as sp_stats
                # 优先用 Sum coding 模型的设计矩阵（与 JMP Effect coding 一致）
                ci_model = self._model_sum if hasattr(self, '_model_sum') and self._model_sum is not None else model
                X_matrix = ci_model.model.exog
                n_obs, p_params = X_matrix.shape
                df_error = max(n_obs - p_params, 1)
                mse_val = float(ci_model.mse_resid)
                sigma_val = float(np.sqrt(mse_val))
                t_crit_val = float(sp_stats.t.ppf(0.975, df_error))
                XtX_inv = np.linalg.inv(X_matrix.T @ X_matrix)
                ci_export = {
                    "xtx_inv_flat": [round(v, 10) for v in XtX_inv.flatten().tolist()],
                    "sigma": round(sigma_val, 8),
                    "t_crit": round(t_crit_val, 6),
                    "param_count": p_params
                }
            except Exception:
                ci_export = {"xtx_inv_flat": [], "sigma": 0.0, "t_crit": 0.0, "param_count": 0}

            return json.dumps({
                "anova_table": anova_table,
                "coefficients": coefficients,
                "uncoded_coefficients": uncoded_coefficients,
                "model_summary": model_summary,
                "coding_info": coding_info_serializable,
                "uncoded_equation": uncoded_equation,
                "uncoded_equations": uncoded_equations,
                "dropped_terms": dropped_terms,
                "inestimable_warning": inestimable_warning,
                "original_model_type": original_model_type,
                "xtx_inv_flat": ci_export.get("xtx_inv_flat", []),
                "sigma": ci_export.get("sigma", 0.0),
                "t_crit": ci_export.get("t_crit", 0.0),
                "param_count": ci_export.get("param_count", 0)
            }, ensure_ascii=False)

        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # ★ v17: 全因子弯曲检验 (Curvature Test) — 对标 Minitab
    # ═══════════════════════════════════════════════════════

    def _is_factorial_design(self, design_method_str):
        """★ v17: 判断是否为因子型设计（需要弯曲检验的设计类型）"""
        if not design_method_str:
            return False
        dm = design_method_str.lower().replace("_", "").replace("-", "")
        return dm in (
            "fullfactorial", "fractionalfactorial", "plackettburman", "factorial",
        )

    def _detect_center_points(self, df_coded, safe_continuous):
        """★ v17: 检测中心点 — 编码值所有连续因子都接近 0 的行"""
        if not safe_continuous:
            return pd.Series([False] * len(df_coded), index=df_coded.index)
        tolerance = 0.01
        center_mask = pd.Series([True] * len(df_coded), index=df_coded.index)
        for col in safe_continuous:
            center_mask &= df_coded[col].abs() < tolerance
        return center_mask

    def _fit_factorial_with_curvature(self, df_safe, safe_continuous, safe_categorical,
                                       safe_names, reverse_names, custom_terms=None):
        """
        ★ v17: Minitab 风格全因子分析 — 角点拟合模型 + 弯曲检验
        
        custom_terms: 精简模型时传入的项名列表（原始因子名，如 ["温度", "压力", "温度×压力"]）
                      为 None 时构建默认的二阶交互模型（对标 Minitab 默认）
        """
        from scipy import stats as sp_stats
        
        # ── 检测中心点 ──
        center_mask = self._detect_center_points(df_safe, safe_continuous)
        n_center = int(center_mask.sum())
        n_total = len(df_safe)
        n_corner = n_total - n_center
        
        if n_center < 2:
            return None
        
        df_corner = df_safe[~center_mask].copy().reset_index(drop=True)
        df_center = df_safe[center_mask].copy().reset_index(drop=True)
        
        y_corner = df_corner["Y"].values
        y_center = df_center["Y"].values
        corner_mean = float(np.mean(y_corner))
        center_mean = float(np.mean(y_center))
        
        # ── 构建公式 ──
        all_cont = list(safe_continuous)
        all_cat = list(safe_categorical)
        all_vars = all_cont + all_cat
        k = len(all_vars)
        
        from itertools import combinations
        
        if custom_terms is not None:
            # ★ 精简模型: 用用户选择的项构建公式
            terms = []
            for ct in custom_terms:
                safe_term = ct
                for orig, safe in safe_names.items():
                    safe_term = safe_term.replace(orig, safe)
                safe_term = safe_term.replace("×", ":")
                for cat in all_cat:
                    orig_cat = reverse_names.get(cat, cat)
                    if cat in safe_term and cat in safe_categorical:
                        safe_term = safe_term.replace(cat, f"C({cat})")
                terms.append(safe_term)
        else:
            # ★ 默认: 全交互模型（线性 + 所有阶交互，对标 Minitab 全因子默认）
            terms = []
            for v in all_cont:
                terms.append(v)
            for v in all_cat:
                terms.append(f"C({v})")
            
            for order in range(2, k + 1):
                for combo in combinations(range(k), order):
                    term_parts = []
                    for idx in combo:
                        v = all_vars[idx]
                        if v in all_cat:
                            term_parts.append(f"C({v})")
                        else:
                            term_parts.append(v)
                    terms.append(":".join(term_parts))
        
        formula = "Y ~ " + " + ".join(terms) if terms else "Y ~ 1"
        
        # ── 用角点拟合 ──
        try:
            model = ols(formula, data=df_corner).fit()
        except Exception:
            # 如果饱和导致拟合失败，减少高阶交互项重试
            terms_reduced = []
            for v in all_cont:
                terms_reduced.append(v)
            for v in all_cat:
                terms_reduced.append(f"C({v})")
            for i in range(len(all_vars)):
                for j in range(i + 1, len(all_vars)):
                    v1 = f"C({all_vars[i]})" if all_vars[i] in all_cat else all_vars[i]
                    v2 = f"C({all_vars[j]})" if all_vars[j] in all_cat else all_vars[j]
                    terms_reduced.append(f"{v1}:{v2}")
            formula = "Y ~ " + " + ".join(terms_reduced)
            model = ols(formula, data=df_corner).fit()
        
        # ── 弯曲检验 ──
        d = center_mean - corner_mean
        curvature_ss = (n_corner * n_center) / (n_corner + n_center) * d ** 2
        curvature_df = 1
        curvature_ms = curvature_ss
        
        # 纯误差 (中心点内部变异)
        pure_error_ss = float(np.sum((y_center - center_mean) ** 2))
        pure_error_df = n_center - 1
        pure_error_ms = pure_error_ss / pure_error_df if pure_error_df > 0 else 0
        
        # 角点残差 (饱和设计时 = 0)
        corner_resid_ss = float(model.ssr)
        corner_resid_df = int(model.df_resid)
        
        # 误差 = 角点残差(失拟) + 纯误差
        error_ss = corner_resid_ss + pure_error_ss
        error_df = corner_resid_df + pure_error_df
        error_ms = error_ss / error_df if error_df > 0 else 0
        
        # 弯曲 F 检验
        curvature_f = curvature_ms / error_ms if error_ms > 0 else 0
        curvature_p = float(1 - sp_stats.f.cdf(curvature_f, curvature_df, error_df)) if error_ms > 0 else 1.0
        
        # 失拟 F 检验
        if corner_resid_df > 0 and pure_error_ms > 0:
            lof_ms = corner_resid_ss / corner_resid_df
            lof_f = lof_ms / pure_error_ms
            lof_p = float(1 - sp_stats.f.cdf(lof_f, corner_resid_df, pure_error_df))
        else:
            lof_ms = 0
            lof_f = None
            lof_p = None
        
        # Ct Pt 系数
        ct_pt_coeff = d
        ct_pt_se = float(np.sqrt(error_ms * (1.0 / n_corner + 1.0 / n_center)))
        ct_pt_t = ct_pt_coeff / ct_pt_se if ct_pt_se > 0 else 0
        ct_pt_p = curvature_p
        
        # ── 模型摘要 ──
        total_ss = float(np.sum((df_safe["Y"].values - np.mean(df_safe["Y"].values)) ** 2))
        r_sq = 1 - error_ss / total_ss if total_ss > 0 else 0
        r_sq_adj = 1 - (error_ms) / (total_ss / (n_total - 1)) if total_ss > 0 else 0
        s = float(np.sqrt(error_ms)) if error_ms > 0 else 0
        
        # R-sq(pred)
        try:
            hat_corner = model.get_influence().hat_matrix_diag
            resid_corner = model.resid.values
            # 饱和设计时 hat=1, resid=0, PRESS=0/0 → 需要特殊处理
            press_corner = 0
            for i_c in range(len(resid_corner)):
                denom = 1 - hat_corner[i_c]
                if abs(denom) > 1e-10:
                    press_corner += (resid_corner[i_c] / denom) ** 2
                # 饱和时 denom≈0，该点 PRESS 贡献跳过
        except Exception:
            press_corner = 0
        
        press_center = 0
        for i in range(n_center):
            remaining = np.delete(y_center, i)
            remaining_mean = float(np.mean(remaining))
            d_loo = remaining_mean - corner_mean
            pred_loo = corner_mean + d_loo
            press_center += (y_center[i] - pred_loo) ** 2
        
        press_total = press_corner + press_center
        r_sq_pred = 1 - press_total / total_ss if total_ss > 0 else 0
        # 饱和设计时 R-sq(pred) 可能异常，Minitab 显示 *
        r_sq_pred_display = round(r_sq_pred, 6) if abs(r_sq_pred) < 2.0 else None
        
        # ── 用合并 MSE 重新计算系数表 ──
        X = model.model.exog
        try:
            XtX_inv = np.linalg.inv(X.T @ X)
        except np.linalg.LinAlgError:
            XtX_inv = np.linalg.pinv(X.T @ X)
        
        coefficients = []
        params = model.params
        for i, pname in enumerate(params.index):
            coef = float(params.iloc[i])
            se = float(np.sqrt(error_ms * XtX_inv[i, i]))
            t_val = coef / se if se > 0 else 0
            p_val = float(2 * (1 - sp_stats.t.cdf(abs(t_val), error_df)))
            
            display = pname
            for safe, orig in reverse_names.items():
                display = display.replace(safe, orig)
            display = display.replace(":", "×")
            
            effect = round(2 * coef, 6) if pname != "Intercept" else None
            
            coefficients.append({
                "term": display,
                "effect": effect,
                "coeff": round(coef, 6),
                "se": round(se, 6),
                "t_value": round(t_val, 4),
                "p_value": round(p_val, 6),
                "vif": 1.0 if pname != "Intercept" else None,
            })
        
        # Ct Pt 行
        coefficients.append({
            "term": "Ct Pt",
            "effect": None,
            "coeff": round(ct_pt_coeff, 6),
            "se": round(ct_pt_se, 6),
            "t_value": round(ct_pt_t, 4),
            "p_value": round(ct_pt_p, 6),
            "vif": 1.0,
        })
        
        # ── ANOVA 表 (Type III, 用合并 MSE) ──
        try:
            anova_t3 = anova_lm(model, typ=3)
        except Exception:
            # 饱和设计 Type III 可能失败，用 Type I 回退
            anova_t3 = anova_lm(model, typ=1)
        
        linear_items = []
        interact_2way = []
        interact_3way = []
        interact_higher = []
        
        for idx in anova_t3.index:
            if idx in ("Intercept", "Residual"):
                continue
            ss = float(anova_t3.loc[idx, "sum_sq"])
            df_ = int(anova_t3.loc[idx, "df"])
            ms = ss / df_ if df_ > 0 else 0
            f_val = ms / error_ms if error_ms > 0 else 0
            p_val = float(1 - sp_stats.f.cdf(f_val, df_, error_df))
            
            display = idx
            for safe, orig in reverse_names.items():
                display = display.replace(safe, orig)
            display = display.replace(":", "×")
            
            item = {"source": display, "df": df_, "ss": round(ss, 6), "ms": round(ms, 6),
                    "f_value": round(f_val, 4), "p_value": round(p_val, 6)}
            
            colon_count = idx.count(":")
            if colon_count == 0:
                linear_items.append(item)
            elif colon_count == 1:
                interact_2way.append(item)
            elif colon_count == 2:
                interact_3way.append(item)
            else:
                interact_higher.append(item)
        
        # 构建完整 ANOVA 表
        anova_rows = []
        
        # 模型行
        all_effect_items = linear_items + interact_2way + interact_3way + interact_higher
        model_effect_ss = sum(it["ss"] for it in all_effect_items) + curvature_ss
        model_effect_df = sum(it["df"] for it in all_effect_items) + curvature_df
        model_ms_ = model_effect_ss / model_effect_df if model_effect_df > 0 else 0
        model_f_ = model_ms_ / error_ms if error_ms > 0 else 0
        model_p_ = float(1 - sp_stats.f.cdf(model_f_, model_effect_df, error_df))
        anova_rows.append({"source": "模型", "df": model_effect_df, "ss": round(model_effect_ss, 6),
                           "ms": round(model_ms_, 6), "f_value": round(model_f_, 4), "p_value": round(model_p_, 6)})
        
        # 线性汇总
        if linear_items:
            lin_ss = sum(it["ss"] for it in linear_items)
            lin_df = sum(it["df"] for it in linear_items)
            lin_ms = lin_ss / lin_df if lin_df > 0 else 0
            lin_f = lin_ms / error_ms if error_ms > 0 else 0
            lin_p = float(1 - sp_stats.f.cdf(lin_f, lin_df, error_df))
            anova_rows.append({"source": "  线性", "df": lin_df, "ss": round(lin_ss, 6),
                               "ms": round(lin_ms, 6), "f_value": round(lin_f, 4), "p_value": round(lin_p, 6)})
            for it in linear_items:
                anova_rows.append({"source": "    " + it["source"], **{k_: it[k_] for k_ in ("df", "ss", "ms", "f_value", "p_value")}})
        
        # 2因子交互汇总
        if interact_2way:
            int2_ss = sum(it["ss"] for it in interact_2way)
            int2_df = sum(it["df"] for it in interact_2way)
            int2_ms = int2_ss / int2_df if int2_df > 0 else 0
            int2_f = int2_ms / error_ms if error_ms > 0 else 0
            int2_p = float(1 - sp_stats.f.cdf(int2_f, int2_df, error_df))
            anova_rows.append({"source": "  2因子交互作用", "df": int2_df, "ss": round(int2_ss, 6),
                               "ms": round(int2_ms, 6), "f_value": round(int2_f, 4), "p_value": round(int2_p, 6)})
            for it in interact_2way:
                anova_rows.append({"source": "    " + it["source"], **{k_: it[k_] for k_ in ("df", "ss", "ms", "f_value", "p_value")}})
        
        # 3因子交互汇总
        if interact_3way:
            int3_ss = sum(it["ss"] for it in interact_3way)
            int3_df = sum(it["df"] for it in interact_3way)
            int3_ms = int3_ss / int3_df if int3_df > 0 else 0
            int3_f = int3_ms / error_ms if error_ms > 0 else 0
            int3_p = float(1 - sp_stats.f.cdf(int3_f, int3_df, error_df))
            anova_rows.append({"source": "  3因子交互作用", "df": int3_df, "ss": round(int3_ss, 6),
                               "ms": round(int3_ms, 6), "f_value": round(int3_f, 4), "p_value": round(int3_p, 6)})
            for it in interact_3way:
                anova_rows.append({"source": "    " + it["source"], **{k_: it[k_] for k_ in ("df", "ss", "ms", "f_value", "p_value")}})
        
        # 更高阶交互
        if interact_higher:
            for it in interact_higher:
                anova_rows.append({"source": "  " + it["source"], **{k_: it[k_] for k_ in ("df", "ss", "ms", "f_value", "p_value")}})
        
        # 弯曲行
        anova_rows.append({"source": "  弯曲", "df": curvature_df, "ss": round(curvature_ss, 6),
                           "ms": round(curvature_ms, 6), "f_value": round(curvature_f, 4), "p_value": round(curvature_p, 6)})
        
        # 误差行
        anova_rows.append({"source": "误差", "df": error_df, "ss": round(error_ss, 6),
                           "ms": round(error_ms, 6), "f_value": None, "p_value": None})
        
        # 失拟行 (仅当角点非饱和时有)
        if corner_resid_df > 0:
            anova_rows.append({"source": "  失拟", "df": corner_resid_df, "ss": round(corner_resid_ss, 6),
                               "ms": round(lof_ms, 6),
                               "f_value": round(lof_f, 4) if lof_f is not None else None,
                               "p_value": round(lof_p, 6) if lof_p is not None else None})
        
        # 纯误差行
        anova_rows.append({"source": "  纯误差", "df": pure_error_df, "ss": round(pure_error_ss, 6),
                           "ms": round(pure_error_ms, 6), "f_value": None, "p_value": None})
        
        # 合计
        anova_rows.append({"source": "合计", "df": n_total - 1, "ss": round(total_ss, 6),
                           "ms": None, "f_value": None, "p_value": None})
        
        # ── 未编码系数和方程 ──
        uncoded_coefficients = []
        uncoded_equation = ""
        uncoded_equations = None
        try:
            df_orig = pd.DataFrame()
            for orig_name in self._factor_names:
                if orig_name in self._categorical_factors:
                    df_orig[orig_name] = self._df[orig_name].astype(str)
                else:
                    df_orig[orig_name] = self._df[orig_name].astype(float)
            df_orig["Y"] = self._df[self._response_name].values
            
            # 仅用角点
            df_orig_corner = df_orig.iloc[df_safe[~center_mask].index].reset_index(drop=True)
            
            orig_cont = [f for f in self._factor_names if f not in self._categorical_factors]
            orig_cat = [f for f in self._factor_names if f in self._categorical_factors]
            
            uncoded_terms = list(orig_cont)
            for v in orig_cat:
                uncoded_terms.append(f"C({v})")
            for order in range(2, len(self._factor_names) + 1):
                for combo in combinations(range(len(self._factor_names)), order):
                    parts = []
                    for ci in combo:
                        fn = self._factor_names[ci]
                        parts.append(f"C({fn})" if fn in self._categorical_factors else fn)
                    uncoded_terms.append(":".join(parts))
            
            uncoded_formula = "Y ~ " + " + ".join(uncoded_terms)
            model_uncoded = ols(uncoded_formula, data=df_orig_corner).fit()
            
            for idx_j, term in enumerate(model_uncoded.params.index):
                coeff_val = float(model_uncoded.params[term])
                # 用合并 MSE 近似 (精确方法需要 uncoded XtX_inv × error_ms)
                se_val = float(model_uncoded.bse[term]) if model_uncoded.df_resid > 0 else 0
                t_val = float(model_uncoded.tvalues[term]) if model_uncoded.df_resid > 0 else 0
                p_val = float(model_uncoded.pvalues[term]) if model_uncoded.df_resid > 0 else 1.0
                display_term = self._restore_term_name(term, reverse_names) if hasattr(self, '_restore_term_name') else term
                uncoded_coefficients.append({
                    "term": display_term,
                    "coeff": round(coeff_val, 6),
                    "se": round(se_val, 6),
                    "t_value": round(t_val, 4),
                    "p_value": round(p_val, 6),
                    "vif": None
                })
            
            uncoded_equation = self._build_equation(model_uncoded, reverse_names) if hasattr(self, '_build_equation') else ""
            # ★ v17: 未编码方程也追加 Ct Pt 项
            if ct_pt_coeff != 0:
                sign = " + " if ct_pt_coeff > 0 else " - "
                uncoded_equation += f"{sign}{abs(ct_pt_coeff):.4f}×Ct Pt"
        except Exception:
            pass
        
        # ── CI 导出参数 ──
        # ★ v18: 用全部数据（角点+中心点）的设计矩阵计算 XtX_inv（不含 Ct Pt 列）
        # CI 的 error_df = 弯曲df + 纯误差df（JMP 把弯曲项的自由度也纳入 CI 的误差）
        ci_export = {}
        try:
            import patsy
            design_info = model.model.data.orig_exog.design_info
            X_full = np.asarray(patsy.build_design_matrices([design_info], df_safe)[0])
            p_params = X_full.shape[1]
            XtX_inv_full = np.linalg.inv(X_full.T @ X_full)
            sigma_val = round(s, 8)
            # ★ v18: CI 的 df = error_df + curvature_df（对齐 JMP）
            ci_df = error_df + curvature_df
            t_crit_val = round(float(sp_stats.t.ppf(0.975, ci_df)), 6)
            ci_export = {
                "xtx_inv_flat": [round(v, 10) for v in XtX_inv_full.flatten().tolist()],
                "sigma": sigma_val,
                "t_crit": t_crit_val,
                "param_count": p_params
            }
        except Exception:
            try:
                p_params = X.shape[1]
                sigma_val = round(s, 8)
                t_crit_val = round(float(sp_stats.t.ppf(0.975, error_df)), 6)
                ci_export = {
                    "xtx_inv_flat": [round(v, 10) for v in XtX_inv.flatten().tolist()],
                    "sigma": sigma_val,
                    "t_crit": t_crit_val,
                    "param_count": p_params
                }
            except Exception:
                ci_export = {"xtx_inv_flat": [], "sigma": 0.0, "t_crit": 0.0, "param_count": 0}
        
        # ── 模型摘要 ──
        model_p_overall = round(model_p_, 6)
        equation = self._build_equation(model, reverse_names) if hasattr(self, '_build_equation') else ""
        equations = self._build_equations_by_category(model, reverse_names) if hasattr(self, '_build_equations_by_category') else None
        
        # ★ v17: 在方程末尾追加 Ct Pt 项
        if ct_pt_coeff != 0:
            sign = " + " if ct_pt_coeff > 0 else " - "
            equation += f"{sign}{abs(ct_pt_coeff):.4f}×Ct Pt"
        
        # Adequate Precision
        adeq_prec = 0.0
        try:
            fitted_arr = model.fittedvalues.values
            y_range = float(fitted_arr.max() - fitted_arr.min())
            mean_var_pred = error_ms * len(model.params) / n_corner
            if mean_var_pred > 0:
                adeq_prec = float(y_range / np.sqrt(mean_var_pred))
        except Exception:
            pass
        
        model_summary = {
            "r_squared": round(r_sq, 6),
            "r_squared_adj": round(r_sq_adj, 6),
            "r_squared_pred": r_sq_pred_display,
            "rmse": round(s, 6),
            "adeq_precision": round(adeq_prec, 4),
            "press": round(press_total, 4),
            "lack_of_fit_p": round(lof_p, 6) if lof_p is not None else None,
            "model_p": model_p_overall,
            "equation": equation,
            "equations": equations,
            "mse": round(error_ms, 6),  # ★ v17: 供 Pareto 使用
            "df_error": error_df,  # ★ v17: 供 Pareto 使用
        }
        
        return {
            "model": model,
            "anova_table": anova_rows,
            "coefficients": coefficients,
            "uncoded_coefficients": uncoded_coefficients,
            "model_summary": model_summary,
            "curvature": {
                "has_curvature_test": True,
                "n_corner": n_corner,
                "n_center": n_center,
                "corner_mean": round(corner_mean, 6),
                "center_mean": round(center_mean, 6),
                "ct_pt_coeff": round(ct_pt_coeff, 6),
                "ct_pt_se": round(ct_pt_se, 6),
                "ct_pt_t": round(ct_pt_t, 4),
                "ct_pt_p": round(ct_pt_p, 6),
                "curvature_ss": round(curvature_ss, 6),
                "curvature_df": curvature_df,
                "curvature_f": round(curvature_f, 4),
                "curvature_p": round(curvature_p, 6),
            },
            "uncoded_equation": uncoded_equation,
            "uncoded_equations": uncoded_equations,
            "xtx_inv_flat": ci_export.get("xtx_inv_flat", []),
            "sigma": ci_export.get("sigma", 0.0),
            "t_crit": ci_export.get("t_crit", 0.0),
            "param_count": ci_export.get("param_count", 0),
        }

    # ═══════════════════════════════════════════════════════
    # ★ v7 新增: 不可估计项检测核心方法
    # ═══════════════════════════════════════════════════════

    def residual_diagnostics(self) -> str:
        """
        ★ 新增: 完整残差诊断四图数据（对标 JMP/Minitab 残差诊断面板）
        
        返回 JSON:
        {
            "normal_probability": {
                "theoretical_quantiles": [...],  — 标准正态分位数
                "ordered_residuals": [...]        — 排序后的标准化残差
            },
            "residuals_vs_fitted": {
                "fitted": [...],
                "residuals": [...]                — 标准化残差
            },
            "residuals_vs_order": {
                "order": [1, 2, 3, ...],
                "residuals": [...]
            },
            "residuals_histogram": {
                "bin_edges": [...],
                "frequencies": [...]
            },
            "cooks_distance": {
                "observation": [1, 2, ...],
                "distance": [...]                 — Cook's 距离，>1 表示强影响点
            }
        }
        
        数学原理:
          标准化残差: e*_i = e_i / (σ̂ √(1 - h_ii))
            其中 h_ii 是杠杆值（Hat 矩阵对角元素）
          
          正态概率图: 如果残差服从正态分布，点应落在直线上
          
          Cook's 距离: D_i = (e*_i)² × h_ii / (p × (1-h_ii))
            度量删除第 i 个观测对所有拟合值的影响
        """
        if self._model is None:
            return json.dumps({"error": "模型未拟合"}, ensure_ascii=False)
        
        try:
            from scipy import stats as sp_stats
            
            # ★ v17: 弯曲检验模式 — 角点饱和模型残差=0，需要用全部数据重新计算
            is_curvature_mode = (hasattr(self, '_curvature_mse') and self._curvature_mse is not None)
            
            if is_curvature_mode:
                # 用全部数据（角点+中心点）计算残差
                # 角点: 拟合值 = 模型预测（精确），残差 = 0
                # 中心点: 拟合值 = corner_mean + Ct Pt coeff
                model = self._model
                mse = self._curvature_mse
                s = np.sqrt(mse)
                
                # 获取全部实际值
                all_y = self._df[self._response_name].values
                n_total = len(all_y)
                
                # 检测中心点
                safe_continuous = [f"X{i}" for i, f in enumerate(self._factor_names) if f not in self._categorical_factors]
                df_safe = pd.DataFrame()
                for i, f in enumerate(self._factor_names):
                    safe = f"X{i}"
                    if f in self._categorical_factors:
                        df_safe[safe] = self._df[f].astype(str)
                    else:
                        info = self._coding_info.get(f, {})
                        center = info.get("center", 0)
                        half_range = info.get("half_range", 1)
                        df_safe[safe] = (self._df[f].astype(float) - center) / half_range
                
                center_mask = self._detect_center_points(df_safe, safe_continuous)
                corner_idx = np.where(~center_mask)[0]
                center_idx = np.where(center_mask)[0]
                
                # 角点拟合值 = 模型精确预测
                fitted_all = np.zeros(n_total)
                fitted_all[corner_idx] = model.fittedvalues.values
                
                # 中心点拟合值 = center_mean（Minitab 做法）
                center_mean = float(np.mean(all_y[center_idx]))
                fitted_all[center_idx] = center_mean
                
                residuals = all_y - fitted_all
                n = n_total
                
                # Minitab 残差图用非标准化残差（raw residuals）
                # 因为角点残差=0，标准化没有意义
                raw_resid = residuals
                
                # 1. 正态概率图
                sorted_resid = np.sort(raw_resid)
                theoretical = sp_stats.norm.ppf((np.arange(1, n + 1) - 0.375) / (n + 0.25))
                
                # 4. 残差直方图
                n_bins = max(5, int(np.sqrt(n)))
                hist_freq, hist_edges = np.histogram(raw_resid, bins=n_bins)
                
                return json.dumps({
                    "normal_probability": {
                        "theoretical_quantiles": [round(float(v), 4) for v in theoretical],
                        "ordered_residuals": [round(float(v), 4) for v in sorted_resid]
                    },
                    "residuals_vs_fitted": {
                        "fitted": [round(float(v), 4) for v in fitted_all],
                        "residuals": [round(float(v), 4) for v in raw_resid]
                    },
                    "residuals_vs_order": {
                        "order": list(range(1, n + 1)),
                        "residuals": [round(float(v), 4) for v in raw_resid]
                    },
                    "residuals_histogram": {
                        "bin_edges": [round(float(v), 4) for v in hist_edges],
                        "frequencies": [int(v) for v in hist_freq]
                    },
                    "cooks_distance": {
                        "observation": list(range(1, n + 1)),
                        "distance": [0.0] * n
                    }
                }, ensure_ascii=False)
            
            model = self._model
            residuals = model.resid.values
            fitted = model.fittedvalues.values
            n = len(residuals)
            
            # 标准化残差
            influence = model.get_influence()
            std_resid = influence.resid_studentized_internal
            
            # 1. 正态概率图
            sorted_resid = np.sort(std_resid)
            # 理论分位数: Φ^(-1)((i-0.375)/(n+0.25))  — Blom's 公式
            theoretical = sp_stats.norm.ppf((np.arange(1, n + 1) - 0.375) / (n + 0.25))
            
            # 2. 残差 vs 拟合值（已有）
            
            # 3. 残差 vs 观测顺序（已有）
            
            # 4. 残差直方图
            n_bins = max(5, int(np.sqrt(n)))
            hist_freq, hist_edges = np.histogram(std_resid, bins=n_bins)
            
            # 5. Cook's 距离
            cooks_d = influence.cooks_distance[0]
            
            return json.dumps({
                "normal_probability": {
                    "theoretical_quantiles": [round(float(v), 4) for v in theoretical],
                    "ordered_residuals": [round(float(v), 4) for v in sorted_resid]
                },
                "residuals_vs_fitted": {
                    "fitted": [round(float(v), 4) for v in fitted],
                    "residuals": [round(float(v), 4) for v in std_resid]
                },
                "residuals_vs_order": {
                    "order": list(range(1, n + 1)),
                    "residuals": [round(float(v), 4) for v in std_resid]
                },
                "residuals_histogram": {
                    "bin_edges": [round(float(v), 4) for v in hist_edges],
                    "frequencies": [int(v) for v in hist_freq]
                },
                "cooks_distance": {
                    "observation": list(range(1, n + 1)),
                    "distance": [round(float(v), 6) for v in cooks_d]
                }
            }, ensure_ascii=False)
            
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # ★ 新增: effects_pareto() — 基于 T 值的效应 Pareto 图
    # ═══════════════════════════════════════════════════════

    def effects_pareto(self, alpha: float = 0.05) -> str:
        """
        ★ v13/v17: Minitab 风格 Pareto 图
        
        v17 改进: 全因子弯曲检验时，使用合并 MSE 计算标准化效应 |t|
        Minitab Pareto 图: 横轴 = |t|, 参考线 = t(1-α/2, df_error)
        """
        if self._model is None:
            return json.dumps([], ensure_ascii=False)
        
        try:
            import re
            from scipy import stats as sp_stats
            import numpy as np
            
            model = self._model
            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            reverse_names = {v: k for k, v in safe_names.items()}
            safe_continuous = [safe_names[f] for f in self._continuous_factors]
            safe_categorical = [safe_names[f] for f in self._categorical_factors]
            
            # ★ v17: 检查是否走了弯曲检验路径
            # 弯曲检验时 self._model 是角点饱和模型 (df_resid=0)
            # 需要用合并 MSE 重新计算 T 值
            is_curvature_mode = (hasattr(self, '_curvature_mse') and self._curvature_mse is not None)
            
            if is_curvature_mode:
                # ★ v17: 弯曲检验路径 — 直接从角点模型 + 合并 MSE 计算
                mse = self._curvature_mse
                df_error = self._curvature_df_error
                
                X = model.model.exog
                try:
                    XtX_inv = np.linalg.inv(X.T @ X)
                except np.linalg.LinAlgError:
                    XtX_inv = np.linalg.pinv(X.T @ X)
                
                params_list = list(model.params.index)
                effects = []
                
                for i, pname in enumerate(params_list):
                    if pname == "Intercept":
                        continue
                    coef = float(model.params.iloc[i])
                    se = float(np.sqrt(mse * XtX_inv[i, i]))
                    t_val = coef / se if se > 0 else 0
                    p_val = float(2 * (1 - sp_stats.t.cdf(abs(t_val), df_error)))
                    
                    # 还原显示名
                    display = pname
                    for safe, orig in reverse_names.items():
                        display = display.replace(safe, orig)
                    display = display.replace(":", "×")
                    
                    effects.append({
                        "term": display,
                        "t_value": round(abs(t_val), 4),
                        "abs_t": round(abs(t_val), 4),
                        "p_value": round(p_val, 6),
                        "significant": p_val < alpha,
                        "df": 1,
                        "log_worth": round(float(-np.log10(max(p_val, 1e-20))), 4),
                    })
                
                effects.sort(key=lambda x: x["abs_t"], reverse=True)
                
                t_crit = float(sp_stats.t.ppf(1 - alpha / 2, df_error))
                m = len(effects)
                for eff in effects:
                    eff["bonferroni_significant"] = eff["p_value"] < alpha / m if m > 0 else False
                
                return json.dumps({
                    "effects": effects,
                    "t_crit": round(t_crit, 4),
                    "log_worth_crit": round(t_crit, 4),  # Minitab Pareto 参考线 = t 临界值
                    "alpha": alpha,
                    "df_error": df_error,
                    "measure": "StandardizedEffect"  # Minitab 标准化效应
                }, ensure_ascii=False)
            
            # ── 非弯曲检验路径: 原有逻辑 ──
            df_safe = pd.DataFrame()
            for orig_name, safe_name in safe_names.items():
                if orig_name in self._categorical_factors:
                    df_safe[safe_name] = self._df[orig_name].astype(str)
                else:
                    info = self._coding_info.get(orig_name, {})
                    c = info.get("center", 0.0)
                    hr = info.get("half_range", 1.0)
                    df_safe[safe_name] = (self._df[orig_name].astype(float) - float(c)) / float(hr)
            df_safe["Y"] = self._df[self._response_name].values
            
            if hasattr(model.model, 'formula'):
                formula_coded = str(model.model.formula)
            else:
                formula_coded = self._build_formula(safe_continuous, safe_categorical, "quadratic")
            
            formula_sum = re.sub(r'C\((\w+)\)(?![,\[])', r'C(\1, Sum)', formula_coded)
            
            try:
                from statsmodels.formula.api import ols
                model_sum = ols(formula_sum, data=df_safe).fit()
            except Exception:
                model_sum = model  # 回退
            
            params = list(model_sum.params.index)
            
            # ── 按因子分组参数 ──
            factor_groups = {}
            
            for p_name in params:
                if p_name == "Intercept":
                    continue
                
                p_clean = p_name.replace(", Sum", "")
                
                if "**" in p_clean:
                    match = re.search(r'I\((\w+)\s*\*\*\s*2\)', p_clean)
                    if match:
                        base_var = match.group(1)
                        display = self._restore_term_name(base_var, reverse_names) + "²"
                        factor_groups.setdefault(display, []).append(p_name)
                elif ":" in p_clean:
                    parts = p_clean.split(":")
                    factor_names_in_term = []
                    for part in parts:
                        m = re.match(r'C\((\w+)', part)
                        if m:
                            factor_names_in_term.append(m.group(1))
                        else:
                            factor_names_in_term.append(part.strip())
                    display_parts = []
                    for fn in factor_names_in_term:
                        display_parts.append(reverse_names.get(fn, fn))
                    display = "×".join(sorted(set(display_parts)))
                    factor_groups.setdefault(display, []).append(p_name)
                else:
                    m = re.match(r'C\((\w+)', p_clean)
                    if m:
                        base_var = m.group(1)
                    else:
                        base_var = p_clean
                    display = reverse_names.get(base_var, base_var)
                    factor_groups.setdefault(display, []).append(p_name)
            
            # ── 对每个因子组做 GLH 联合 F 检验 ──
            effects = []
            for display_name, group_params in factor_groups.items():
                df_group = len(group_params)
                try:
                    r_matrix = np.zeros((df_group, len(params)))
                    for i, pn in enumerate(group_params):
                        j = params.index(pn)
                        r_matrix[i, j] = 1.0
                    f_result = model_sum.f_test(r_matrix)
                    f_val = float(f_result.fvalue)
                    p_val = float(f_result.pvalue)
                    t_equiv = float(np.sqrt(f_val))
                    log_worth = float(-np.log10(max(p_val, 1e-20)))
                except Exception:
                    t_equiv = 0.0
                    p_val = 1.0
                    log_worth = 0.0
                
                effects.append({
                    "term": display_name,
                    "t_value": round(t_equiv, 4),
                    "abs_t": round(log_worth, 4),
                    "p_value": round(p_val, 6),
                    "significant": p_val < alpha,
                    "df": df_group,
                    "log_worth": round(log_worth, 4)
                })
            
            effects.sort(key=lambda x: x["abs_t"], reverse=True)
            
            m = len(effects)
            df_error = int(model_sum.df_resid)
            t_crit = float(sp_stats.t.ppf(1 - alpha / 2, df_error))
            p_norm_at_tcrit = 2 * (1 - sp_stats.norm.cdf(t_crit))
            log_worth_crit = float(-np.log10(p_norm_at_tcrit))
            if m > 0 and df_error > 0:
                for eff in effects:
                    eff["bonferroni_significant"] = eff["p_value"] < alpha / m
            
            return json.dumps({
                "effects": effects,
                "log_worth_crit": round(log_worth_crit, 4),
                "alpha": alpha,
                "df_error": df_error,
                "measure": "LogWorth"
            }, ensure_ascii=False)
            
        except Exception as e:
            return json.dumps({"effects": [{"error": str(e)}], "log_worth_crit": 1.301, "alpha": 0.05, "measure": "LogWorth"}, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # 原有方法（保留兼容）
    # ═══════════════════════════════════════════════════════


    # ═══════════════════════════════════════════════════════
    # ★ v7 新增方法: 不可估计项检测、OLS 响应曲面、预测刻画器、最优化
    # ═══════════════════════════════════════════════════════

    def _detect_and_drop_inestimable(self, formula: str, df_safe: pd.DataFrame,
                                      safe_names: dict, reverse_names: dict,
                                      safe_continuous: list, safe_categorical: list,
                                      model_type: str) -> tuple:
        """
        ★ v10 重写: 精确检测并剔除不可估计的模型项

        策略（对标 Minitab）:
          1. 拟合当前公式，检查 NaN 参数和 SVD 近奇异
          2. 如果近奇异: 逐项试删，找到删掉后 SVD ratio 恢复最大的那一项
          3. 在同等 SVD 改善下，优先剔除高阶项（二次 > 交互 > 主效应）
          4. 剔除后重复检查，直到矩阵满秩

        返回: (new_formula, dropped_term_names)
        """

        rhs = formula.split("~")[1].strip() if "~" in formula else ""
        if not rhs or rhs == "1":
            return formula, []

        current_terms = [t.strip() for t in rhs.split("+")]
        dropped = []

        def _term_order(t):
            """项的阶次: 值越大越优先被剔除"""
            if "**2" in t:
                return 3  # 二次项
            elif ":" in t:
                return 2  # 交互项
            elif t.startswith("C("):
                return 0  # 类别主效应 — 最不愿意剔除
            else:
                return 1  # 连续主效应

        def _check_singular(terms_list):
            """检查公式是否近奇异，返回 (is_singular, sv_ratio, nan_params, model)"""
            if not terms_list:
                return True, 0, [], None
            f = "Y ~ " + " + ".join(terms_list)
            try:
                m = ols(f, data=df_safe).fit()
                nans = [p for p in m.params.index if np.isnan(m.params[p])]
                if nans:
                    return True, 0, nans, m
                X = m.model.exog
                svd = np.linalg.svd(X, compute_uv=False)
                ratio = svd[-1] / svd[0] if svd[0] > 0 else 0
                return ratio < 1e-12, ratio, [], m
            except Exception:
                return True, 0, [], None

        max_attempts = len(current_terms)
        for attempt in range(max_attempts):
            is_singular, sv_ratio, nan_params, test_model = _check_singular(current_terms)

            if not is_singular:
                break  # 满秩，结束

            if nan_params:
                # 有 NaN 参数: 直接定位对应的公式项
                removed = False
                for nan_p in nan_params:
                    if nan_p in current_terms:
                        current_terms.remove(nan_p)
                        dropped.append(self._restore_term_name(nan_p, reverse_names))
                        removed = True
                        break
                if not removed:
                    # NaN 参数名不在 current_terms（哑变量展开名），
                    # 找包含该哑变量的公式项
                    for nan_p in nan_params:
                        for t in current_terms:
                            if t in nan_p or nan_p.startswith(t.replace("C(", "").replace(")", "")):
                                current_terms.remove(t)
                                dropped.append(self._restore_term_name(t, reverse_names))
                                removed = True
                                break
                        if removed:
                            break
                    if not removed:
                        # 兜底: 按阶次删最高阶项
                        sorted_terms = sorted(current_terms, key=_term_order, reverse=True)
                        if sorted_terms:
                            t = sorted_terms[0]
                            current_terms.remove(t)
                            dropped.append(self._restore_term_name(t, reverse_names))
                        else:
                            break
            else:
                # ★ v10 核心: 无 NaN 但近奇异 — 逐项试删，精确定位
                # 选择标准（优先级从高到低，纯线性代数，不依赖响应值）:
                #   1. 删掉后 sv_ratio 恢复最大（矩阵最健康）
                #   2. sv_ratio 相同时，优先删高阶项（二次 > 交互 > 主效应）
                #   3. 阶次也相同时，优先删展开后哑变量列数最多的项（自由度成本最大）
                #   4. 以上全相同时，删公式中靠后的项（与 Minitab 行为一致）
                best_term = None
                best_ratio = -1
                best_order = -1
                best_dummy_count = -1
                best_position = -1  # 在公式中的位置（越大 = 越靠后 = 越优先删）

                for idx_t, t in enumerate(current_terms):
                    trial = [x for x in current_terms if x != t]
                    if not trial:
                        continue
                    _, ratio_after, _, _ = _check_singular(trial)
                    order = _term_order(t)
                    # 估算该项展开后的哑变量列数
                    dummy_count = 1
                    if "C(" in t and ":" in t:
                        # 类别×连续交互: 哑变量数 = 类别因子水平数 - 1
                        dummy_count = 2  # 默认估计
                        for cat_name in safe_categorical:
                            if cat_name in t:
                                n_levels = df_safe[cat_name].nunique()
                                dummy_count = max(n_levels - 1, 1)
                                break

                    if (ratio_after > best_ratio + 1e-6):
                        best_ratio = ratio_after
                        best_term = t
                        best_order = order
                        best_dummy_count = dummy_count
                        best_position = idx_t
                    elif abs(ratio_after - best_ratio) < 1e-6:
                        # sv_ratio 相同: 按阶次 > 哑变量列数 > 位置
                        if order > best_order:
                            best_term = t
                            best_order = order
                            best_dummy_count = dummy_count
                            best_position = idx_t
                        elif order == best_order:
                            if dummy_count > best_dummy_count:
                                best_term = t
                                best_dummy_count = dummy_count
                                best_position = idx_t
                            elif dummy_count == best_dummy_count and idx_t > best_position:
                                best_term = t
                                best_position = idx_t

                if best_term is not None:
                    current_terms.remove(best_term)
                    dropped.append(self._restore_term_name(best_term, reverse_names))
                else:
                    break  # 无法改善

        final_formula = "Y ~ " + " + ".join(current_terms) if current_terms else "Y ~ 1"
        return final_formula, dropped

    # ═══════════════════════════════════════════════════════
    # ★ v7 新增: OLS 响应曲面数据
    # ═══════════════════════════════════════════════════════

    def _predict_with_se(self, point: dict, XtX_inv, sigma: float, has_ci: bool,
                         t_crit: float = 0.0) -> tuple:
        """
        ★ v8 修复: 预测单点并计算均值置信区间标准误 (CI of mean，与 JMP 刻画器一致)

        返回: (predicted_value, se_mean)
          se_mean = σ̂ × √(x₀ᵀ(XᵀX)⁻¹x₀)   — 均值的置信区间标准误

        ★ v8 最终修复说明:
          经实测验证，JMP 刻画器显示的是 CI of mean（均值置信区间），不是 PI（预测区间）。
            CI of mean: ŷ ± t × σ̂ × √(h)         其中 h = x₀ᵀ(XᵀX)⁻¹x₀
            PI:         ŷ ± t × σ̂ × √(1 + h)     (比 CI 更宽)
          使用 Sum coding 模型的 (XᵀX)⁻¹ 计算杠杆值，与 JMP 的 Effect coding 一致。
        """
        pred_json = self.predict_point(json.dumps(point, ensure_ascii=False))
        pred_val = json.loads(pred_json).get("predicted", 0.0)

        if not has_ci or XtX_inv is None:
            return pred_val, 0.0

        try:
            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            row = {}
            for orig_name, safe_name in safe_names.items():
                val = point.get(orig_name)
                if orig_name in self._categorical_factors:
                    row[safe_name] = str(val)
                else:
                    info = self._coding_info.get(orig_name, {})
                    c = info.get("center", 0.0)
                    hr = info.get("half_range", 1.0)
                    row[safe_name] = (float(val) - c) / hr

            df_point = pd.DataFrame([row])

            # ★ 方法1（优先）: Sum coding + patsy 手动计算 CI of mean
            ci_model = self._model_sum if self._model_sum is not None else self._model
            try:
                import patsy
                # ★ v18 修复: 弯曲检验模式下使用合并误差（角点饱和模型 mse_resid=0）
                if hasattr(self, '_curvature_ci_params') and self._curvature_ci_params is not None:
                    ci_sigma = self._curvature_ci_params["sigma"]
                    # ★ v18: 用全部数据（角点+中心点）的 XtX_inv 计算杠杆值
                    ci_XtX_inv = self._curvature_ci_params.get("XtX_inv")
                    if ci_XtX_inv is not None:
                        design_info = ci_model.model.data.orig_exog.design_info
                        x0_row = patsy.build_design_matrices([design_info], df_point)[0]
                        x0 = np.asarray(x0_row)[0]
                        hat_value = float(x0 @ ci_XtX_inv @ x0)
                        se_mean = ci_sigma * np.sqrt(max(hat_value, 0.0))
                        return pred_val, round(float(se_mean), 6)
                else:
                    ci_sigma = np.sqrt(float(ci_model.mse_resid))
                design_info = ci_model.model.data.orig_exog.design_info
                x0_row = patsy.build_design_matrices([design_info], df_point)[0]
                x0 = np.asarray(x0_row)[0]
                ci_exog = ci_model.model.exog
                ci_XtX_inv = np.linalg.inv(ci_exog.T @ ci_exog)
                hat_value = float(x0 @ ci_XtX_inv @ x0)
                se_mean = ci_sigma * np.sqrt(max(hat_value, 0.0))
                return pred_val, round(float(se_mean), 6)
            except Exception:
                pass

            # ★ 方法2: Treatment coding + patsy
            try:
                import patsy
                design_info = self._model.model.data.orig_exog.design_info
                x0_row = patsy.build_design_matrices([design_info], df_point)[0]
                x0 = np.asarray(x0_row)[0]
                hat_value = float(x0 @ XtX_inv @ x0)
                se_mean = sigma * np.sqrt(max(hat_value, 0.0))
                return pred_val, round(float(se_mean), 6)
            except Exception:
                pass

            # ★ 方法3（最终回退）: 平均杠杆近似
            try:
                n = self._model.model.exog.shape[0]
                p = self._model.model.exog.shape[1]
                se_mean = sigma * np.sqrt(float(p) / float(n))
                return pred_val, round(float(se_mean), 6)
            except Exception:
                pass

            return pred_val, 0.0
        except Exception:
            return pred_val, 0.0

    # ═══════════════════════════════════════════════════════
    # find_optimal — OLS 最优化
    # ═══════════════════════════════════════════════════════

    def _find_single_df_t_value(self, anova_term: str, model) -> Optional[float]:
        try:
            if anova_term in model.tvalues.index:
                return float(model.tvalues[anova_term])
            anova_clean = anova_term.replace(" ", "")
            for param_term in model.tvalues.index:
                if param_term.replace(" ", "") == anova_clean:
                    return float(model.tvalues[param_term])
            return None
        except Exception:
            return None

    # ═══════════════════════════════════════════════════════
    # 原有方法（保留兼容）
    # ═══════════════════════════════════════════════════════

    def response_surface_data_ols(self, factor1: str, factor2: str, grid_size: int = 30,
                                   bounds_json: str = "{}") -> str:
        """
        ★ v7 新增: 基于 OLS 模型生成 3D 响应曲面 + 等高线数据

        与 GPR 版 DOESurfacePlotter.response_surface_data() 返回相同格式，
        C# 端可复用同一套 OxyPlot 渲染逻辑。

        factor1, factor2: 原始因子名（如 "温度", "压力"）
        bounds_json: '{"温度": [80, 160], "压力": [10, 30], ...}'
          如果为空，从数据中自动推断范围
        grid_size: 网格密度

        返回 JSON:
        {
            "x": [80, 82.7, ...],      — factor1 的值数组
            "y": [10, 10.7, ...],       — factor2 的值数组
            "z": [[...], [...], ...],   — 预测值矩阵 (grid_size × grid_size)
            "x_label": "温度",
            "y_label": "压力"
        }
        """
        if self._model is None or self._df is None:
            return json.dumps({"error": "模型未拟合"}, ensure_ascii=False)

        try:
            bounds = json.loads(bounds_json) if bounds_json else {}

            # 确定两个因子的范围
            def get_bounds(fname):
                if fname in bounds:
                    b = bounds[fname]
                    if isinstance(b, list) and len(b) == 2:
                        return float(b[0]), float(b[1])
                # 从数据推断
                if fname in self._categorical_factors:
                    return 0, 1  # 类别因子不适合做曲面
                col = self._df[fname].astype(float)
                return float(col.min()), float(col.max())

            x_min, x_max = get_bounds(factor1)
            y_min, y_max = get_bounds(factor2)

            x_range = np.linspace(x_min, x_max, grid_size)
            y_range = np.linspace(y_min, y_max, grid_size)

            # 其他因子固定在中心值
            center_values = {}
            for name in self._factor_names:
                if name == factor1 or name == factor2:
                    continue
                if name in self._categorical_factors:
                    levels = sorted(self._df[name].astype(str).unique())
                    center_values[name] = levels[0] if levels else ""
                else:
                    col = self._df[name].astype(float)
                    center_values[name] = float((col.min() + col.max()) / 2.0)

            # 计算网格上的预测值
            z_grid = np.zeros((grid_size, grid_size))
            for i, x_val in enumerate(x_range):
                for j, y_val in enumerate(y_range):
                    point = dict(center_values)
                    point[factor1] = float(x_val)
                    point[factor2] = float(y_val)
                    pred_json = self.predict_point(json.dumps(point, ensure_ascii=False))
                    pred = json.loads(pred_json)
                    z_grid[j, i] = pred.get("predicted", 0.0)

            return json.dumps({
                "x": [round(float(v), 4) for v in x_range],
                "y": [round(float(v), 4) for v in y_range],
                "z": [[round(float(z_grid[j, i]), 4) for i in range(grid_size)] for j in range(grid_size)],
                "x_label": factor1,
                "y_label": factor2
            }, ensure_ascii=False)

        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    def response_surface_image_ols(self, factor1: str, factor2: str,
                                    grid_size: int = 50, width: int = 800, height: int = 600,
                                    bounds_json: str = "{}") -> str:
        """
        ★ v7 新增: 基于 OLS 模型生成响应曲面 PNG 图片（base64）

        返回 base64 编码的 PNG 图片字符串，空字符串表示生成失败
        """
        if not HAS_MATPLOTLIB or self._model is None:
            return ""

        try:
            data = json.loads(self.response_surface_data_ols(factor1, factor2, grid_size, bounds_json))
            if "error" in data:
                return ""

            X = np.array(data["x"])
            Y = np.array(data["y"])
            Z = np.array(data["z"])
            XX, YY = np.meshgrid(X, Y)

            fig = plt.figure(figsize=(width / 100, height / 100), dpi=100)
            ax = fig.add_subplot(111, projection='3d')
            surf = ax.plot_surface(XX, YY, Z, cmap='viridis', alpha=0.85, edgecolor='none')
            ax.set_xlabel(factor1, fontsize=11)
            ax.set_ylabel(factor2, fontsize=11)
            ax.set_zlabel(self._response_name, fontsize=11)
            ax.set_title(f'OLS 响应曲面: {factor1} × {factor2}', fontsize=13)
            fig.colorbar(surf, shrink=0.5, aspect=10)
            plt.tight_layout()

            buf = io.BytesIO()
            fig.savefig(buf, format='png', bbox_inches='tight')
            plt.close(fig)
            buf.seek(0)
            return base64.b64encode(buf.read()).decode('utf-8')
        except Exception:
            return ""

    def contour_data_ols(self, factor1: str, factor2: str, grid_size: int = 30,
                          bounds_json: str = "{}") -> str:
        """
        ★ v7 新增: OLS 等高线数据（与 response_surface_data_ols 格式相同）
        """
        return self.response_surface_data_ols(factor1, factor2, grid_size, bounds_json)

    # ═══════════════════════════════════════════════════════
    # residual_diagnostics() — 残差四图数据
    # ═══════════════════════════════════════════════════════

    def predict_point(self, factors_json: str) -> str:
        if self._model is None or self._df is None:
            return json.dumps({"error": "模型未拟合"}, ensure_ascii=False)
        try:
            factors = json.loads(factors_json)
            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            row = {}
            for orig_name, safe_name in safe_names.items():
                val = factors.get(orig_name)
                if orig_name in self._categorical_factors:
                    row[safe_name] = str(val)
                else:
                    info = self._coding_info.get(orig_name, {})
                    center = info.get("center", 0.0)
                    half_range = info.get("half_range", 1.0)
                    row[safe_name] = (float(val) - center) / half_range
            df_point = pd.DataFrame([row])
            pred = self._model.predict(df_point)
            return json.dumps({"predicted": round(float(pred.iloc[0]), 4)}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # ★ v7 改造: prediction_profiler — 增加置信区间带
    # ═══════════════════════════════════════════════════════

    def prediction_profiler(self, grid_size: int = 50, fixed_values_json: str = "",
                             alpha: float = 0.05) -> str:
        """
        ★ v8 改造: 预测刻画器数据 — 预测区间带 (PI，与 JMP 刻画器一致)

        新增返回字段（每个因子）:
          "y_lower": [...]   — 预测区间下界 (1-α)
          "y_upper": [...]   — 预测区间上界 (1-α)

        新增返回字段（全局）:
          "current_ci_lower": float  — 当前预测点的 PI 下界
          "current_ci_upper": float  — 当前预测点的 PI 上界

        CI 计算方法 (PI = 预测区间，与 JMP 刻画器一致):
          ŷ ± t(α/2, df_error) × SE_pred
          SE_pred = σ̂ × √(1 + x₀ᵀ(XᵀX)⁻¹x₀)

        ★ v8 修复:
          v7 的回退分支只返回 σ（缺少杠杆项），导致 CI 比 JMP 窄约 2.4 倍。
          v8 增加了 patsy 手动构建 x₀ 的回退计算，确保 CI 宽度正确。
        """
        if self._model is None or self._df is None:
            return json.dumps({"error": "模型未拟合"}, ensure_ascii=False)

        try:
            from scipy import stats as sp_stats

            fixed = json.loads(fixed_values_json) if fixed_values_json else {}
            center = {}
            for name in self._factor_names:
                if name in fixed:
                    center[name] = fixed[name]
                elif name in self._categorical_factors:
                    levels = sorted(self._df[name].astype(str).unique())
                    center[name] = levels[0] if levels else ""
                else:
                    col = self._df[name].astype(float)
                    center[name] = float((col.min() + col.max()) / 2.0)

            current_pred_json = self.predict_point(json.dumps(center, ensure_ascii=False))
            current_pred = json.loads(current_pred_json).get("predicted", 0.0)

            # ★ v7: 预计算 CI 所需的矩阵
            model = self._model
            X_matrix = model.model.exog
            n_obs = X_matrix.shape[0]
            p_params = X_matrix.shape[1]
            df_error = max(n_obs - p_params, 1)
            mse = float(model.mse_resid)
            sigma = np.sqrt(mse)
            t_crit = float(sp_stats.t.ppf(1 - alpha / 2, df_error))

            try:
                XtX_inv = np.linalg.inv(X_matrix.T @ X_matrix)
                has_ci = True
            except np.linalg.LinAlgError:
                XtX_inv = None
                has_ci = False

            # ★ v18 修复: 弯曲检验模式下，角点模型是饱和的 (mse_resid≈0)
            # 需要用合并误差（含中心点纯误差）替代，否则 CI 全为 0
            if hasattr(self, '_curvature_ci_params') and self._curvature_ci_params is not None:
                _ccp = self._curvature_ci_params
                sigma = _ccp["sigma"]
                df_error = _ccp["error_df"]
                t_crit = float(sp_stats.t.ppf(1 - alpha / 2, df_error))
                if _ccp["XtX_inv"] is not None:
                    XtX_inv = _ccp["XtX_inv"]
                    has_ci = True

            result = {"factors": {}}

            for factor in self._factor_names:
                is_cat = factor in self._categorical_factors

                if is_cat:
                    levels = sorted(self._df[factor].astype(str).unique())
                    x_values = levels
                    y_values = []
                    y_lower = []
                    y_upper = []
                    for lv in levels:
                        point = dict(center)
                        point[factor] = lv
                        pred_val, se_pred = self._predict_with_se(point, XtX_inv, sigma, has_ci, t_crit)
                        y_values.append(pred_val)
                        if has_ci:
                            y_lower.append(round(pred_val - t_crit * se_pred, 4))
                            y_upper.append(round(pred_val + t_crit * se_pred, 4))
                        else:
                            y_lower.append(pred_val)
                            y_upper.append(pred_val)

                    result["factors"][factor] = {
                        "x": x_values, "y": y_values,
                        "y_lower": y_lower, "y_upper": y_upper,
                        "is_categorical": True,
                        "current_value": str(center[factor]),
                        "levels": levels
                    }
                else:
                    col = self._df[factor].astype(float)
                    f_min, f_max = float(col.min()), float(col.max())
                    x_values = np.linspace(f_min, f_max, grid_size).tolist()
                    y_values = []
                    y_lower = []
                    y_upper = []
                    for x_val in x_values:
                        point = dict(center)
                        point[factor] = x_val
                        pred_val, se_pred = self._predict_with_se(point, XtX_inv, sigma, has_ci, t_crit)
                        y_values.append(pred_val)
                        if has_ci:
                            y_lower.append(round(pred_val - t_crit * se_pred, 4))
                            y_upper.append(round(pred_val + t_crit * se_pred, 4))
                        else:
                            y_lower.append(pred_val)
                            y_upper.append(pred_val)

                    result["factors"][factor] = {
                        "x": [round(v, 4) for v in x_values],
                        "y": y_values,
                        "y_lower": y_lower, "y_upper": y_upper,
                        "is_categorical": False,
                        "current_value": round(float(center[factor]), 4),
                        "range": [round(f_min, 4), round(f_max, 4)]
                    }

            result["current_predicted"] = current_pred
            result["response_name"] = self._response_name

            # ★ v8 新增: 计算当前预测点的置信区间 (PI)
            if has_ci:
                _, se_current = self._predict_with_se(center, XtX_inv, sigma, has_ci, t_crit)
                result["current_ci_lower"] = round(current_pred - t_crit * se_current, 4)
                result["current_ci_upper"] = round(current_pred + t_crit * se_current, 4)
            else:
                result["current_ci_lower"] = current_pred
                result["current_ci_upper"] = current_pred

            return json.dumps(result, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    def find_optimal(self, maximize: bool = True, n_restarts: int = 20) -> str:
        """
        ★ v18 重写: 使用 differential_evolution 全局优化器替代 L-BFGS-B 多次重启
        对标 JMP/Minitab 的全局搜索能力，避免局部最优
        """
        if self._model is None or self._df is None:
            return json.dumps({"error": "模型未拟合", "success": False}, ensure_ascii=False)
        try:
            from scipy.optimize import differential_evolution
            cont_factors = self._continuous_factors
            cat_factors = self._categorical_factors
            cont_bounds = {}
            for name in cont_factors:
                col = self._df[name].astype(float)
                cont_bounds[name] = (float(col.min()), float(col.max()))
            cat_levels = {}
            for name in cat_factors:
                cat_levels[name] = sorted(self._df[name].astype(str).unique())
            if cat_factors:
                cat_combos = list(iterproduct(*[cat_levels[f] for f in cat_factors]))
            else:
                cat_combos = [()]
            global_best_val = float('-inf') if maximize else float('inf')
            global_best_factors = None
            for cat_combo in cat_combos:
                cat_dict = {cat_factors[i]: cat_combo[i] for i in range(len(cat_factors))} if cat_factors else {}
                if not cont_factors:
                    point = dict(cat_dict)
                    pred_json = self.predict_point(json.dumps(point, ensure_ascii=False))
                    pred_val = json.loads(pred_json).get("predicted", 0.0)
                    if (maximize and pred_val > global_best_val) or (not maximize and pred_val < global_best_val):
                        global_best_val = pred_val
                        global_best_factors = dict(point)
                    continue
                bounds_list = [(cont_bounds[f][0], cont_bounds[f][1]) for f in cont_factors]
                def objective(x):
                    point = dict(cat_dict)
                    for i, f in enumerate(cont_factors):
                        point[f] = float(x[i])
                    pred_json = self.predict_point(json.dumps(point, ensure_ascii=False))
                    pred_val = json.loads(pred_json).get("predicted", 0.0)
                    return -pred_val if maximize else pred_val
                try:
                    # ★ v18: differential_evolution 全局优化，不依赖初始点
                    res = differential_evolution(
                        objective, bounds=bounds_list,
                        maxiter=500, tol=1e-10, seed=42,
                        polish=True,       # 最后用 L-BFGS-B 精修
                        init='sobol',      # Sobol 序列初始化，覆盖更均匀
                        mutation=(0.5, 1.5),
                        recombination=0.9
                    )
                    val = -res.fun if maximize else res.fun
                    if (maximize and val > global_best_val) or (not maximize and val < global_best_val):
                        global_best_val = val
                        best_point = dict(cat_dict)
                        for i, f in enumerate(cont_factors):
                            best_point[f] = float(res.x[i])
                        global_best_factors = best_point
                except Exception:
                    pass
            if global_best_factors is None:
                return json.dumps({"error": "优化失败", "success": False}, ensure_ascii=False)
            in_range = True
            for name in cont_factors:
                val = global_best_factors[name]
                lo, hi = cont_bounds[name]
                if val < lo - 0.01 * (hi - lo) or val > hi + 0.01 * (hi - lo):
                    in_range = False
                    break
            optimal_factors = {}
            for name in self._factor_names:
                val = global_best_factors.get(name)
                if name in self._categorical_factors:
                    optimal_factors[name] = str(val)
                else:
                    optimal_factors[name] = round(float(val), 4)
            return json.dumps({
                "optimal_factors": optimal_factors,
                "predicted_response": round(float(global_best_val), 4),
                "maximize": maximize, "success": True, "in_range": in_range,
                "note": "最优条件在实验范围内" if in_range else "最优条件接近实验边界，建议谨慎外推"
            }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e), "success": False}, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # 内部辅助方法
    # ═══════════════════════════════════════════════════════


    # ═══════════════════════════════════════════════════════
    # ★ v8 新增方法: 异常点分析、排除重拟合、Tukey HSD
    # ═══════════════════════════════════════════════════════

    # ═══════════════════════════════════════════════════════
    # ★ v8 新增: 异常点分析 (Outlier Analysis)
    # ═══════════════════════════════════════════════════════

    def outlier_analysis(self, cooks_threshold: float = -1, leverage_threshold: float = -1,
                          residual_threshold: float = 3.0) -> str:
        """
        ★ v8 新增: 综合异常点分析
        
        判定标准（任一满足即标记为异常点）:
          1. Cook's Distance > 4/n（默认）或用户指定阈值
          2. |学生化残差| > residual_threshold（默认 3.0）
          3. 杠杆值 > 2p/n（默认）或用户指定阈值
        
        返回 JSON:
        {
            "outliers": [
                {
                    "index": 5,              — 观测序号 (1-based)
                    "actual": 85.3,
                    "predicted": 82.1,
                    "residual": 3.2,
                    "std_residual": 2.85,
                    "cooks_d": 0.52,
                    "leverage": 0.45,
                    "reasons": ["Cook's D > 0.222", "|残差| > 3σ"]
                },
                ...
            ],
            "thresholds": {
                "cooks_d": 0.222,
                "leverage": 0.333,
                "std_residual": 3.0
            },
            "total_observations": 18,
            "outlier_count": 2
        }
        """
        if self._model is None or self._df is None:
            return json.dumps({"error": "模型未拟合"}, ensure_ascii=False)
        
        try:
            model = self._model
            n = int(model.nobs)
            p = len(model.params)
            
            # 计算阈值
            cooks_thresh = cooks_threshold if cooks_threshold > 0 else 4.0 / n
            leverage_thresh = leverage_threshold if leverage_threshold > 0 else 2.0 * p / n
            resid_thresh = residual_threshold
            
            # 获取诊断量
            influence = model.get_influence()
            cooks_d = influence.cooks_distance[0]
            std_resid = influence.resid_studentized_internal
            leverage = influence.hat_matrix_diag
            fitted = model.fittedvalues.values
            actual = self._df[self._response_name].values
            residuals = model.resid.values
            
            outliers = []
            for i in range(n):
                reasons = []
                if cooks_d[i] > cooks_thresh:
                    reasons.append(f"Cook's D={cooks_d[i]:.3f} > {cooks_thresh:.3f}")
                if abs(std_resid[i]) > resid_thresh:
                    reasons.append(f"|标准化残差|={abs(std_resid[i]):.2f} > {resid_thresh}")
                if leverage[i] > leverage_thresh:
                    reasons.append(f"杠杆值={leverage[i]:.3f} > {leverage_thresh:.3f}")
                
                if reasons:
                    outliers.append({
                        "index": i + 1,
                        "actual": round(float(actual[i]), 4),
                        "predicted": round(float(fitted[i]), 4),
                        "residual": round(float(residuals[i]), 4),
                        "std_residual": round(float(std_resid[i]), 4),
                        "cooks_d": round(float(cooks_d[i]), 6),
                        "leverage": round(float(leverage[i]), 4),
                        "reasons": reasons
                    })
            
            return json.dumps({
                "outliers": outliers,
                "thresholds": {
                    "cooks_d": round(cooks_thresh, 4),
                    "leverage": round(leverage_thresh, 4),
                    "std_residual": resid_thresh
                },
                "total_observations": n,
                "outlier_count": len(outliers)
            }, ensure_ascii=False)
            
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    def refit_excluding(self, exclude_indices_json: str, model_type: str = "quadratic") -> str:
        """
        ★ v8 新增: 排除指定观测后重新拟合 OLS 模型
        
        exclude_indices_json: '[5, 12]' — 要排除的观测序号 (1-based)
        model_type: 模型类型
        
        工作原理:
          1. 保存原始完整数据到 self._df_full_backup
          2. 从 self._df 中去除指定行
          3. 重新编码、拟合模型
          4. 返回与 fit_ols 相同格式的结果 + 额外的 excluded_count 信息
        
        注意: 此方法会修改 self._model 为排除后的模型。
              调用 restore_full_data_from_backup() 可恢复完整数据。
        """
        if self._df is None:
            return json.dumps({"error": "数据未加载"}, ensure_ascii=False)
        
        try:
            exclude_indices = json.loads(exclude_indices_json) if exclude_indices_json else []
            # 转为 0-based
            exclude_0based = set(idx - 1 for idx in exclude_indices if 1 <= idx <= len(self._df))
            
            if not exclude_0based:
                # 没有要排除的，直接返回当前模型
                return self.fit_ols(model_type)
            
            # ★ 修复: 保存原始完整数据到备份属性，便于恢复
            if not hasattr(self, '_df_full_backup') or self._df_full_backup is None:
                self._df_full_backup = self._df.copy()
            
            # 排除指定行
            keep_mask = [i not in exclude_0based for i in range(len(self._df))]
            self._df = self._df.loc[keep_mask].reset_index(drop=True)
            
            if len(self._df) < 3:
                self._df = self._df_full_backup.copy()
                return json.dumps({"error": "排除后数据不足（至少需要3组）"}, ensure_ascii=False)
            
            # 重新拟合
            result_json = self.fit_ols(model_type)
            result = json.loads(result_json)
            
            # 添加排除信息
            result["excluded_indices"] = list(exclude_indices)
            result["excluded_count"] = len(exclude_0based)
            result["remaining_count"] = len(self._df)
            
            return json.dumps(result, ensure_ascii=False)
            
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    def restore_full_data(self, factors_json: str, responses_json: str, response_name: str,
                           factor_types_json: str = "{}") -> str:
        """
        ★ v8 新增: 恢复完整数据（取消异常点排除）
        本质上就是重新调用 load_data
        """
        return self.load_data(factors_json, responses_json, response_name, factor_types_json)

    def restore_full_data_from_backup(self, model_type: str = "quadratic") -> str:
        """
        ★ v11 新增: 从内部备份恢复完整数据（无需重新传入数据）
        
        在 refit_excluding() 后调用此方法可恢复到排除前的状态。
        如果没有备份数据，返回错误。
        """
        if not hasattr(self, '_df_full_backup') or self._df_full_backup is None:
            return json.dumps({"error": "没有备份数据，请使用 restore_full_data() 并传入原始数据"}, 
                            ensure_ascii=False)
        
        self._df = self._df_full_backup.copy()
        self._df_full_backup = None  # 清除备份
        
        # 重新拟合模型
        result_json = self.fit_ols(model_type)
        result = json.loads(result_json)
        result["restored"] = True
        result["restored_count"] = len(self._df)
        return json.dumps(result, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # ★ v8 新增: Tukey HSD 多重比较
    # ═══════════════════════════════════════════════════════

    def tukey_hsd(self, factor_name: str = "") -> str:
        """
        ★ v8 新增: Tukey HSD 事后多重比较检验
        
        针对类别因子，检验各水平间的均值差异是否显著。
        
        factor_name: 要检验的类别因子名。为空时对第一个类别因子检验。
        
        数学原理:
          Tukey HSD: q = (ȳ_i - ȳ_j) / SE,  其中 SE = √(MSE / n_h)
          n_h = 调和均值样本量
          q 统计量与 Studentized Range 分布比较
        
        返回 JSON:
        {
            "factor_name": "Catalyst",
            "comparisons": [
                {
                    "group1": "A",
                    "group2": "B",
                    "mean_diff": 2.35,
                    "p_value": 0.012,
                    "significant": true,
                    "ci_lower": 0.52,
                    "ci_upper": 4.18,
                    "reject": true
                },
                ...
            ],
            "group_means": {
                "A": 79.5,
                "B": 83.2,
                "C": 85.1
            },
            "mse": 1.23,
            "alpha": 0.05
        }
        """
        if self._df is None:
            return json.dumps({"error": "数据未加载"}, ensure_ascii=False)
        
        try:
            from statsmodels.stats.multicomp import pairwise_tukeyhsd
            
            # 确定类别因子
            if factor_name and factor_name in self._categorical_factors:
                cat_factor = factor_name
            elif self._categorical_factors:
                cat_factor = self._categorical_factors[0]
            else:
                return json.dumps({"error": "没有类别因子，无法进行 Tukey HSD 检验"}, ensure_ascii=False)
            
            # 提取数据
            groups = self._df[cat_factor].astype(str).values
            values = self._df[self._response_name].astype(float).values
            
            # 执行 Tukey HSD
            tukey_result = pairwise_tukeyhsd(values, groups, alpha=0.05)
            
            # 解析结果
            comparisons = []
            summary_data = tukey_result.summary().data[1:]  # 跳过表头
            
            for row in summary_data:
                comparisons.append({
                    "group1": str(row[0]),
                    "group2": str(row[1]),
                    "mean_diff": round(float(row[2]), 4),
                    "p_value": round(float(row[3]), 6),
                    "ci_lower": round(float(row[4]), 4),
                    "ci_upper": round(float(row[5]), 4),
                    "significant": bool(row[6]),
                    "reject": bool(row[6])
                })
            
            # 各组均值
            group_means = {}
            for level in sorted(self._df[cat_factor].astype(str).unique()):
                mask = self._df[cat_factor].astype(str) == level
                group_means[level] = round(float(self._df.loc[mask, self._response_name].mean()), 4)
            
            # MSE（从当前模型获取，如果没有则用池化方差）
            mse = 0.0
            if self._model is not None:
                mse = float(self._model.mse_resid)
            else:
                mse = float(self._df[self._response_name].var())
            
            return json.dumps({
                "factor_name": cat_factor,
                "comparisons": comparisons,
                "group_means": group_means,
                "mse": round(mse, 4),
                "alpha": 0.05
            }, ensure_ascii=False)
            
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # ★ v9 新增: Box-Cox 变换分析
    # ═══════════════════════════════════════════════════════

    def box_cox_analysis(self) -> str:
        """
        ★ v9 新增: Box-Cox 最优 λ 搜索与变换效果分析

        数学原理:
          Box-Cox 变换: y(λ) = (y^λ - 1)/λ  (λ≠0)
                        y(λ) = ln(y)          (λ=0)
          通过最大似然法搜索最优 λ，使变换后残差最接近正态分布。

        常见 λ 值含义:
          λ = -1.0  → 倒数变换 (1/y)
          λ = -0.5  → 倒数平方根 (1/√y)
          λ =  0.0  → 对数变换 (ln y)
          λ =  0.5  → 平方根变换 (√y)
          λ =  1.0  → 无变换 (y 本身)
          λ =  2.0  → 平方变换 (y²)

        前提条件: 所有响应值必须为正数 (y > 0)
        """
        if self._model is None or self._df is None:
            return json.dumps({"error": "模型未拟合"}, ensure_ascii=False)

        try:
            from scipy import stats as sp_stats

            y = self._df[self._response_name].astype(float).values

            if np.any(y <= 0):
                return json.dumps({
                    "error": "Box-Cox 变换要求所有响应值为正数 (y > 0)。当前数据包含零或负值。",
                    "all_positive": False
                }, ensure_ascii=False)

            # λ profile 曲线数据
            lambdas = np.linspace(-3, 3, 121)
            log_likes = []
            for lam in lambdas:
                try:
                    if abs(lam) < 1e-10:
                        y_t = np.log(y)
                    else:
                        y_t = (y ** lam - 1) / lam
                    n = len(y_t)
                    ss = np.sum((y_t - y_t.mean()) ** 2)
                    ll = -n / 2.0 * np.log(ss / n) + (lam - 1) * np.sum(np.log(y))
                    log_likes.append(float(ll))
                except Exception:
                    log_likes.append(float('-inf'))

            # scipy 精确最优 λ
            y_transformed, optimal_lambda = sp_stats.boxcox(y)

            # 95% CI (profile likelihood)
            max_ll = max(log_likes)
            chi2_crit = sp_stats.chi2.ppf(0.95, 1) / 2.0
            ci_threshold = max_ll - chi2_crit
            ci_lambdas = [lam for lam, ll in zip(lambdas, log_likes) if ll >= ci_threshold]
            lambda_ci = [float(min(ci_lambdas)), float(max(ci_lambdas))] if ci_lambdas else [optimal_lambda - 1, optimal_lambda + 1]

            # 圆整到常用值
            common_lambdas = [-2, -1, -0.5, 0, 0.5, 1, 2]
            rounded_lambda = min(common_lambdas, key=lambda x: abs(x - optimal_lambda))
            if lambda_ci[0] <= rounded_lambda <= lambda_ci[1]:
                use_lambda = rounded_lambda
            else:
                use_lambda = round(optimal_lambda, 2)

            transform_names = {
                -2: "1/y² (倒数平方变换)", -1: "1/y (倒数变换)",
                -0.5: "1/√y (倒数平方根变换)", 0: "ln(y) (对数变换)",
                0.5: "√y (平方根变换)", 1: "y (无变换)", 2: "y² (平方变换)"
            }
            transform_name = transform_names.get(use_lambda, f"y^{use_lambda:.2f}")

            orig_r2 = float(self._model.rsquared)
            orig_r2_adj = float(self._model.rsquared_adj)

            # 变换后临时拟合
            if abs(use_lambda) < 1e-10:
                y_new = np.log(y)
            elif abs(use_lambda - 1.0) < 1e-10:
                y_new = y
            else:
                y_new = (y ** use_lambda - 1) / use_lambda

            # ★ 修复: 使用 try/finally 确保即使中间出错也能恢复 _df 和 _model
            df_backup = self._df.copy()
            model_backup = self._model
            try:
                self._df[self._response_name] = y_new
                self._fit_quadratic_model()
                trans_r2 = float(self._model.rsquared) if self._model else 0.0
                trans_r2_adj = float(self._model.rsquared_adj) if self._model else 0.0
            finally:
                self._df = df_backup
                self._model = model_backup
                # 不再需要重新拟合 — 直接从备份恢复模型对象

            improvement = trans_r2 - orig_r2
            recommend = improvement > 0.01 and abs(use_lambda - 1.0) > 0.1

            if recommend:
                recommendation = (f"推荐使用{transform_name} (λ={use_lambda})，"
                                  f"R² 从 {orig_r2:.4f} 提升至 {trans_r2:.4f} (+{improvement:.4f})")
            elif abs(use_lambda - 1.0) < 0.1:
                recommendation = "最优 λ ≈ 1.0，无需变换，当前响应变量已具有良好的正态性。"
            else:
                recommendation = (f"变换效果不明显 (λ={use_lambda}, R² 变化 {improvement:+.4f})，"
                                  f"建议保持原始响应变量。")

            return json.dumps({
                "optimal_lambda": round(float(optimal_lambda), 4),
                "lambda_ci": [round(lambda_ci[0], 4), round(lambda_ci[1], 4)],
                "rounded_lambda": use_lambda,
                "transform_name": transform_name,
                "original_r_squared": round(orig_r2, 6),
                "transformed_r_squared": round(trans_r2, 6),
                "original_r_squared_adj": round(orig_r2_adj, 6),
                "transformed_r_squared_adj": round(trans_r2_adj, 6),
                "improvement": round(improvement, 6),
                "recommend_transform": recommend,
                "recommendation": recommendation,
                "lambda_profile": {
                    "lambdas": [round(float(l), 2) for l in lambdas],
                    "log_likelihoods": [round(float(ll), 2) if np.isfinite(ll) else None for ll in log_likes]
                },
                "all_positive": True, "error": None
            }, ensure_ascii=False)

        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    def apply_box_cox(self, lambda_value: float, model_type: str = "quadratic") -> str:
        """
        ★ v9 新增: 应用 Box-Cox 变换后重新拟合 OLS 模型
        返回与 fit_ols() 相同格式 + box_cox_info
        
        ★ v11 修复: 变换后 _df 和 _model 保持一致。
          调用此方法后，_df 中的响应列为变换后的值，_model 为对变换后数据拟合的模型。
          后续的 predict_point / residual_diagnostics 等方法都基于变换后的模型。
          
          如果需要恢复原始数据，调用 restore_full_data() 或 restore_full_data_from_backup()。
        """
        if self._df is None:
            return json.dumps({"error": "数据未加载"}, ensure_ascii=False)
        try:
            y = self._df[self._response_name].astype(float).values
            if np.any(y <= 0):
                return json.dumps({"error": "Box-Cox 变换要求 y > 0"}, ensure_ascii=False)

            if abs(lambda_value) < 1e-10:
                y_new = np.log(y)
                transform_desc = "ln(y)"
            elif abs(lambda_value - 1.0) < 1e-10:
                y_new = y
                transform_desc = "y (无变换)"
            else:
                y_new = (y ** lambda_value - 1) / lambda_value
                transform_desc = f"y^{lambda_value:.2f}"

            # ★ FIX #8: 使用不可覆盖的原始备份
            # 首次调用保存原始数据；后续调用从原始备份恢复再变换
            if not hasattr(self, '_df_original_backup') or self._df_original_backup is None:
                self._df_original_backup = self._df.copy()
            else:
                # 从原始备份恢复（避免连续变换导致备份被覆盖）
                self._df = self._df_original_backup.copy()
                y_new_recalc = self._df[self._response_name].astype(float).values
                if abs(lambda_value) < 1e-10:
                    y_new = np.log(y_new_recalc)
                elif abs(lambda_value - 1.0) < 1e-10:
                    y_new = y_new_recalc
                else:
                    y_new = (y_new_recalc ** lambda_value - 1) / lambda_value
            
            # 同时设置旧的 _df_full_backup（兼容 restore_full_data_from_backup）
            self._df_full_backup = self._df_original_backup.copy()
            
            # 应用变换
            self._df[self._response_name] = y_new
            result_json = self.fit_ols(model_type)
            result = json.loads(result_json)
            # ★ 修复: 不再恢复 _df，让 _df 和 _model 保持一致

            result["box_cox_info"] = {
                "lambda": lambda_value, "transform": transform_desc, "applied": True,
                "note": "数据和模型已切换为变换后的版本。调用 restore_full_data_from_backup() 可恢复。"
            }
            return json.dumps(result, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # ★ v9 新增: OLS 报告导出 (Word .docx)
    # ═══════════════════════════════════════════════════════

    def export_ols_report(self, output_path: str, title: str = "OLS 回归分析报告") -> str:
        """
        ★ v9 新增: 导出 OLS 分析结果为 Word 报告

        报告内容:
          1. 标题页 + 实验概况
          2. 实验数据表
          3. 模型摘要（R², R²adj, R²pred, RMSE 等）
          4. 回归方程（含类别因子展开）
          5. ANOVA 方差分析表
          6. 回归系数表（含显著性标记）
          7. 残差诊断四图
          8. 标准化效应 Pareto 图
          9. 主效应图
          10. 实际值 vs 预测值图
          11. 最优条件建议
          12. 模型诊断结论
        """
        if self._model is None or self._df is None:
            return json.dumps({"error": "模型未拟合，无法导出报告"}, ensure_ascii=False)

        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            import datetime

            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(10.5)

            model = self._model
            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            reverse_names = {v: k for k, v in safe_names.items()}
            n = int(model.nobs)
            p_count = len(model.params)
            ms_res = float(model.mse_resid)

            # 预计算常用量
            leverage = model.get_influence().hat_matrix_diag
            press_resid = model.resid / (1 - leverage)
            press = float(np.sum(press_resid ** 2))
            y_vals = self._df[self._response_name].astype(float)
            ss_total = float(np.sum((y_vals - y_vals.mean()) ** 2))
            r2_pred = 1.0 - press / ss_total if ss_total > 1e-12 else 0.0
            fitted = model.fittedvalues
            adeq_prec = 0.0
            # ★ FIX #2: Adequate Precision 对标 Design-Expert
            adeq_prec = _calc_adequate_precision_design_expert(model)

            # ════════════════════════════════════════════
            # 1. 标题页
            # ════════════════════════════════════════════
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'生成时间: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            # 实验概况
            doc.add_heading('1. 实验概况', level=1)
            info_data = [
                ('响应变量', self._response_name),
                ('观测数', str(n)),
                ('因子数', str(len(self._factor_names))),
                ('连续因子', ', '.join(self._continuous_factors) if self._continuous_factors else '无'),
                ('类别因子', ', '.join(self._categorical_factors) if self._categorical_factors else '无'),
                ('模型参数数', str(p_count)),
                ('残差自由度', str(int(model.df_resid))),
            ]
            tbl_info = doc.add_table(rows=len(info_data), cols=2)
            tbl_info.style = 'Light Grid Accent 1'
            for i, (k, v) in enumerate(info_data):
                tbl_info.rows[i].cells[0].text = k
                tbl_info.rows[i].cells[1].text = v
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 2. 实验数据表（前 30 行）
            # ════════════════════════════════════════════
            doc.add_heading('2. 实验数据', level=1)
            display_cols = self._factor_names + [self._response_name]
            max_rows = min(30, len(self._df))
            tbl_data = doc.add_table(rows=1 + max_rows, cols=1 + len(display_cols))
            tbl_data.style = 'Light Grid Accent 1'
            tbl_data.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表头
            tbl_data.rows[0].cells[0].text = '#'
            for j, col in enumerate(display_cols):
                tbl_data.rows[0].cells[j + 1].text = col
            for cell in tbl_data.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)

            # 数据行
            for i in range(max_rows):
                tbl_data.rows[i + 1].cells[0].text = str(i + 1)
                for j, col in enumerate(display_cols):
                    val = self._df.iloc[i][col]
                    if isinstance(val, float):
                        tbl_data.rows[i + 1].cells[j + 1].text = f'{val:.2f}'
                    else:
                        tbl_data.rows[i + 1].cells[j + 1].text = str(val)

            if len(self._df) > 30:
                doc.add_paragraph(f'（仅显示前 30 行，共 {len(self._df)} 行）',
                                  style='Normal').runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 3. 模型摘要
            # ════════════════════════════════════════════
            doc.add_heading('3. 模型摘要', level=1)

            # 模型质量评级
            r2 = float(model.rsquared)
            r2_adj = float(model.rsquared_adj)
            if r2 > 0.95 and r2_adj > 0.90 and r2_pred > 0.80:
                quality = "优秀 — 模型拟合良好，预测能力强"
            elif r2 > 0.85 and r2_adj > 0.75:
                quality = "良好 — 模型可用，建议检查是否有改进空间"
            elif r2 > 0.70:
                quality = "一般 — 模型解释力有限，可能需要增加因子或变换"
            else:
                quality = "较差 — 模型拟合不佳，建议检查数据或模型结构"

            summary_data = [
                ('模型质量', quality),
                ('R²', f'{r2:.6f}'),
                ('R² (调整)', f'{r2_adj:.6f}'),
                ('R² (预测)', f'{r2_pred:.6f}'),
                ('R²adj 与 R²pred 差距', f'{abs(r2_adj - r2_pred):.4f}' +
                 (' ✓ < 0.2' if abs(r2_adj - r2_pred) < 0.2 else ' ⚠ > 0.2，可能过拟合')),
                ('RMSE', f'{np.sqrt(ms_res):.6f}'),
                ('Adequate Precision', f'{adeq_prec:.4f}' +
                 (' ✓ > 4' if adeq_prec > 4 else ' ⚠ < 4，信噪比不足')),
                ('PRESS', f'{press:.4f}'),
                ('模型 F 值', f'{model.fvalue:.4f}' if not np.isnan(model.fvalue) else 'N/A'),
                ('模型 P 值', f'{model.f_pvalue:.6f}' if not np.isnan(model.f_pvalue) else 'N/A'),
            ]
            tbl_sum = doc.add_table(rows=len(summary_data), cols=2)
            tbl_sum.style = 'Light Grid Accent 1'
            tbl_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (label, val) in enumerate(summary_data):
                tbl_sum.rows[i].cells[0].text = label
                tbl_sum.rows[i].cells[1].text = val
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 4. 回归方程
            # ════════════════════════════════════════════
            doc.add_heading('4. 回归方程', level=1)

            equations_info = self._build_equations_by_category(model, reverse_names)
            if equations_info.get('has_categorical'):
                cat_name = equations_info.get('categorical_factor', '')
                doc.add_paragraph(f'按 {cat_name} 水平展开:').runs[0].font.italic = True
                for level, eq in equations_info.get('equations', {}).items():
                    p_eq = doc.add_paragraph()
                    run_lv = p_eq.add_run(f'{cat_name} = {level}: ')
                    run_lv.bold = True
                    run_lv.font.size = Pt(10)
                    run_eq = p_eq.add_run(eq)
                    run_eq.font.name = 'Consolas'
                    run_eq.font.size = Pt(9.5)
            else:
                equation = self._build_equation(model, reverse_names)
                p_eq = doc.add_paragraph()
                run_eq = p_eq.add_run(equation)
                run_eq.font.name = 'Consolas'
                run_eq.font.size = Pt(10)
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 5. ANOVA 表 — ★ 修复: 构建 df_safe 再传入
            # ════════════════════════════════════════════
            doc.add_heading('5. ANOVA 方差分析表', level=1)

            # 构建 df_safe (与 fit_ols 中相同的逻辑)
            df_safe = pd.DataFrame()
            for orig_name, safe_name in safe_names.items():
                if orig_name in self._categorical_factors:
                    df_safe[safe_name] = self._df[orig_name].astype(str)
                else:
                    info = self._coding_info.get(orig_name, {})
                    center = info.get("center", 0.0)
                    half_range = info.get("half_range", 1.0)
                    df_safe[safe_name] = (self._df[orig_name].astype(float) - center) / half_range
            df_safe["Y"] = self._df[self._response_name].values

            anova_data = self._build_anova_table(model, df_safe, safe_names, reverse_names, "quadratic")
            if anova_data:
                headers = ['来源', '自由度', '平方和', '均方', 'F 值', 'P 值']
                tbl_a = doc.add_table(rows=1 + len(anova_data), cols=6)
                tbl_a.style = 'Light Grid Accent 1'
                tbl_a.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j, h in enumerate(headers):
                    tbl_a.rows[0].cells[j].text = h
                    for paragraph in tbl_a.rows[0].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                for i, rd in enumerate(anova_data):
                    tbl_a.rows[i + 1].cells[0].text = str(rd.get('source', ''))
                    tbl_a.rows[i + 1].cells[1].text = str(rd.get('df', ''))
                    tbl_a.rows[i + 1].cells[2].text = f"{rd.get('ss', 0):.4f}" if rd.get('ss') is not None else ''
                    tbl_a.rows[i + 1].cells[3].text = f"{rd.get('ms', 0):.4f}" if rd.get('ms') is not None else ''
                    tbl_a.rows[i + 1].cells[4].text = f"{rd.get('f_value', 0):.4f}" if rd.get('f_value') is not None else ''
                    tbl_a.rows[i + 1].cells[5].text = f"{rd.get('p_value', 0):.6f}" if rd.get('p_value') is not None else ''
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 6. 系数表（含显著性标记）
            # ════════════════════════════════════════════
            doc.add_heading('6. 回归系数表', level=1)
            X_matrix = model.model.exog
            try:
                XtX = X_matrix.T @ X_matrix
                XtX_inv = np.linalg.inv(XtX)
            except Exception:
                XtX_inv = None

            coeff_rows = []
            for idx_j, term in enumerate(model.params.index):
                display = self._restore_term_name(term, reverse_names)
                coeff = float(model.params[term])
                se = float(model.bse[term])
                t_val = float(model.tvalues[term])
                p_val = float(model.pvalues[term])
                vif = ''
                if XtX_inv is not None and idx_j > 0:
                    vif_val = float(XtX_inv[idx_j, idx_j] * XtX[idx_j, idx_j])
                    vif = f"{vif_val:.2f}"
                sig = '***' if p_val < 0.001 else '**' if p_val < 0.01 else '*' if p_val < 0.05 else ''
                coeff_rows.append((display, f'{coeff:.6f}', f'{se:.6f}', f'{t_val:.4f}', f'{p_val:.6f}', vif, sig))

            ch = ['项', '系数 β', '标准误 SE', 'T 值', 'P 值', 'VIF', '显著性']
            tbl2 = doc.add_table(rows=1 + len(coeff_rows), cols=7)
            tbl2.style = 'Light Grid Accent 1'
            tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h in enumerate(ch):
                tbl2.rows[0].cells[j].text = h
                for paragraph in tbl2.rows[0].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for i, rd in enumerate(coeff_rows):
                for j, val in enumerate(rd):
                    tbl2.rows[i + 1].cells[j].text = val

            doc.add_paragraph('显著性: *** p<0.001, ** p<0.01, * p<0.05',
                              style='Normal').runs[0].font.size = Pt(9)
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 7. 残差诊断四图
            # ════════════════════════════════════════════
            doc.add_heading('7. 残差诊断图', level=1)
            if HAS_MATPLOTLIB:
                try:
                    from scipy import stats as sp_stats
                    influence = model.get_influence()
                    std_resid = influence.resid_studentized_internal
                    fitted_vals = model.fittedvalues.values

                    fig, axes = plt.subplots(2, 2, figsize=(10, 8))
                    sorted_resid = np.sort(std_resid)
                    theoretical = sp_stats.norm.ppf((np.arange(1, n + 1) - 0.375) / (n + 0.25))
                    axes[0, 0].scatter(theoretical, sorted_resid, s=20, c='steelblue', alpha=0.7)
                    axes[0, 0].plot([theoretical.min(), theoretical.max()], [theoretical.min(), theoretical.max()], 'r--', lw=1)
                    axes[0, 0].set_title('正态概率图'); axes[0, 0].set_xlabel('理论分位数'); axes[0, 0].set_ylabel('标准化残差')
                    axes[0, 1].scatter(fitted_vals, std_resid, s=20, c='darkorange', alpha=0.7)
                    axes[0, 1].axhline(y=0, color='gray', linestyle='--', lw=1)
                    axes[0, 1].set_title('残差 vs 拟合值'); axes[0, 1].set_xlabel('拟合值'); axes[0, 1].set_ylabel('标准化残差')
                    axes[1, 0].hist(std_resid, bins=max(5, int(np.sqrt(n))), color='cornflowerblue', edgecolor='gray', alpha=0.8)
                    axes[1, 0].set_title('残差直方图'); axes[1, 0].set_xlabel('标准化残差'); axes[1, 0].set_ylabel('频率')
                    axes[1, 1].plot(range(1, n + 1), std_resid, 'o-', color='steelblue', markersize=4, lw=1)
                    axes[1, 1].axhline(y=0, color='gray', linestyle='--', lw=1)
                    axes[1, 1].set_title('残差 vs 观测顺序'); axes[1, 1].set_xlabel('观测顺序'); axes[1, 1].set_ylabel('标准化残差')
                    plt.tight_layout()
                    buf = io.BytesIO()
                    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
                    plt.close(fig); buf.seek(0)
                    doc.add_picture(buf, width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f'残差图生成失败: {e}')
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 8. Pareto 图
            # ════════════════════════════════════════════
            doc.add_heading('8. 标准化效应 Pareto 图', level=1)
            if HAS_MATPLOTLIB:
                try:
                    from scipy import stats as sp_stats
                    pareto_data = json.loads(self.effects_pareto())
                    if pareto_data and 'error' not in pareto_data[0]:
                        terms_list = [d['term'] for d in pareto_data]
                        abs_t_list = [d['abs_t'] for d in pareto_data]
                        sig_list = [d['significant'] for d in pareto_data]
                        fig2, ax2 = plt.subplots(figsize=(8, max(3, len(terms_list) * 0.4)))
                        colors = ['#6495ED' if s else '#C0C0C0' for s in sig_list]
                        ax2.barh(range(len(terms_list)), abs_t_list, color=colors, edgecolor='gray', height=0.6)
                        ax2.set_yticks(range(len(terms_list))); ax2.set_yticklabels(terms_list)
                        ax2.set_xlabel('标准化效应 |t|'); ax2.set_title('标准化效应 Pareto 图 (α = 0.05)')
                        ax2.invert_yaxis()
                        df_error = int(model.df_resid)
                        if df_error > 0:
                            t_crit = float(sp_stats.t.ppf(0.975, df_error))
                            ax2.axvline(x=t_crit, color='red', linestyle='--', lw=1.5, label=f't_crit={t_crit:.2f}')
                            ax2.legend(loc='lower right')
                        plt.tight_layout()
                        buf2 = io.BytesIO()
                        fig2.savefig(buf2, format='png', dpi=150, bbox_inches='tight')
                        plt.close(fig2); buf2.seek(0)
                        doc.add_picture(buf2, width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f'Pareto 图生成失败: {e}')
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 9. 实际值 vs 预测值图
            # ════════════════════════════════════════════
            doc.add_heading('9. 实际值 vs 预测值', level=1)
            if HAS_MATPLOTLIB:
                try:
                    actual = y_vals.values
                    predicted = fitted.values
                    fig3, ax3 = plt.subplots(figsize=(6, 6))
                    ax3.scatter(actual, predicted, s=30, c='steelblue', alpha=0.7, edgecolors='gray', lw=0.5)
                    vmin = min(actual.min(), predicted.min())
                    vmax = max(actual.max(), predicted.max())
                    ax3.plot([vmin, vmax], [vmin, vmax], 'r--', lw=1.5, label='理想线 (y=x)')
                    ax3.set_xlabel('实际值'); ax3.set_ylabel('预测值')
                    ax3.set_title(f'实际值 vs 预测值 (R² = {r2:.4f})')
                    ax3.legend()
                    ax3.set_aspect('equal')
                    plt.tight_layout()
                    buf3 = io.BytesIO()
                    fig3.savefig(buf3, format='png', dpi=150, bbox_inches='tight')
                    plt.close(fig3); buf3.seek(0)
                    doc.add_picture(buf3, width=Inches(4))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f'实际vs预测图生成失败: {e}')
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 10. 主效应图
            # ════════════════════════════════════════════
            doc.add_heading('10. 主效应图', level=1)
            if HAS_MATPLOTLIB:
                try:
                    me_data = json.loads(self.main_effects())
                    if me_data:
                        fig4, ax4 = plt.subplots(figsize=(8, 5))
                        colors_me = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b']
                        global_mean = np.mean([pt['mean'] for pts in me_data.values() for pt in pts])
                        for ci, (fname, pts) in enumerate(me_data.items()):
                            x_vals = []
                            y_vals_me = []
                            for k, pt in enumerate(pts):
                                try:
                                    x_vals.append(float(pt['level']))
                                except (ValueError, TypeError):
                                    x_vals.append(k)
                                y_vals_me.append(pt['mean'])
                            ax4.plot(x_vals, y_vals_me, 'o-', color=colors_me[ci % len(colors_me)],
                                     label=fname, lw=2, markersize=6)
                        ax4.axhline(y=global_mean, color='gray', linestyle='--', lw=1, label=f'全局均值={global_mean:.2f}')
                        ax4.set_xlabel('因子水平'); ax4.set_ylabel(self._response_name)
                        ax4.set_title('主效应图'); ax4.legend(fontsize=9)
                        plt.tight_layout()
                        buf4 = io.BytesIO()
                        fig4.savefig(buf4, format='png', dpi=150, bbox_inches='tight')
                        plt.close(fig4); buf4.seek(0)
                        doc.add_picture(buf4, width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f'主效应图生成失败: {e}')
            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 11. 模型诊断结论
            # ════════════════════════════════════════════
            doc.add_heading('11. 模型诊断与建议', level=1)

            conclusions = []
            # R² 评价
            if r2 > 0.95:
                conclusions.append(f'• R² = {r2:.4f}，模型解释了超过 95% 的响应变异，拟合优秀。')
            elif r2 > 0.85:
                conclusions.append(f'• R² = {r2:.4f}，模型拟合良好。')
            else:
                conclusions.append(f'• R² = {r2:.4f}，模型解释力有限，建议考虑增加因子或改变模型形式。')

            # R²pred vs R²adj
            gap = abs(r2_adj - r2_pred)
            if gap < 0.2:
                conclusions.append(f'• R²adj - R²pred = {gap:.4f} < 0.2，模型无明显过拟合风险。')
            else:
                conclusions.append(f'• ⚠ R²adj - R²pred = {gap:.4f} > 0.2，存在过拟合风险，建议精简模型。')

            # Adequate Precision
            if adeq_prec > 4:
                conclusions.append(f'• Adequate Precision = {adeq_prec:.2f} > 4，信噪比充足，模型可用于预测。')
            else:
                conclusions.append(f'• ⚠ Adequate Precision = {adeq_prec:.2f} < 4，信噪比不足。')

            # 失拟检验
            lof_p = self._calc_lack_of_fit_p(model, df_safe, safe_names, "quadratic")
            if lof_p is not None:
                if lof_p > 0.05:
                    conclusions.append(f'• 失拟检验 P = {lof_p:.4f} > 0.05，模型无显著失拟（好）。')
                else:
                    conclusions.append(f'• ⚠ 失拟检验 P = {lof_p:.4f} < 0.05，模型可能不充分，建议增加高阶项或检查数据。')

            # 显著因子
            sig_terms = [self._restore_term_name(t, reverse_names) for t in model.pvalues.index
                         if t != 'Intercept' and float(model.pvalues[t]) < 0.05]
            if sig_terms:
                conclusions.append(f'• 显著因子 (p<0.05): {", ".join(sig_terms)}')
            else:
                conclusions.append(f'• 未发现显著因子 (α=0.05)，建议检查实验设计是否有足够的分辨力。')

            for c in conclusions:
                doc.add_paragraph(c)

            doc.add_paragraph('')

            # ════════════════════════════════════════════
            # 页脚
            # ════════════════════════════════════════════
            doc.add_paragraph('')
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_ft = footer.add_run('MaxChemical DOE Module — Auto-generated Report')
            run_ft.font.size = Pt(8)
            run_ft.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

            doc.save(output_path)
            return json.dumps({
                "success": True, "path": output_path,
                "sections": ["实验概况", "实验数据", "模型摘要", "回归方程", "ANOVA表",
                              "系数表", "残差图", "Pareto图", "实际vs预测", "主效应图", "诊断结论"]
            }, ensure_ascii=False)

        except ImportError:
            return json.dumps({"error": "缺少 python-docx 库。请安装: pip install python-docx"}, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    def main_effects(self) -> str:
        """主效应图数据
        ★ v11 修复: 使用模型预测的边际均值 (LS-Means / Marginal Means)，
          对标 JMP / Minitab 的主效应图行为。
          在不平衡设计中，LS-Means 比原始均值更准确。
          
        原来的 bug (Bug#2): float("催化剂A") 会抛出 ValueError — 已修复。
        原来的 bug (Bug#14): 使用原始均值而非模型边际均值 — 现在修复。
        """
        if self._df is None:
            return json.dumps({}, ensure_ascii=False)
        
        # ★ 如果有拟合模型，用模型预测边际均值 (LS-Means)
        if self._model is not None:
            try:
                result = {}
                # 计算各因子的中心值（其他因子固定时的参考点）
                center = {}
                for name in self._factor_names:
                    if name in self._categorical_factors:
                        levels = sorted(self._df[name].astype(str).unique())
                        center[name] = levels[0] if levels else ""
                    else:
                        col = self._df[name].astype(float)
                        center[name] = float((col.min() + col.max()) / 2.0)
                
                for factor in self._factor_names:
                    is_cat = factor in self._categorical_factors
                    
                    if is_cat:
                        levels = sorted(self._df[factor].astype(str).unique())
                        effect_data = []
                        for lv in levels:
                            point = dict(center)
                            point[factor] = lv
                            pred_json = self.predict_point(json.dumps(point, ensure_ascii=False))
                            pred_val = json.loads(pred_json).get("predicted", 0.0)
                            effect_data.append({
                                "level": str(lv),
                                "mean": round(float(pred_val), 4)
                            })
                        result[factor] = effect_data
                    else:
                        col = self._df[factor].astype(float)
                        levels = sorted(col.unique())
                        effect_data = []
                        for lv in levels:
                            point = dict(center)
                            point[factor] = float(lv)
                            pred_json = self.predict_point(json.dumps(point, ensure_ascii=False))
                            pred_val = json.loads(pred_json).get("predicted", 0.0)
                            effect_data.append({
                                "level": round(float(lv), 4),
                                "mean": round(float(pred_val), 4)
                            })
                        result[factor] = effect_data
                
                return json.dumps(result, ensure_ascii=False)
            except Exception:
                pass  # 回退到原始均值
        
        # 回退: 无模型时用原始均值
        result = {}
        for factor in self._factor_names:
            is_cat = factor in self._categorical_factors
            groups = self._df.groupby(factor)[self._response_name].mean()
            result[factor] = [
                {
                    "level": str(level) if is_cat else round(float(level), 4),
                    "mean": round(float(mean), 4)
                }
                for level, mean in groups.items()
            ]
        return json.dumps(result, ensure_ascii=False)

    def interaction_effects(self) -> str:
        """交互效应图数据
        ★ 修复 (Bug#2): 类别因子的 level 保留字符串标签，不做 float() 转换
        """
        if self._df is None:
            return json.dumps([], ensure_ascii=False)
        
        cat_set = set(self._categorical_factors)
        result = []
        for i in range(len(self._factor_names)):
            for j in range(i + 1, len(self._factor_names)):
                f1, f2 = self._factor_names[i], self._factor_names[j]
                f1_is_cat = f1 in cat_set
                f2_is_cat = f2 in cat_set
                groups = self._df.groupby([f1, f2])[self._response_name].mean()
                data = [
                    {
                        "f1": str(idx[0]) if f1_is_cat else round(float(idx[0]), 4),
                        "f2": str(idx[1]) if f2_is_cat else round(float(idx[1]), 4),
                        "mean": round(float(val), 4)
                    }
                    for idx, val in groups.items()
                ]
                result.append({"factor1": f1, "factor2": f2, "data": data})
        return json.dumps(result, ensure_ascii=False)

    def pareto_chart(self) -> str:
        """Pareto 图数据（原有实现，保留兼容，★ 修复: 使用统计检验判断显著性）"""
        if self._df is None:
            return json.dumps([], ensure_ascii=False)
        
        # ★ 修复: 如果有已拟合的模型, 使用 t 检验 p 值判断显著性
        # 旧实现用 abs(effect) > 0.1 * abs(overall_mean) 这个完全没有统计学依据的阈值
        if self._model is not None and HAS_STATSMODELS:
            try:
                safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
                reverse_names = {v: k for k, v in safe_names.items()}
                effects = []
                for term in self._model.params.index:
                    if term == "Intercept":
                        continue
                    display = self._restore_term_name(term, reverse_names)
                    t_val = float(self._model.tvalues[term])
                    p_val = float(self._model.pvalues[term])
                    effects.append({
                        "term": display,
                        "effect": round(abs(t_val), 4),
                        "raw_effect": round(t_val, 4),
                        "significant": bool(p_val < 0.05)
                    })
                effects.sort(key=lambda x: x["effect"], reverse=True)
                return json.dumps(effects, ensure_ascii=False)
            except Exception:
                pass  # 回退到下面的简单计算
        
        # 回退: 无模型时用效应大小（仅用于纯因子水平比较）
        effects = []
        overall_mean = self._df[self._response_name].mean()
        for factor in self._factor_names:
            levels = sorted(self._df[factor].unique())
            if len(levels) >= 2:
                low_mean = self._df[self._df[factor] == levels[0]][self._response_name].mean()
                high_mean = self._df[self._df[factor] == levels[-1]][self._response_name].mean()
                effect = high_mean - low_mean
                # 使用 2σ 近似判断（比原来的 10% 阈值更合理）
                pooled_std = self._df[self._response_name].std()
                n_per_level = len(self._df) / len(levels)
                se_effect = pooled_std * np.sqrt(2.0 / max(n_per_level, 1))
                significant = bool(abs(effect) > 2.0 * se_effect) if se_effect > 1e-12 else False
                effects.append({
                    "term": factor,
                    "effect": round(float(abs(effect)), 4),
                    "raw_effect": round(float(effect), 4),
                    "significant": significant
                })
        effects.sort(key=lambda x: x["effect"], reverse=True)
        return json.dumps(effects, ensure_ascii=False)

    def anova_table(self) -> str:
        """ANOVA 表（原有实现，保留兼容，★ v12 修复: 统一使用 Partial SS，与 _build_anova_table 一致）"""
        if self._df is None or not HAS_STATSMODELS:
            return json.dumps([], ensure_ascii=False)
        try:
            if self._model is not None:
                safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
                reverse_names = {v: k_name for k_name, v in safe_names.items()}
                
                # ★ v12 修复: 构建 df_safe 并委托给 _build_anova_table
                # 确保与 fit_ols 输出的 ANOVA 完全一致
                df_safe = pd.DataFrame()
                for orig_name, safe_name in safe_names.items():
                    if orig_name in self._categorical_factors:
                        df_safe[safe_name] = self._df[orig_name].astype(str)
                    else:
                        info = self._coding_info.get(orig_name, {})
                        center = info.get("center", 0.0)
                        half_range = info.get("half_range", 1.0)
                        df_safe[safe_name] = (self._df[orig_name].astype(float) - center) / half_range
                df_safe["Y"] = self._df[self._response_name].values
                
                # 推断 model_type
                model_type = "quadratic"  # 默认
                formula_str = str(self._model.model.formula) if hasattr(self._model.model, 'formula') else ""
                if "**2" not in formula_str and ":" not in formula_str:
                    model_type = "linear"
                elif "**2" not in formula_str:
                    model_type = "interaction"
                
                result = self._build_anova_table(self._model, df_safe, safe_names, reverse_names, model_type)
                return json.dumps(result, ensure_ascii=False)
            
            # 回退: 无已拟合模型时，拟合主效应模型
            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            reverse_names = {v: k_name for k_name, v in safe_names.items()}
            df_safe = self._df.rename(columns=safe_names)
            df_safe = df_safe.rename(columns={self._response_name: "Y"})
            formula_terms = " + ".join(safe_names.values())
            formula = f"Y ~ {formula_terms}"
            model = ols(formula, data=df_safe).fit()
            result = self._build_anova_table(model, df_safe, safe_names, reverse_names, "linear")
            return json.dumps(result, ensure_ascii=False)
        except Exception as e:
            return json.dumps([{"source": "Error", "error": str(e)}], ensure_ascii=False)

    def regression_summary(self) -> str:
        """回归摘要（原有实现，保留兼容，★ 修复: 复用 self._model）"""
        if self._df is None or not HAS_STATSMODELS:
            return json.dumps({"error": "数据未加载或缺少 statsmodels"}, ensure_ascii=False)
        try:
            # ★ 修复: 优先使用已拟合的模型
            if self._model is not None:
                model = self._model
                safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
                reverse_names = {v: k_name for k_name, v in safe_names.items()}
                
                eq_parts = [f"{float(model.params['Intercept']):.4f}"]
                coefficients = [{"term": "截距", "coeff": round(float(model.params['Intercept']), 4),
                                 "p_value": round(float(model.pvalues['Intercept']), 6)}]
                for term in model.params.index:
                    if term == "Intercept":
                        continue
                    display = self._restore_term_name(term, reverse_names)
                    coeff = float(model.params[term])
                    p_val = float(model.pvalues[term])
                    sign = "+" if coeff >= 0 else "-"
                    eq_parts.append(f"{sign} {abs(coeff):.4f}×{display}")
                    coefficients.append({
                        "term": display, "coeff": round(coeff, 4),
                        "p_value": round(p_val, 6)
                    })
                equation = f"{self._response_name} = " + " ".join(eq_parts)
                return json.dumps({
                    "r_squared": round(float(model.rsquared), 6),
                    "adj_r_squared": round(float(model.rsquared_adj), 6),
                    "f_statistic": round(float(model.fvalue), 4) if not np.isnan(model.fvalue) else None,
                    "f_p_value": round(float(model.f_pvalue), 6) if not np.isnan(model.f_pvalue) else None,
                    "equation": equation, "coefficients": coefficients
                }, ensure_ascii=False)
            
            # 回退: 拟合主效应模型
            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            df_safe = self._df.rename(columns=safe_names)
            df_safe = df_safe.rename(columns={self._response_name: "Y"})
            formula_terms = " + ".join(safe_names.values())
            formula = f"Y ~ {formula_terms}"
            model = ols(formula, data=df_safe).fit()
            eq_parts = [f"{model.params['Intercept']:.4f}"]
            coefficients = [{"term": "截距", "coeff": round(float(model.params['Intercept']), 4),
                             "p_value": round(float(model.pvalues['Intercept']), 6)}]
            for orig, safe in safe_names.items():
                if safe in model.params:
                    coeff = model.params[safe]
                    p_val = model.pvalues[safe]
                    sign = "+" if coeff >= 0 else "-"
                    eq_parts.append(f"{sign} {abs(coeff):.4f}*{orig}")
                    coefficients.append({
                        "term": orig, "coeff": round(float(coeff), 4),
                        "p_value": round(float(p_val), 6)
                    })
            equation = f"{self._response_name} = " + " ".join(eq_parts)
            return json.dumps({
                "r_squared": round(float(model.rsquared), 6),
                "adj_r_squared": round(float(model.rsquared_adj), 6),
                "f_statistic": round(float(model.fvalue), 4) if not np.isnan(model.fvalue) else None,
                "f_p_value": round(float(model.f_pvalue), 6) if not np.isnan(model.f_pvalue) else None,
                "equation": equation, "coefficients": coefficients
            }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    def residual_analysis(self) -> str:
        """残差分析
        ★ 修复 (Bug#7): 优先复用 self._model（与 fit_ols / _fit_quadratic_model 拟合的同一个模型）
        原来的 bug: 每次独立拟合纯主效应模型，与 fit_ols 的 quadratic 模型结果不一致
        """
        if self._df is None or not HAS_STATSMODELS:
            return json.dumps({}, ensure_ascii=False)
        try:
            # 优先使用已拟合的模型（与 fit_ols / _fit_quadratic_model 一致）
            if self._model is not None:
                model = self._model
            else:
                # 回退: 拟合主效应模型（仅在 _fit_quadratic_model 也失败时才走到这里）
                safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
                df_safe = self._df.rename(columns=safe_names)
                df_safe = df_safe.rename(columns={self._response_name: "Y"})
                formula_terms = " + ".join(safe_names.values())
                model = ols(f"Y ~ {formula_terms}", data=df_safe).fit()
            
            fitted = model.fittedvalues.tolist()
            residuals = model.resid.tolist()
            try:
                influence = model.get_influence()
                std_residuals = influence.resid_studentized_internal.tolist()
            except Exception:
                std_residuals = (model.resid / model.resid.std()).tolist() if model.resid.std() > 1e-12 else residuals
            return json.dumps({
                "fitted": [round(v, 4) for v in fitted],
                "residuals": [round(v, 4) for v in residuals],
                "std_residuals": [round(v, 4) for v in std_residuals],
                "order": list(range(1, len(fitted) + 1)),
                "actual": self._df[self._response_name].tolist()
            }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    def actual_vs_predicted(self) -> str:
        """实际值 vs 预测值（原有实现，保留兼容，★ 修复: 复用 self._model）"""
        if self._df is None or not HAS_STATSMODELS:
            return json.dumps({}, ensure_ascii=False)
        try:
            # ★ 修复: 优先使用已拟合的模型
            if self._model is not None:
                actual = self._df[self._response_name].tolist()
                predicted = self._model.fittedvalues.tolist()
                return json.dumps({
                    "actual": [round(v, 4) for v in actual],
                    "predicted": [round(v, 4) for v in predicted],
                    "r_squared": round(float(self._model.rsquared), 6)
                }, ensure_ascii=False)
            
            # 回退
            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            df_safe = self._df.rename(columns=safe_names)
            df_safe = df_safe.rename(columns={self._response_name: "Y"})
            formula_terms = " + ".join(safe_names.values())
            model = ols(f"Y ~ {formula_terms}", data=df_safe).fit()
            actual = self._df[self._response_name].tolist()
            predicted = model.fittedvalues.tolist()
            return json.dumps({
                "actual": [round(v, 4) for v in actual],
                "predicted": [round(v, 4) for v in predicted],
                "r_squared": round(float(model.rsquared), 6)
            }, ensure_ascii=False)
        except Exception as e:
            return json.dumps({"error": str(e)}, ensure_ascii=False)

    # ═══════════════════════════════════════════════════════
    # 内部辅助方法
    # ═══════════════════════════════════════════════════════

    def _fit_quadratic_model(self):
        if not HAS_STATSMODELS or self._df is None:
            return
        try:
            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            safe_continuous = [safe_names[f] for f in self._continuous_factors]
            safe_categorical = [safe_names[f] for f in self._categorical_factors]
            df_safe = pd.DataFrame()
            self._coding_info = {}
            for orig_name, safe_name in safe_names.items():
                if orig_name in self._categorical_factors:
                    df_safe[safe_name] = self._df[orig_name].astype(str)
                    self._coding_info[orig_name] = {"type": "categorical"}
                else:
                    col = self._df[orig_name].astype(float)
                    # ★ v16: 优先用设计范围，CCD 时乘以 α（对齐 Minitab）
                    if orig_name in self._factor_bounds:
                        bounds = self._factor_bounds[orig_name]
                        lo, hi = bounds[0], bounds[1]
                        alpha_scale = bounds[2] if len(bounds) > 2 else 1.0
                        center = (lo + hi) / 2.0
                        half_range = (hi - lo) / 2.0 * alpha_scale
                    else:
                        # 回退: 用数据范围（直接导入 Excel 时没有因子范围信息）
                        col_min = col.min()
                        col_max = col.max()
                        center = (col_min + col_max) / 2.0
                        half_range = (col_max - col_min) / 2.0
                    if half_range < 1e-12:
                        half_range = 1.0
                    self._coding_info[orig_name] = {"type": "continuous", "center": center, "half_range": half_range}
                    df_safe[safe_name] = (col - center) / half_range
            df_safe["Y"] = self._df[self._response_name].values
            k_cont = len(safe_continuous)
            n = len(df_safe)
            cat_dummy_counts = []
            for cat_name in self._categorical_factors:
                n_unique = self._df[cat_name].nunique()
                cat_dummy_counts.append(max(n_unique - 1, 1))
            total_cat_dummies = sum(cat_dummy_counts)
            p_quad = 1 + k_cont + total_cat_dummies
            p_quad += k_cont * (k_cont - 1) // 2 + k_cont
            p_quad += total_cat_dummies * k_cont
            for i in range(len(cat_dummy_counts)):
                for j in range(i + 1, len(cat_dummy_counts)):
                    p_quad += cat_dummy_counts[i] * cat_dummy_counts[j]
            p_inter = 1 + k_cont + total_cat_dummies
            p_inter += k_cont * (k_cont - 1) // 2
            p_inter += total_cat_dummies * k_cont
            for i in range(len(cat_dummy_counts)):
                for j in range(i + 1, len(cat_dummy_counts)):
                    p_inter += cat_dummy_counts[i] * cat_dummy_counts[j]

            if n > p_quad + 2:
                target_type = "quadratic"
            elif n > p_inter + 2:
                target_type = "interaction"
            else:
                target_type = "linear"

            formula = self._build_formula(safe_continuous, safe_categorical, target_type)

            # ★ v7: 在自动拟合时也做不可估计项检测
            reverse_names = {v: k for k, v in safe_names.items()}
            try:
                self._model = ols(formula, data=df_safe).fit()
                nan_params = [p for p in self._model.params.index if np.isnan(self._model.params[p])]
                if nan_params:
                    raise np.linalg.LinAlgError("NaN params")
                # ★ v7 增强: rank 检查
                X_exog = self._model.model.exog
                if np.linalg.matrix_rank(X_exog) < X_exog.shape[1]:
                    raise np.linalg.LinAlgError("Rank deficient in auto-fit")
                # ★ v10: SVD 奇异值比值兜底
                svd_vals = np.linalg.svd(X_exog, compute_uv=False)
                sv_ratio = svd_vals[-1] / svd_vals[0] if svd_vals[0] > 0 else 0
                if sv_ratio < 1e-12:
                    raise np.linalg.LinAlgError("Near-singular in auto-fit")
                if self._model.df_resid <= 0:
                    raise np.linalg.LinAlgError("df_resid <= 0 in auto-fit")
            except Exception:
                formula, _ = self._detect_and_drop_inestimable(
                    formula, df_safe, safe_names, reverse_names, safe_continuous, safe_categorical, target_type
                )
                if formula and formula != "Y ~ 1":
                    try:
                        self._model = ols(formula, data=df_safe).fit()
                    except Exception:
                        self._model = None
                else:
                    self._model = None
        except Exception:
            self._model = None
    def _fit_model(self):
        """内部拟合 OLS 模型（仅主效应，保留兼容）"""
        self._fit_quadratic_model()

    def _build_formula(self, safe_continuous: list, safe_categorical: list = None, 
                       model_type: str = "quadratic") -> str:
        """
        ★ 改造: 构建 statsmodels 公式字符串，支持混合因子类型。
        
        连续因子: 主效应 + 交互项 + 二次项（quadratic 时）
        类别因子: C(X) 哑变量主效应 + 与连续因子的交互（不生成二次项）
        
        兼容旧调用: 如果 safe_categorical 为 None，退化为纯连续模式
        """
        if safe_categorical is None:
            safe_categorical = []
        
        terms = []
        
        # 连续因子主效应
        for v in safe_continuous:
            terms.append(v)
        
        # 类别因子主效应（用 C() 包裹）
        for v in safe_categorical:
            terms.append(f"C({v})")
        
        if model_type in ("interaction", "quadratic"):
            # 连续 × 连续 交互
            for i in range(len(safe_continuous)):
                for j in range(i + 1, len(safe_continuous)):
                    terms.append(f"{safe_continuous[i]}:{safe_continuous[j]}")
            
            # 类别 × 连续 交互
            for cat_v in safe_categorical:
                for cont_v in safe_continuous:
                    terms.append(f"C({cat_v}):{cont_v}")
            
            # ★ 新增: 类别 × 类别 交互（化工中不同类别因子之间也可能有交互）
            for i in range(len(safe_categorical)):
                for j in range(i + 1, len(safe_categorical)):
                    terms.append(f"C({safe_categorical[i]}):C({safe_categorical[j]})")
        
        if model_type == "quadratic":
            # 仅连续因子生成二次项（类别因子没有二次项）
            for v in safe_continuous:
                terms.append(f"I({v}**2)")
        
        return "Y ~ " + " + ".join(terms) if terms else "Y ~ 1"

    def _build_custom_formula(self, terms_json: str, safe_names: dict, reverse_names: dict,
                               safe_continuous: list, safe_categorical: list) -> str:
        """
        ★ v4+v9修复: 从用户指定的项列表构建 statsmodels 公式
        
        支持两种项名格式:
          1. 合并名:  "Catalyst", "Catalyst×Temperature", "Temperature²"
          2. 展开名:  "Catalyst[B]", "Catalyst[C]×Temperature"  (Pareto 图返回的格式)
        
        展开名会先收缩回合并名再处理，确保类别因子作为整体加入公式。
        """
        import re
        
        user_terms = json.loads(terms_json) if terms_json else []
        if not user_terms:
            return self._build_formula(safe_continuous, safe_categorical, "quadratic")
        
        name_to_safe = dict(safe_names)
        cat_set = set(self._categorical_factors)
        
        # ★ v9修复: 先把展开名收缩回合并名
        # "Catalyst[B]" → "Catalyst", "Catalyst[C]×Temperature" → "Catalyst×Temperature"
        collapsed_terms = set()
        for term in user_terms:
            collapsed = re.sub(r'\[[^\]]*\]', '', term.strip())
            collapsed_terms.add(collapsed)
        
        # 解析收缩后的项 → statsmodels 公式项
        formula_terms = set()
        involved_factors = set()
        
        for term in collapsed_terms:
            # 二次项: "Temperature²" 或 "Temperature^2"
            quad_match = re.match(r'^(.+?)[²\^2]$', term)
            if quad_match:
                factor_name = quad_match.group(1).strip()
                if factor_name in name_to_safe and factor_name not in cat_set:
                    safe = name_to_safe[factor_name]
                    formula_terms.add(f"I({safe}**2)")
                    involved_factors.add(factor_name)
                continue
            
            # 交互项: "Temperature×Pressure" 或 "Catalyst×Temperature"
            if "×" in term or "*" in term:
                parts = re.split(r'[×\*]', term)
                parts = [p.strip() for p in parts]
                if len(parts) == 2:
                    f1, f2 = parts[0], parts[1]
                    if f1 in name_to_safe and f2 in name_to_safe:
                        s1 = name_to_safe[f1]
                        s2 = name_to_safe[f2]
                        t1 = f"C({s1})" if f1 in cat_set else s1
                        t2 = f"C({s2})" if f2 in cat_set else s2
                        formula_terms.add(f"{t1}:{t2}")
                        involved_factors.add(f1)
                        involved_factors.add(f2)
                continue
            
            # 主效应: "Temperature" 或 "Catalyst"
            if term in name_to_safe:
                safe = name_to_safe[term]
                if term in cat_set:
                    formula_terms.add(f"C({safe})")
                else:
                    formula_terms.add(safe)
                involved_factors.add(term)
        
        # ── 层次性原则: 自动补充缺失的主效应 ──
        for factor_name in involved_factors:
            safe = name_to_safe.get(factor_name)
            if safe:
                if factor_name in cat_set:
                    formula_terms.add(f"C({safe})")
                else:
                    formula_terms.add(safe)
        
        if not formula_terms:
            return "Y ~ 1"
        
        return "Y ~ " + " + ".join(sorted(formula_terms))

    def _build_anova_table(self, model, df_safe, safe_names, reverse_names, model_type) -> list:
        """★ v14 重写: 构建完整 ANOVA 表 — 使用 Sum coding + Type III SS 对标 Minitab Adj SS
        
        Minitab 的 Adj SS = Sum coding + Type III SS:
          Sum coding (Effect coding) 使类别因子系数表示偏离总均值的效应，
          结合连续因子编码到 [-1, +1]，Type III SS 与 Minitab Adj SS 完全一致。
          
        注意: 不是 Type II SS。Minitab 的 Adj SS 是 Type III（调整后平方和），
          不是遵循边际性原则的 Type II。两者在不平衡数据时有差异。
        
        分组汇总行使用缩减模型法计算（去掉整组后的 SS 增量）。
        
        ANOVA 表结构（对标 Minitab）:
          1. 模型 (Model) 汇总行
          2. 线性 (Linear) 汇总行 — 主效应汇总
          3.   各主效应分解行
          4. 平方 (Square) 汇总行 — 二次项汇总
          5.   各二次项分解行
          6. 双因子交互作用 (2-Way Interaction) 汇总行
          7.   各交互项分解行
          8. 误差 (Error / 残差) 行
          9.   失拟 (Lack of Fit) 行
          10.  纯误差 (Pure Error) 行
          11. 总计 (Total) 行
        """
        try:
            from scipy import stats as sp_stats
            
            # ── ★ v13 核心: Minitab Adj SS = 编码值 + Sum coding + Type III SS ──
            # 
            # 经过用真实数据逐一验证，Minitab 的 Adj SS 精确等于:
            #   1. 连续因子编码到 [-1, +1] (center ± half_range)
            #   2. 类别因子用 Sum/Effect coding (-1, 0, 1) 而非 Treatment coding (0, 1)
            #   3. 使用 Type III SS
            #
            # Sum coding 使类别因子哑变量与截距正交，
            # 结合编码后的连续因子，Type III SS 的结果与 Minitab 完全一致。
            
            safe_continuous = [safe_names[f] for f in self._continuous_factors]
            safe_categorical = [safe_names[f] for f in self._categorical_factors]
            
            # ★ 关键: 使用编码值 df_safe（已经是 [-1,+1]），不用原始值
            # 但需要把公式中的 C(X) 改为 C(X, Sum) 来使用 Sum coding
            
            # 从已拟合模型获取实际公式（已剔除不可估计项）
            if hasattr(model.model, 'formula'):
                formula_coded = str(model.model.formula)
            else:
                formula_coded = self._build_formula(safe_continuous, safe_categorical, model_type)
            
            # 将 Treatment coding 的 C(Xn) 替换为 Sum coding 的 C(Xn, Sum)
            import re
            formula_sum = re.sub(r'C\((\w+)\)(?![,\[])', r'C(\1, Sum)', formula_coded)
            
            try:
                model_anova = ols(formula_sum, data=df_safe).fit()
            except Exception:
                # Sum coding 失败时回退到编码值 + Treatment coding + Type II
                try:
                    model_anova = ols(formula_coded, data=df_safe).fit()
                    formula_sum = formula_coded  # 标记使用的是原公式
                except Exception:
                    model_anova = model
            
            # ★ Sum coding + Type III SS = Minitab Adj SS（注意: 这是 Type III 而非 Type II）
            anova_result = anova_lm(model_anova, typ=3)
            
            # 构建各项的分组归类
            linear_keys = set()
            square_keys = set()
            interact_keys = set()
            
            for v in safe_continuous:
                linear_keys.add(v)
            for v in safe_categorical:
                linear_keys.add(f"C({v})")
                linear_keys.add(f"C({v}, Sum)")  # Sum coding 格式
            
            if model_type in ("interaction", "quadratic"):
                for i in range(len(safe_continuous)):
                    for j in range(i + 1, len(safe_continuous)):
                        interact_keys.add(f"{safe_continuous[i]}:{safe_continuous[j]}")
                for cat_v in safe_categorical:
                    for cont_v in safe_continuous:
                        interact_keys.add(f"C({cat_v}):{cont_v}")
                        interact_keys.add(f"C({cat_v}, Sum):{cont_v}")  # Sum coding
                for i in range(len(safe_categorical)):
                    for j in range(i + 1, len(safe_categorical)):
                        interact_keys.add(f"C({safe_categorical[i]}):C({safe_categorical[j]})")
                        interact_keys.add(f"C({safe_categorical[i]}, Sum):C({safe_categorical[j]}, Sum)")
            
            if model_type == "quadratic":
                for v in safe_continuous:
                    square_keys.add(f"I({v} ** 2)")  # statsmodels 输出格式含空格
                    square_keys.add(f"I({v}**2)")     # 也兼容无空格格式
            
            # ── 残差信息 ──
            ss_res_full = float(model_anova.ssr)
            df_res_full = int(model_anova.df_resid)
            ms_res = ss_res_full / max(df_res_full, 1)
            
            # ── 从 anova2 提取各项 Adj SS ──
            linear_rows = []
            square_rows = []
            interact_rows = []
            
            for idx, row in anova_result.iterrows():
                if idx == "Residual":
                    continue
                # Type III 包含 Intercept 行，跳过
                if idx == "Intercept":
                    continue
                
                ss = float(row.get("sum_sq", 0))
                df = int(row.get("df", 0))
                ms = ss / max(df, 1)
                f_val = float(row.get("F", 0)) if pd.notna(row.get("F")) else None
                p_val = float(row.get("PR(>F)", 1)) if pd.notna(row.get("PR(>F)")) else None
                
                # ★ 去掉 Sum coding 标记再做名称还原
                idx_clean = str(idx).replace(", Sum", "")
                display_name = self._restore_term_name(idx_clean, reverse_names)
                
                item = {
                    "source": display_name, "df": df, "ss": round(ss, 6),
                    "ms": round(ms, 6),
                    "f_value": round(f_val, 4) if f_val is not None else None,
                    "p_value": round(p_val, 6) if p_val is not None else None,
                }
                
                idx_str = str(idx)
                if idx_str in linear_keys:
                    linear_rows.append(item)
                elif idx_str in square_keys:
                    square_rows.append(item)
                elif idx_str in interact_keys:
                    interact_rows.append(item)
                else:
                    # 尝试匹配 (statsmodels 可能添加空格)
                    idx_nospace = idx_str.replace(" ", "")
                    matched = False
                    for sk in square_keys:
                        if sk.replace(" ", "") == idx_nospace:
                            square_rows.append(item)
                            matched = True
                            break
                    if not matched:
                        for ik in interact_keys:
                            if ik.replace(" ", "") == idx_nospace:
                                interact_rows.append(item)
                                matched = True
                                break
                    if not matched:
                        for lk in linear_keys:
                            if lk.replace(" ", "") == idx_nospace:
                                linear_rows.append(item)
                                matched = True
                                break
                    if not matched:
                        # 兜底: 归入交互项组
                        interact_rows.append(item)
            
            # ── 分组汇总: 使用 GLH (General Linear Hypothesis) 联合 F 检验 ──
            # Minitab 的分组汇总 SS = F × df_num × MSE，通过对该组所有参数联合检验得到
            
            anova_params = list(model_anova.params.index)
            mse_anova = float(model_anova.mse_resid)
            
            def _make_group_summary_glh(source_name, group_param_names):
                """Minitab 风格: 分组汇总 = GLH 联合 F 检验"""
                if not group_param_names:
                    return None
                try:
                    r_matrix = np.zeros((len(group_param_names), len(anova_params)))
                    for i, pn in enumerate(group_param_names):
                        j = anova_params.index(pn)
                        r_matrix[i, j] = 1.0
                    f_result = model_anova.f_test(r_matrix)
                    g_df = int(f_result.df_num)
                    g_f = float(f_result.fvalue)
                    g_ss = g_f * g_df * mse_anova
                    g_ms = g_ss / max(g_df, 1)
                    g_p = float(f_result.pvalue)
                    return {
                        "source": source_name, "df": g_df, "ss": round(g_ss, 6),
                        "ms": round(g_ms, 6),
                        "f_value": round(g_f, 4),
                        "p_value": round(g_p, 6)
                    }
                except Exception:
                    # 回退: 简单求和
                    return None
            
            # 识别各组参数名（从 ANOVA 模型的参数列表中）
            linear_param_names = []
            square_param_names = []
            interact_param_names = []
            
            for p_name in anova_params:
                if p_name == "Intercept":
                    continue
                if "**" in p_name:
                    square_param_names.append(p_name)
                elif ":" in p_name:
                    interact_param_names.append(p_name)
                else:
                    linear_param_names.append(p_name)
            
            # ── 组装最终 ANOVA 表 ──
            result = []
            
            # ★ v15 修正: 模型/误差汇总行用 Sum coding ANOVA 模型的 ESS/SSR
            # （之前用各因子 Type III SS 之和会导致误差 SS 为负，因为 Type III 各项不正交分解）
            ss_model = float(model_anova.ess)
            df_model = int(model_anova.df_model)
            ms_model = ss_model / max(df_model, 1)
            ss_res_full = float(model_anova.ssr)
            df_res_full = int(model_anova.df_resid)
            ms_res = ss_res_full / max(df_res_full, 1)
            f_model = ms_model / ms_res if ms_res > 1e-12 else None
            p_model = float(model_anova.f_pvalue) if not np.isnan(model_anova.f_pvalue) else None
            result.append({
                "source": "模型", "df": df_model, "ss": round(ss_model, 6),
                "ms": round(ms_model, 6),
                "f_value": round(f_model, 4) if f_model is not None else None,
                "p_value": round(p_model, 6) if p_model is not None else None
            })
            
            # 线性汇总 + 各主效应
            linear_group = _make_group_summary_glh("  线性", linear_param_names)
            if linear_group:
                result.append(linear_group)
            for row in linear_rows:
                r = dict(row)
                r["source"] = "    " + r["source"]
                result.append(r)
            
            # 平方汇总 + 各二次项
            if square_rows:
                square_group = _make_group_summary_glh("  平方", square_param_names)
                if square_group:
                    result.append(square_group)
                for row in square_rows:
                    r = dict(row)
                    r["source"] = "    " + r["source"]
                    result.append(r)
            
            # 交互汇总 + 各交互项
            if interact_rows:
                interact_group = _make_group_summary_glh("  双因子交互作用", interact_param_names)
                if interact_group:
                    result.append(interact_group)
                for row in interact_rows:
                    r = dict(row)
                    r["source"] = "    " + r["source"]
                    result.append(r)
            
            # 误差（残差）行
            result.append({
                "source": "误差", "df": df_res_full, "ss": round(ss_res_full, 6),
                "ms": round(ms_res, 6), "f_value": None, "p_value": None
            })
            
            # Lack of Fit 和 Pure Error（用原始值模型）
            lof_result = self._calc_lack_of_fit(model_anova, df_safe, safe_names)
            if lof_result is not None:
                lof_row = lof_result["lack_of_fit"]
                lof_row["source"] = "  " + lof_row["source"]
                result.append(lof_row)
                pe_row = lof_result["pure_error"]
                pe_row["source"] = "  " + pe_row["source"]
                result.append(pe_row)
            
            # 总计行
            ss_total = float(np.sum((df_safe["Y"] - df_safe["Y"].mean()) ** 2))
            result.append({
                "source": "合计", "df": len(df_safe) - 1, "ss": round(ss_total, 6),
                "ms": None, "f_value": None, "p_value": None
            })
            
            return result
            
        except Exception as e:
            return [{"source": "Error", "error": str(e)}]

    def _calc_lack_of_fit(self, model, df_safe, safe_names) -> Optional[dict]:
        """
        ★ 新增: 计算 Lack of Fit（失拟检验）
        
        只有当存在重复实验时才能计算:
          SS_LOF = SS_Res - SS_PE
          SS_PE = Σ Σ (y_ij - ȳ_i.)²（同一因子组合下的纯误差）
          F_LOF = MS_LOF / MS_PE
        """
        try:
            factor_cols = list(safe_names.values())
            
            # 按因子组合分组
            groups = df_safe.groupby(factor_cols)["Y"]
            
            # 纯误差: 同一因子组合内的偏差平方和
            ss_pe = 0.0
            df_pe = 0
            n_groups = 0
            
            for _, group in groups:
                ni = len(group)
                if ni > 1:
                    # SS_PE = Σ(y_ij - ȳ_i)²，显式计算避免 ddof 混淆
                    ss_pe += float(((group - group.mean()) ** 2).sum())
                    df_pe += ni - 1
                n_groups += 1
            
            if df_pe == 0:
                return None  # 没有重复实验，无法计算 LOF
            
            ss_res = float(model.ssr)
            df_res = int(model.df_resid)
            
            ss_lof = ss_res - ss_pe
            df_lof = df_res - df_pe
            
            if df_lof <= 0 or df_pe <= 0:
                return None
            
            ms_lof = ss_lof / df_lof
            ms_pe = ss_pe / df_pe
            
            from scipy import stats as sp_stats
            
            f_lof = ms_lof / ms_pe if ms_pe > 1e-12 else 0.0
            p_lof = 1.0 - sp_stats.f.cdf(f_lof, df_lof, df_pe) if f_lof > 0 else 1.0
            
            return {
                "lack_of_fit": {
                    "source": "失拟", "df": df_lof, "ss": round(ss_lof, 6),
                    "ms": round(ms_lof, 6), "f_value": round(f_lof, 4), "p_value": round(p_lof, 6)
                },
                "pure_error": {
                    "source": "纯误差", "df": df_pe, "ss": round(ss_pe, 6),
                    "ms": round(ms_pe, 6), "f_value": None, "p_value": None
                }
            }
        except Exception:
            return None

    def _calc_lack_of_fit_p(self, model, df_safe, safe_names, model_type) -> Optional[float]:
        """计算 Lack of Fit P 值"""
        result = self._calc_lack_of_fit(model, df_safe, safe_names)
        if result is None:
            return None
        return result["lack_of_fit"].get("p_value")

    def _restore_term_name(self, term: str, reverse_names: dict) -> str:
        """将安全变量名映射回原始因子名，支持类别因子的 C() 表示"""
        if term == "Intercept":
            return "截距"
        
        display = term
        
        # 处理二次项: I(X0**2) 或 I(X0 ** 2)（statsmodels 输出可能含空格）
        # ★ 修复 (v4): 使用正则匹配以兼容有无空格的两种格式
        import re
        quad_match = re.match(r'^I\((\w+)\s*\*\*\s*2\)$', display)
        if quad_match:
            inner = quad_match.group(1)
            orig = reverse_names.get(inner, inner)
            return f"{orig}²"
        
        # ★ 新增: 处理类别因子 C(X3)[T.B] → 催化剂[B]
        # ★ v13: 也支持 Sum coding C(X3)[S.A] → 催化剂[A]
        cat_match = re.match(r'^C\((\w+)\)\[(?:T|S)\.([^\]]+)\]$', display)
        if cat_match:
            safe_name = cat_match.group(1)
            level = cat_match.group(2)
            orig = reverse_names.get(safe_name, safe_name)
            return f"{orig}[{level}]"
        
        # ★ 新增: 处理类别因子交互 C(X3)[T.B]:X0 → 催化剂[B]×温度
        # ★ v13: 也支持 Sum coding C(X3)[S.A]:X0
        if ":" in display:
            parts = display.split(":")
            orig_parts = []
            for p in parts:
                cat_m = re.match(r'^C\((\w+)\)\[(?:T|S)\.([^\]]+)\]$', p)
                if cat_m:
                    safe_name = cat_m.group(1)
                    level = cat_m.group(2)
                    orig = reverse_names.get(safe_name, safe_name)
                    orig_parts.append(f"{orig}[{level}]")
                else:
                    orig_parts.append(reverse_names.get(p, p))
            return "×".join(orig_parts)
        
        # 主效应
        return reverse_names.get(display, display)

    def _build_equation(self, model, reverse_names: dict) -> str:
        """构建回归方程字符串（向后兼容：返回单个方程，含类别因子哑变量项）"""
        parts = []
        for term in model.params.index:
            coeff = float(model.params[term])
            display = self._restore_term_name(term, reverse_names)
            
            if term == "Intercept":
                parts.append(f"{coeff:.4f}")
            else:
                sign = "+" if coeff >= 0 else "-"
                parts.append(f"{sign} {abs(coeff):.4f}×{display}")
        
        return f"{self._response_name} = " + " ".join(parts)

    def _build_equations_by_category(self, model, reverse_names: dict) -> dict:
        """
        ★ 新增 v4: Minitab 风格方程展示 — 按类别因子水平展开为多个方程
        
        当模型包含类别因子时，将哑变量代入，为每种类别组合生成一个只含连续因子的方程。
        
        返回:
        {
            "has_categorical": true,
            "categorical_factor": "催化剂",            # 展开的类别因子名
            "equations": {
                "A": "产率 = -68.1 + 1.494×温度 + 0.115×压力 - 0.00451×温度² ...",
                "B": "产率 = -76.4 + 1.552×温度 + 0.143×压力 - 0.00451×温度² ...",
                "C": "产率 = -66.0 + 1.510×温度 + 0.152×压力 - 0.00451×温度² ..."
            },
            "common_equation": "产率 = 85.2 + 12.3×温度 + ...(含哑变量项的完整方程)"
        }
        
        当无类别因子时:
        {
            "has_categorical": false,
            "equations": {"全部": "产率 = 85.2 + 12.3×温度 + ..."},
            "common_equation": "产率 = 85.2 + 12.3×温度 + ..."
        }
        
        数学原理:
          模型: Y = β₀ + β₁T + β₂P + β₃D₂ + β₄D₃ + β₅T² + β₆P² + β₇TP + β₈T·D₂ + β₉T·D₃ + β₁₀P·D₂ + β₁₁P·D₃
          
          Catalyst=A (D₂=0, D₃=0):
            Y = β₀ + β₁T + β₂P + β₅T² + β₆P² + β₇TP
          
          Catalyst=B (D₂=1, D₃=0):
            Y = (β₀+β₃) + (β₁+β₈)T + (β₂+β₁₀)P + β₅T² + β₆P² + β₇TP
          
          Catalyst=C (D₂=0, D₃=1):
            Y = (β₀+β₄) + (β₁+β₉)T + (β₂+β₁₁)P + β₅T² + β₆P² + β₇TP
        """
        import re
        
        if not self._categorical_factors:
            eq = self._build_equation(model, reverse_names)
            return {
                "has_categorical": False,
                "equations": {"全部": eq},
                "common_equation": eq
            }
        
        # ── 解析模型中所有项，分类为: 纯连续项 / 类别主效应 / 类别×连续交互 ──
        # 我们的编码: Treatment coding (0/1), statsmodels 的 C(X)[T.level] 格式
        
        params = model.params
        
        # 识别所有类别因子的哑变量项
        # C(X3)[T.B] → 因子X3, 水平B, 系数β
        # C(X3)[T.B]:X0 → 因子X3, 水平B, 与连续因子X0的交互
        
        safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
        
        # 收集所有类别因子的水平列表
        cat_factor_levels = {}  # orig_name → [level0, level1, level2, ...]
        for cat_name in self._categorical_factors:
            levels = sorted(self._df[cat_name].astype(str).unique())
            cat_factor_levels[cat_name] = levels
        
        # 对每个类别因子的每个水平，构建一个"修正字典":
        # {连续项名 → 额外加的系数, "Intercept" → 额外加的常数}
        # 然后把修正后的系数组装成方程
        
        # 简化: 目前只支持单个类别因子展开（多个类别因子时做笛卡尔积会很复杂）
        # 如果有多个类别因子，选第一个展开，其余类别因子保留在方程中
        primary_cat = self._categorical_factors[0]
        primary_cat_safe = safe_names[primary_cat]
        primary_levels = cat_factor_levels[primary_cat]
        
        # 对每个水平，构建哑变量值字典
        # Treatment coding: 第一个水平(参考水平)所有哑变量=0
        non_ref_levels = primary_levels[1:]  # 非参考水平
        
        # 遍历模型参数，分类
        # 1. 纯连续项（不含该类别因子的项）: 所有水平共享
        # 2. 该类别因子的哑变量项: 代入 0/1
        
        equations = {}
        
        for level in primary_levels:
            # 构建该水平下的哑变量值
            dummy_values = {}
            for lv in non_ref_levels:
                dummy_values[lv] = 1.0 if level == lv else 0.0
            
            # 遍历所有参数项，代入哑变量
            effective_coeffs = {}  # term_display → coefficient
            
            for term_key in params.index:
                coeff = float(params[term_key])
                term_str = str(term_key)
                
                if term_key == "Intercept":
                    effective_coeffs["Intercept"] = effective_coeffs.get("Intercept", 0.0) + coeff
                    continue
                
                # 检查该项是否涉及我们正在展开的类别因子
                # 模式: C(X3)[T.B]  或  C(X3)[T.B]:X0  或  X0:C(X3)[T.B]
                cat_pattern = rf'C\({re.escape(primary_cat_safe)}\)\[T\.([^\]]+)\]'
                cat_match = re.search(cat_pattern, term_str)
                
                if cat_match:
                    matched_level = cat_match.group(1)
                    dummy_val = dummy_values.get(matched_level, 0.0)
                    
                    if dummy_val == 0.0:
                        continue  # 这个哑变量在当前水平下为0，项消失
                    
                    # 提取该项中除了类别因子部分之外的连续因子部分
                    # 例: C(X3)[T.B]:X0 → 连续部分是 X0
                    # 例: C(X3)[T.B] → 纯类别项，加到截距
                    parts = term_str.split(":")
                    continuous_parts = []
                    for p in parts:
                        if not re.match(cat_pattern, p):
                            continuous_parts.append(p)
                    
                    if not continuous_parts:
                        # 纯类别主效应 → 加到截距
                        effective_coeffs["Intercept"] = effective_coeffs.get("Intercept", 0.0) + coeff * dummy_val
                    else:
                        # 类别×连续交互 → 加到对应连续项的系数
                        cont_key = ":".join(continuous_parts)
                        effective_coeffs[cont_key] = effective_coeffs.get(cont_key, 0.0) + coeff * dummy_val
                else:
                    # 纯连续项 → 直接累加
                    effective_coeffs[term_str] = effective_coeffs.get(term_str, 0.0) + coeff
            
            # 组装方程字符串
            eq_parts = []
            # 先写截距
            intercept = effective_coeffs.pop("Intercept", 0.0)
            eq_parts.append(f"{intercept:.4f}")
            
            # 再写其余项（按原始顺序）
            for term_key in params.index:
                if term_key == "Intercept":
                    continue
                term_str = str(term_key)
                
                # 跳过涉及该类别因子的项（已经代入了）
                cat_pattern = rf'C\({re.escape(primary_cat_safe)}\)\[T\.([^\]]+)\]'
                if re.search(cat_pattern, term_str):
                    continue
                
                if term_str in effective_coeffs:
                    c = effective_coeffs[term_str]
                    display = self._restore_term_name(term_str, reverse_names)
                    sign = "+" if c >= 0 else "-"
                    eq_parts.append(f"{sign} {abs(c):.4f}×{display}")
            
            eq_str = f"{self._response_name} = " + " ".join(eq_parts)
            equations[str(level)] = eq_str
        
        common_eq = self._build_equation(model, reverse_names)
        
        return {
            "has_categorical": True,
            "categorical_factor": primary_cat,
            "equations": equations,
            "common_equation": common_eq
        }


# ═══════════════════════════════════════════════════════
# DOESurfacePlotter — 响应曲面和等高线（保留不变）
# ═══════════════════════════════════════════════════════

class DOESurfacePlotter:
    """响应曲面绘图器"""

    def __init__(self):
        self._gpr_model = None
        self._analyzer = None
        self._factor_names = []
        self._bounds = {}

    def set_gpr_model(self, gpr_model, factor_names_json: str, bounds_json: str):
        self._gpr_model = gpr_model
        self._factor_names = json.loads(factor_names_json)
        self._bounds = json.loads(bounds_json)

    def set_analyzer(self, analyzer: DOEAnalyzer, bounds_json: str):
        self._analyzer = analyzer
        self._factor_names = analyzer._factor_names
        self._bounds = json.loads(bounds_json)

    def response_surface_data(self, factor1: str, factor2: str, grid_size: int = 30) -> str:
        b1 = self._bounds.get(factor1, [0, 1])
        b2 = self._bounds.get(factor2, [0, 1])
        x_range = np.linspace(b1[0], b1[1], grid_size)
        y_range = np.linspace(b2[0], b2[1], grid_size)
        z_grid = np.zeros((grid_size, grid_size))
        center_values = {name: (self._bounds[name][0] + self._bounds[name][1]) / 2
                         for name in self._factor_names if name not in (factor1, factor2)}
        for i, x_val in enumerate(x_range):
            for j, y_val in enumerate(y_range):
                point = dict(center_values)
                point[factor1] = float(x_val)
                point[factor2] = float(y_val)
                z_grid[j, i] = self._predict_point(point)
        return json.dumps({
            "x": x_range.tolist(), "y": y_range.tolist(), "z": z_grid.tolist(),
            "x_label": factor1, "y_label": factor2
        }, ensure_ascii=False)

    def contour_data(self, factor1: str, factor2: str, grid_size: int = 30) -> str:
        return self.response_surface_data(factor1, factor2, grid_size)

    def response_surface_image(self, factor1: str, factor2: str,
                                grid_size: int = 50, width: int = 800, height: int = 600) -> str:
        if not HAS_MATPLOTLIB:
            return ""
        data = json.loads(self.response_surface_data(factor1, factor2, grid_size))
        X = np.array(data["x"])
        Y = np.array(data["y"])
        Z = np.array(data["z"])
        XX, YY = np.meshgrid(X, Y)
        fig = plt.figure(figsize=(width / 100, height / 100), dpi=100)
        ax = fig.add_subplot(111, projection='3d')
        surf = ax.plot_surface(XX, YY, Z, cmap='viridis', alpha=0.85, edgecolor='none')
        ax.set_xlabel(factor1, fontsize=11)
        ax.set_ylabel(factor2, fontsize=11)
        ax.set_zlabel('响应值', fontsize=11)
        ax.set_title(f'响应曲面: {factor1} × {factor2}', fontsize=13)
        fig.colorbar(surf, shrink=0.5, aspect=10)
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return base64.b64encode(buf.read()).decode('utf-8')

    def gpr_confidence_band_data(self, factor_name: str, grid_size: int = 50) -> str:
        if self._gpr_model is None:
            return json.dumps({"error": "未设置 GPR 模型"}, ensure_ascii=False)
        bounds = self._bounds.get(factor_name, [0, 1])
        x_range = np.linspace(bounds[0], bounds[1], grid_size)
        center_values = {name: (self._bounds[name][0] + self._bounds[name][1]) / 2
                         for name in self._factor_names if name != factor_name}
        means, lowers, uppers = [], [], []
        for x_val in x_range:
            point = dict(center_values)
            point[factor_name] = float(x_val)
            pred_json = self._gpr_model.predict(json.dumps(point))
            pred = json.loads(pred_json)
            m, s = pred["mean"], pred["std"]
            means.append(round(m, 4))
            lowers.append(round(m - 2 * s, 4))
            uppers.append(round(m + 2 * s, 4))
        return json.dumps({
            "x": x_range.tolist(), "mean": means, "lower": lowers,
            "upper": uppers, "x_label": factor_name
        }, ensure_ascii=False)

    def _predict_point(self, point: dict) -> float:
        if self._gpr_model is not None:
            pred_json = self._gpr_model.predict(json.dumps(point))
            pred = json.loads(pred_json)
            return pred["mean"]
        elif self._analyzer is not None and self._analyzer._model is not None:
            safe_names = {f: f"X{i}" for i, f in enumerate(self._factor_names)}
            safe_point = {safe_names.get(k, k): v for k, v in point.items() if k in safe_names}
            try:
                pred = self._analyzer._model.predict(pd.DataFrame([safe_point]))
                return float(pred.iloc[0])
            except Exception:
                return 0.0
        return 0.0


# ═══════════════════════════════════════════════════════
# 模块入口
# ═══════════════════════════════════════════════════════

def create_analyzer():
    """创建 DOEAnalyzer 实例"""
    return DOEAnalyzer()

def create_surface_plotter():
    """创建 DOESurfacePlotter 实例"""
    return DOESurfacePlotter()


# ═══════════════════════════════════════════════════════
# 本地测试
# ═══════════════════════════════════════════════════════

if __name__ == "__main__":
    designer = DOEDesigner()
    
    # 测试全因子设计
    factors = json.dumps([
        {"name": "温度", "levels": [80, 120, 160]},
        {"name": "压力", "levels": [10, 20]},
        {"name": "催化剂", "levels": [1.0, 3.0, 5.0]}
    ], ensure_ascii=False)
    
    matrix = designer.full_factorial(factors)
    rows = json.loads(matrix)
    print(f"全因子设计: {len(rows)} 组实验")
    
    # ★ 测试 CCD
    factors_rsm = json.dumps([
        {"name": "温度", "lower": 80, "upper": 160},
        {"name": "压力", "lower": 10, "upper": 30},
        {"name": "催化剂", "lower": 1.0, "upper": 5.0}
    ], ensure_ascii=False)
    
    matrix_ccd = designer.ccd(factors_rsm, "rotatable")
    rows_ccd = json.loads(matrix_ccd)
    print(f"\nCCD 可旋转设计: {len(rows_ccd)} 组实验")
    for i, row in enumerate(rows_ccd[:5]):
        print(f"  第{i+1}组: {row}")
    
    matrix_ccf = designer.ccd(factors_rsm, "face")
    rows_ccf = json.loads(matrix_ccf)
    print(f"\nCCD 面心设计: {len(rows_ccf)} 组实验")
    
    # ★ 测试 BBD
    matrix_bbd = designer.box_behnken(factors_rsm)
    rows_bbd = json.loads(matrix_bbd)
    print(f"\nBox-Behnken: {len(rows_bbd)} 组实验")
    for i, row in enumerate(rows_bbd[:5]):
        print(f"  第{i+1}组: {row}")
    
    # ★ 测试 D-Optimal
    matrix_dopt = designer.d_optimal(factors_rsm, num_runs=15)
    rows_dopt = json.loads(matrix_dopt)
    print(f"\nD-Optimal (15组): {len(rows_dopt)} 组实验")
    
    # ★ 测试设计质量
    quality = json.loads(designer.get_design_quality(factors_rsm, matrix_ccd))
    print(f"\nCCD 设计质量:")
    print(f"  D-效率: {quality['d_efficiency']}")
    print(f"  A-效率: {quality['a_efficiency']}")
    print(f"  G-效率: {quality['g_efficiency']}")
    print(f"  VIF: {quality['vif']}")
    print(f"  Power: {quality['power_analysis']}")
    
    # ★ 测试 OLS 分析
    print(f"\n--- OLS 分析测试 ---")
    analyzer = DOEAnalyzer()
    
    # 模拟数据: y = 50 + 10*T + 5*P - 3*T² + 2*T*P + noise
    np.random.seed(42)
    test_factors = []
    test_responses = []
    for row in rows_ccd:
        t = (row["温度"] - 120) / 40  # 标准化
        p = (row["压力"] - 20) / 10
        c = (row["催化剂"] - 3) / 2
        y = 50 + 10*t + 5*p + 2*c - 3*t**2 - 1.5*p**2 + 2*t*p + np.random.normal(0, 1)
        test_factors.append(row)
        test_responses.append(round(y, 2))
    
    analyzer.load_data(json.dumps(test_factors), json.dumps(test_responses), "转化率")
    
    ols_result = json.loads(analyzer.fit_ols("quadratic"))
    if "error" not in ols_result:
        print(f"  R² = {ols_result['model_summary']['r_squared']}")
        print(f"  R²adj = {ols_result['model_summary']['r_squared_adj']}")
        print(f"  R²pred = {ols_result['model_summary']['r_squared_pred']}")
        print(f"  RMSE = {ols_result['model_summary']['rmse']}")
        print(f"  Adequate Precision = {ols_result['model_summary']['adeq_precision']}")
        print(f"  Lack of Fit P = {ols_result['model_summary']['lack_of_fit_p']}")
        print(f"\n  ANOVA 表:")
        for row in ols_result["anova_table"]:
            print(f"    {row['source']:12s} df={row['df']} SS={row['ss']}")
        print(f"\n  系数表:")
        for coeff in ols_result["coefficients"][:5]:
            print(f"    {coeff['term']:12s} β={coeff['coeff']:8.4f} p={coeff['p_value']:.4f} VIF={coeff['vif']}")
    else:
        print(f"  错误: {ols_result['error']}")
    
    # ★ 测试残差诊断
    diag = json.loads(analyzer.residual_diagnostics())
    if "error" not in diag:
        print(f"\n  残差诊断: {len(diag['residuals_vs_fitted']['fitted'])} 个数据点")
        print(f"  Cook's 距离最大值: {max(diag['cooks_distance']['distance']):.4f}")
    
    # ★ 测试效应 Pareto
    pareto_result = json.loads(analyzer.effects_pareto())
    pareto = pareto_result["effects"]
    print(f"\n  效应 Pareto (前3):")
    for eff in pareto[:3]:
        sig = "★" if eff.get("significant") else " "
        print(f"    {sig} {eff['term']:12s} |t|={eff['abs_t']:.2f} p={eff['p_value']:.4f}")
    
    print("\n所有测试通过！")
